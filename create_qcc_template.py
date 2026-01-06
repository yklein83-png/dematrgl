"""
Script pour creer le template QCC (Questionnaire Connaissance Client) complet
Base sur les exigences AMF/ACPR - 120+ champs
"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Configuration des styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

# ============================================
# EN-TETE
# ============================================
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run("QUESTIONNAIRE DE CONNAISSANCE CLIENT")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Document conforme aux exigences AMF/ACPR").italic = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Date : ").bold = True
p.add_run("$DATE_JOUR$")

p = doc.add_paragraph()
p.add_run("Conseiller : ").bold = True
p.add_run("$TITRE_CONSEILLER$ $PRENOM_CONSEILLER$ $NOM_CONSEILLER$")

doc.add_paragraph()

# ============================================
# SECTION 1 : IDENTITE TITULAIRE 1
# ============================================
def add_section_title(title):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)

def add_field_table(fields):
    """Ajoute un tableau de champs label: valeur"""
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    for i, (label, placeholder) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = placeholder
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        # Largeur colonnes
        table.rows[i].cells[0].width = Cm(6)
        table.rows[i].cells[1].width = Cm(10)
    doc.add_paragraph()

add_section_title("1. IDENTITE TITULAIRE 1 (PRINCIPAL)")

fields_t1 = [
    ("Civilite", "$T1_CIVILITE$"),
    ("Nom", "$T1_NOM$"),
    ("Nom de jeune fille", "$T1_NOM_JEUNE_FILLE$"),
    ("Prenom", "$T1_PRENOM$"),
    ("Date de naissance", "$T1_DATE_NAISSANCE$"),
    ("Lieu de naissance", "$T1_LIEU_NAISSANCE$"),
    ("Nationalite", "$T1_NATIONALITE$"),
    ("Adresse", "$T1_ADRESSE$"),
    ("Email", "$T1_EMAIL$"),
    ("Telephone", "$T1_TELEPHONE$"),
    ("Residence fiscale", "$T1_RESIDENCE_FISCALE$"),
    ("US Person", "$T1_US_PERSON$"),
    ("Regime de protection juridique", "$T1_REGIME_PROTECTION$"),
    ("Representant legal (si applicable)", "$T1_REPRESENTANT_LEGAL$"),
]
add_field_table(fields_t1)

# Situation professionnelle T1
p = doc.add_paragraph()
p.add_run("Situation professionnelle :").bold = True

fields_pro_t1 = [
    ("Profession actuelle", "$T1_PROFESSION$"),
    ("Retraite depuis le", "$T1_RETRAITE_DEPUIS$"),
    ("Chomage depuis le", "$T1_CHOMAGE_DEPUIS$"),
    ("Ancienne profession", "$T1_ANCIENNE_PROFESSION$"),
    ("Chef d'entreprise", "$T1_CHEF_ENTREPRISE$"),
    ("Denomination entreprise", "$T1_ENTREPRISE_DENOMINATION$"),
    ("Forme juridique", "$T1_ENTREPRISE_FORME_JURIDIQUE$"),
    ("Siege social", "$T1_ENTREPRISE_SIEGE_SOCIAL$"),
]
add_field_table(fields_pro_t1)

# ============================================
# SECTION 2 : IDENTITE TITULAIRE 2 (Optionnel)
# ============================================
add_section_title("2. IDENTITE TITULAIRE 2 (SI APPLICABLE)")

fields_t2 = [
    ("Civilite", "$T2_CIVILITE$"),
    ("Nom", "$T2_NOM$"),
    ("Nom de jeune fille", "$T2_NOM_JEUNE_FILLE$"),
    ("Prenom", "$T2_PRENOM$"),
    ("Date de naissance", "$T2_DATE_NAISSANCE$"),
    ("Lieu de naissance", "$T2_LIEU_NAISSANCE$"),
    ("Nationalite", "$T2_NATIONALITE$"),
    ("Adresse", "$T2_ADRESSE$"),
    ("Email", "$T2_EMAIL$"),
    ("Telephone", "$T2_TELEPHONE$"),
    ("Residence fiscale", "$T2_RESIDENCE_FISCALE$"),
    ("US Person", "$T2_US_PERSON$"),
    ("Profession", "$T2_PROFESSION$"),
]
add_field_table(fields_t2)

# ============================================
# SECTION 3 : SITUATION FAMILIALE
# ============================================
add_section_title("3. SITUATION FAMILIALE")

fields_famille = [
    ("Situation familiale", "$SITUATION_FAMILIALE$"),
    ("Date de mariage", "$DATE_MARIAGE$"),
    ("Contrat de mariage", "$CONTRAT_MARIAGE$"),
    ("Regime matrimonial", "$REGIME_MATRIMONIAL$"),
    ("Date PACS", "$DATE_PACS$"),
    ("Convention PACS", "$CONVENTION_PACS$"),
    ("Regime PACS", "$REGIME_PACS$"),
    ("Donation entre epoux", "$DONATION_ENTRE_EPOUX$"),
    ("Date donation epoux", "$DONATION_EPOUX_DATE$"),
    ("Montant donation epoux", "$DONATION_EPOUX_MONTANT$"),
    ("Donation aux enfants", "$DONATION_ENFANTS$"),
    ("Nombre d'enfants", "$NOMBRE_ENFANTS$"),
    ("Nombre d'enfants a charge", "$NOMBRE_ENFANTS_CHARGE$"),
]
add_field_table(fields_famille)

# ============================================
# SECTION 4 : SITUATION FINANCIERE
# ============================================
add_section_title("4. SITUATION FINANCIERE")

fields_finance = [
    ("Revenus annuels du foyer", "$REVENUS_ANNUELS$"),
    ("Patrimoine global", "$PATRIMOINE_GLOBAL$"),
    ("Charges annuelles (%)", "$CHARGES_POURCENT$"),
    ("Charges annuelles (montant)", "$CHARGES_MONTANT$"),
    ("Capacite d'epargne mensuelle", "$CAPACITE_EPARGNE$"),
    ("Assujetti a l'IR", "$IMPOT_REVENU$"),
    ("Assujetti a l'IFI", "$IMPOT_FORTUNE$"),
]
add_field_table(fields_finance)

# Repartition patrimoine
p = doc.add_paragraph()
p.add_run("Repartition du patrimoine :").bold = True

fields_repartition = [
    ("Patrimoine financier (%)", "$PATRIMOINE_FIN_PCT$"),
    ("Patrimoine immobilier (%)", "$PATRIMOINE_IMMO_PCT$"),
    ("Patrimoine professionnel (%)", "$PATRIMOINE_PRO_PCT$"),
    ("Autres actifs (%)", "$PATRIMOINE_AUTRES_PCT$"),
]
add_field_table(fields_repartition)

# ============================================
# SECTION 5 : ORIGINE DES FONDS
# ============================================
add_section_title("5. ORIGINE DES FONDS")

fields_origine = [
    ("Nature des fonds", "$ORIGINE_NATURE$"),
    ("Montant prevu", "$ORIGINE_MONTANT$"),
    ("Revenus professionnels", "$ORIGINE_REVENUS$"),
    ("Epargne", "$ORIGINE_EPARGNE$"),
    ("Heritage/Donation", "$ORIGINE_HERITAGE$"),
    ("Cession professionnelle", "$ORIGINE_CESSION_PRO$"),
    ("Cession immobiliere", "$ORIGINE_CESSION_IMMO$"),
    ("Cession mobiliere", "$ORIGINE_CESSION_MOBILIERE$"),
    ("Gains de jeu", "$ORIGINE_GAINS_JEU$"),
    ("Rachat assurance-vie", "$ORIGINE_ASSURANCE_VIE$"),
    ("Autres (preciser)", "$ORIGINE_AUTRES$"),
    ("Etablissement de provenance", "$ORIGINE_ETABLISSEMENT$"),
]
add_field_table(fields_origine)

# ============================================
# SECTION 6 : CONNAISSANCE ET EXPERIENCE (KYC)
# ============================================
add_section_title("6. CONNAISSANCE ET EXPERIENCE EN INVESTISSEMENTS")

doc.add_paragraph("Pour chaque type de produit, indiquez votre experience :")
doc.add_paragraph()

# Tableau KYC produits
table_kyc = doc.add_table(rows=10, cols=6)
table_kyc.style = 'Table Grid'

# En-tetes
headers = ["Produit", "Detention", "Nb operations/an", "Duree experience", "Volume moyen", "Connaissance"]
for i, h in enumerate(headers):
    table_kyc.rows[0].cells[i].text = h
    table_kyc.rows[0].cells[i].paragraphs[0].runs[0].bold = True

# Produits
produits = [
    ("Monetaires", "$KYC_MONETAIRES_DETENTION$", "$KYC_MONETAIRES_OPERATIONS$", "$KYC_MONETAIRES_DUREE$", "$KYC_MONETAIRES_VOLUME$", "$KYC_MONETAIRES_Q1$"),
    ("Obligations", "$KYC_OBLIGATIONS_DETENTION$", "$KYC_OBLIGATIONS_OPERATIONS$", "$KYC_OBLIGATIONS_DUREE$", "$KYC_OBLIGATIONS_VOLUME$", "$KYC_OBLIGATIONS_Q1$"),
    ("Actions", "$KYC_ACTIONS_DETENTION$", "$KYC_ACTIONS_OPERATIONS$", "$KYC_ACTIONS_DUREE$", "$KYC_ACTIONS_VOLUME$", "$KYC_ACTIONS_Q1$"),
    ("SCPI/OPCI", "$KYC_SCPI_DETENTION$", "$KYC_SCPI_OPERATIONS$", "$KYC_SCPI_DUREE$", "$KYC_SCPI_VOLUME$", "$KYC_SCPI_Q1$"),
    ("Private Equity", "$KYC_PE_DETENTION$", "$KYC_PE_OPERATIONS$", "$KYC_PE_DUREE$", "$KYC_PE_VOLUME$", "$KYC_PE_Q1$"),
    ("ETF/Trackers", "$KYC_ETF_DETENTION$", "$KYC_ETF_OPERATIONS$", "$KYC_ETF_DUREE$", "$KYC_ETF_VOLUME$", "$KYC_ETF_Q1$"),
    ("Produits derives", "$KYC_DERIVES_DETENTION$", "$KYC_DERIVES_OPERATIONS$", "$KYC_DERIVES_DUREE$", "$KYC_DERIVES_VOLUME$", "$KYC_DERIVES_Q1$"),
    ("Produits structures", "$KYC_STRUCTURES_DETENTION$", "$KYC_STRUCTURES_OPERATIONS$", "$KYC_STRUCTURES_DUREE$", "$KYC_STRUCTURES_VOLUME$", "$KYC_STRUCTURES_Q1$"),
    ("Assurance-vie UC", "$KYC_ASSURANCE_VIE_DETENTION$", "$KYC_ASSURANCE_VIE_OPERATIONS$", "$KYC_ASSURANCE_VIE_DUREE$", "$KYC_ASSURANCE_VIE_VOLUME$", "$KYC_ASSURANCE_VIE_Q1$"),
]

for i, (produit, *placeholders) in enumerate(produits, start=1):
    table_kyc.rows[i].cells[0].text = produit
    for j, ph in enumerate(placeholders):
        table_kyc.rows[i].cells[j+1].text = ph

doc.add_paragraph()

# Gestion portefeuille
p = doc.add_paragraph()
p.add_run("Mode de gestion :").bold = True

fields_gestion = [
    ("Mandat de gestion", "$KYC_MANDAT_GESTION$"),
    ("Gestion personnelle", "$KYC_GESTION_PERSONNELLE$"),
    ("Gestion avec conseiller", "$KYC_GESTION_CONSEILLER$"),
    ("Experience professionnelle finance", "$KYC_EXPERIENCE_PRO$"),
]
add_field_table(fields_gestion)

# Culture financiere
p = doc.add_paragraph()
p.add_run("Culture financiere :").bold = True

fields_culture = [
    ("Lecture presse financiere", "$KYC_PRESSE_FINANCIERE$"),
    ("Suivi de la bourse", "$KYC_SUIVI_BOURSE$"),
    ("Lecture reguliere releves", "$KYC_RELEVES_BANCAIRES$"),
]
add_field_table(fields_culture)

# ============================================
# SECTION 7 : PROFIL DE RISQUE
# ============================================
add_section_title("7. PROFIL DE RISQUE")

fields_risque = [
    ("Objectifs d'investissement", "$OBJECTIFS$"),
    ("Horizon de placement", "$HORIZON$"),
    ("Tolerance au risque", "$TOLERANCE_RISQUE$"),
    ("Pertes maximales acceptables", "$PERTES_MAX$"),
    ("Experience de perte en capital", "$EXPERIENCE_PERTE$"),
    ("Niveau de perte subi", "$EXPERIENCE_PERTE_NIVEAU$"),
    ("Reaction en cas de baisse", "$REACTION_PERTE$"),
    ("Reaction en cas de hausse", "$REACTION_GAIN$"),
    ("Importance de la liquidite", "$LIQUIDITE_IMPORTANTE$"),
    ("% du patrimoine a investir", "$POURCENTAGE_PATRIMOINE$"),
]
add_field_table(fields_risque)

# Profil calcule
p = doc.add_paragraph()
run = p.add_run("PROFIL DE RISQUE DETERMINE : ")
run.bold = True
p.add_run("$PROFIL_RISQUE$")

p = doc.add_paragraph()
p.add_run("Score : ").bold = True
p.add_run("$PROFIL_SCORE$ / 100")

p = doc.add_paragraph()
p.add_run("Date de calcul : ").bold = True
p.add_run("$PROFIL_DATE_CALCUL$")

# ============================================
# SECTION 8 : PREFERENCES DURABILITE (ESG)
# ============================================
add_section_title("8. PREFERENCES EN MATIERE DE DURABILITE (ESG)")

fields_esg = [
    ("Souhait d'investissement durable", "$DURABILITE_SOUHAIT$"),
    ("Niveau de preference ESG", "$DURABILITE_NIVEAU$"),
    ("Importance Environnement (1-10)", "$DURABILITE_ENV$"),
    ("Importance Social (1-10)", "$DURABILITE_SOCIAL$"),
    ("Importance Gouvernance (1-10)", "$DURABILITE_GOUV$"),
    ("Investissement a impact", "$DURABILITE_IMPACT$"),
    ("Investissement solidaire", "$DURABILITE_SOLIDAIRE$"),
    ("% minimum taxonomie UE", "$DURABILITE_TAXONOMIE$"),
    ("Prise en compte des PAI", "$DURABILITE_PAI$"),
]
add_field_table(fields_esg)

p = doc.add_paragraph()
p.add_run("Secteurs exclus : ").bold = True
p.add_run("$DURABILITE_EXCLUSIONS$")

# ============================================
# SECTION 9 : LCB-FT
# ============================================
add_section_title("9. LUTTE CONTRE LE BLANCHIMENT (LCB-FT)")

fields_lcb = [
    ("Niveau de risque LCB-FT", "$LCB_NIVEAU_RISQUE$"),
    ("Personne Politiquement Exposee (PPE)", "$LCB_PPE$"),
    ("Fonction PPE", "$LCB_PPE_FONCTION$"),
    ("Famille de PPE", "$LCB_PPE_FAMILLE$"),
    ("Verification gel des avoirs", "$LCB_GEL_AVOIRS$"),
    ("Date verification", "$LCB_GEL_DATE$"),
]
add_field_table(fields_lcb)

# ============================================
# SIGNATURES
# ============================================
doc.add_paragraph()
doc.add_paragraph()
add_section_title("ATTESTATION ET SIGNATURES")

doc.add_paragraph(
    "Je/Nous certifie(ons) que les informations fournies dans ce questionnaire sont exactes et completes. "
    "Je/Nous m'engage(ons) a informer le cabinet de tout changement significatif dans ma/notre situation."
)

doc.add_paragraph()
doc.add_paragraph()

# Tableau signatures
table_sig = doc.add_table(rows=4, cols=2)
table_sig.rows[0].cells[0].text = "Fait a Papeete, le $DATE_JOUR$"
table_sig.rows[0].cells[1].text = ""

table_sig.rows[1].cells[0].text = ""
table_sig.rows[1].cells[1].text = ""

table_sig.rows[2].cells[0].text = "Signature du/des client(s)"
table_sig.rows[2].cells[1].text = "Signature du conseiller"
table_sig.rows[2].cells[0].paragraphs[0].runs[0].bold = True
table_sig.rows[2].cells[1].paragraphs[0].runs[0].bold = True

table_sig.rows[3].cells[0].text = "$T1_CIVILITE$ $T1_PRENOM$ $T1_NOM$\n$T2_CIVILITE$ $T2_PRENOM$ $T2_NOM$"
table_sig.rows[3].cells[1].text = "$TITRE_CONSEILLER$ $PRENOM_CONSEILLER$ $NOM_CONSEILLER$"

# Sauvegarder
doc.save('/app/templates/QCC_TEMPLATE.docx')
print("Template QCC cree avec succes!")

# Afficher les placeholders
import re
doc2 = Document('/app/templates/QCC_TEMPLATE.docx')
placeholders = set()
for para in doc2.paragraphs:
    found = re.findall(r'\$[A-Z0-9_]+\$', para.text)
    placeholders.update(found)
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                found = re.findall(r'\$[A-Z0-9_]+\$', para.text)
                placeholders.update(found)

print(f"\nNombre de placeholders: {len(placeholders)}")
print("\nPlaceholders dans le template:")
for p in sorted(placeholders):
    print(f"  {p}")

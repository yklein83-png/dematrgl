"""
Script pour créer le template DER à partir du document original
"""
from docx import Document
import shutil
import os

# Copier le document original
src = '/app/templates/DER_TEMPLATE_ORIGINAL.docx'
dst = '/app/templates/DER_TEMPLATE.docx'

# Charger et modifier
doc = Document(src)

# Fonction pour remplacer dans les paragraphes
def replace_in_paragraphs(paragraphs, replacements):
    for para in paragraphs:
        for old, new in replacements.items():
            if old in para.text:
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)

# Fonction pour remplacer dans les tableaux
def replace_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            for run in para.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

# Remplacements à effectuer - ces textes seront remplacés par des placeholders
replacements = {
    # Le document contient déjà des placeholders comme $TITRE_CONSEILLER$ etc.
    # On va juste s'assurer qu'ils sont corrects
}

# Le document original contient déjà les placeholders corrects
# On le copie simplement comme template
shutil.copy(src, dst)

print(f"Template DER copié vers {dst}")

# Vérifier le contenu
doc = Document(dst)
placeholders_found = []
for para in doc.paragraphs:
    text = para.text
    # Chercher les placeholders $...$
    import re
    found = re.findall(r'\$[A-Z_]+\$', text)
    placeholders_found.extend(found)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                found = re.findall(r'\$[A-Z_]+\$', para.text)
                placeholders_found.extend(found)

print(f"\nPlaceholders trouvés dans le template:")
for p in set(placeholders_found):
    print(f"  - {p}")

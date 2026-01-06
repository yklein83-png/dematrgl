"""Script pour créer les templates DOCX"""
from docx import Document
import os

templates_dir = '/app/templates'
os.makedirs(templates_dir, exist_ok=True)

# DER TEMPLATE
doc = Document()
doc.add_heading('DOCUMENT D\'ENTREE EN RELATION', 0)
doc.add_paragraph('')
doc.add_paragraph('Date: $DATE_JOUR$')
doc.add_paragraph('')
doc.add_heading('1. INFORMATIONS CLIENT', level=1)
doc.add_paragraph('$TITRE_CONTACT$ $PRENOM_CONTACT$ $NOM_CONTACT$')
doc.add_paragraph('')
doc.add_heading('2. CABINET CONSEILLER', level=1)
doc.add_paragraph('$CABINET_NOM$')
doc.add_paragraph('N ORIAS: $CABINET_ORIAS$')
doc.add_paragraph('Adresse: $CABINET_ADRESSE$')
doc.add_paragraph('Email: $CABINET_EMAIL$')
doc.add_paragraph('Tel: $CABINET_TEL$')
doc.add_paragraph('')
doc.add_heading('3. CONSEILLER', level=1)
doc.add_paragraph('$TITRE_CONSEILLER$ $PRENOM_CONSEILLER$ $NOM_CONSEILLER$')
doc.add_paragraph('')
doc.add_paragraph('Fait a Papeete, le $DATE_JOUR_JJJ_MM_AAAA$')
doc.add_paragraph('')
doc.add_paragraph('Signature Client                    Signature Conseiller')
doc.save(f'{templates_dir}/DER_TEMPLATE.docx')
print('DER_TEMPLATE.docx cree')

# KYC TEMPLATE
doc = Document()
doc.add_heading('QUESTIONNAIRE DE CONNAISSANCE CLIENT (KYC)', 0)
doc.add_paragraph('Date: $DATE_JOUR$ - Lieu: $LIEU$')
doc.add_paragraph('')

doc.add_heading('1. IDENTITE TITULAIRE 1', level=1)
table = doc.add_table(rows=7, cols=2)
table.style = 'Table Grid'
cells = [
    ('Civilite', '$T1_CIVILITE$'),
    ('Nom', '$T1_NOM$'),
    ('Nom de jeune fille', '$T1_NOM_JF$'),
    ('Prenom', '$T1_PRENOM$'),
    ('Date de naissance', '$T1_DATE_NAISSANCE$'),
    ('Lieu de naissance', '$T1_LIEU_NAISSANCE$'),
    ('Nationalite', '$T1_NATIONALITE$'),
]
for i, (label, value) in enumerate(cells):
    table.rows[i].cells[0].text = label
    table.rows[i].cells[1].text = value

doc.add_paragraph('')
table2 = doc.add_table(rows=4, cols=2)
table2.style = 'Table Grid'
cells2 = [
    ('Adresse', '$T1_ADRESSE$'),
    ('Email', '$T1_EMAIL$'),
    ('Telephone', '$T1_TELEPHONE$'),
    ('Profession', '$T1_PROFESSION$'),
]
for i, (label, value) in enumerate(cells2):
    table2.rows[i].cells[0].text = label
    table2.rows[i].cells[1].text = value

doc.add_paragraph('US Person: $T1_US_PERSON$')
doc.add_paragraph('')

doc.add_heading('2. IDENTITE TITULAIRE 2 (si applicable)', level=1)
table3 = doc.add_table(rows=6, cols=2)
table3.style = 'Table Grid'
cells3 = [
    ('Civilite', '$T2_CIVILITE$'),
    ('Nom', '$T2_NOM$'),
    ('Prenom', '$T2_PRENOM$'),
    ('Date de naissance', '$T2_DATE_NAISSANCE$'),
    ('Nationalite', '$T2_NATIONALITE$'),
    ('Profession', '$T2_PROFESSION$'),
]
for i, (label, value) in enumerate(cells3):
    table3.rows[i].cells[0].text = label
    table3.rows[i].cells[1].text = value
doc.add_paragraph('')

doc.add_heading('3. SITUATION FAMILIALE', level=1)
doc.add_paragraph('Situation: $SITUATION_FAMILIALE$')
doc.add_paragraph('Date mariage: $DATE_MARIAGE$ | Contrat: $CONTRAT_MARIAGE$')
doc.add_paragraph('Regime matrimonial: $REGIME_MATRIMONIAL$')
doc.add_paragraph('Nombre enfants: $NOMBRE_ENFANTS$ | A charge: $NOMBRE_ENFANTS_CHARGE$')
doc.add_paragraph('')

doc.add_heading('4. SITUATION FINANCIERE', level=1)
doc.add_paragraph('Revenus annuels foyer: $REVENUS_ANNUELS$')
doc.add_paragraph('Patrimoine global: $PATRIMOINE_GLOBAL$')
doc.add_paragraph('Charges: $CHARGES_POURCENT$ soit $CHARGES_MONTANT$')
doc.add_paragraph('Capacite epargne mensuelle: $CAPACITE_EPARGNE$')
doc.add_paragraph('IR: $IR$ | IFI: $IFI$')
doc.add_paragraph('')

doc.add_heading('5. REPARTITION DU PATRIMOINE', level=1)
doc.add_paragraph('Financier: $PATRIMOINE_FIN_PCT$')
doc.add_paragraph('Immobilier: $PATRIMOINE_IMMO_PCT$')
doc.add_paragraph('Professionnel: $PATRIMOINE_PRO_PCT$')
doc.add_paragraph('Autres: $PATRIMOINE_AUTRES_PCT$')
doc.add_paragraph('')

doc.add_heading('6. ORIGINE DES FONDS', level=1)
doc.add_paragraph('Nature: $ORIGINE_NATURE$')
doc.add_paragraph('Montant prevu: $ORIGINE_MONTANT$')
doc.add_paragraph('Sources: Revenus $ORIGINE_REVENUS$ | Epargne $ORIGINE_EPARGNE$ | Heritage $ORIGINE_HERITAGE$')
doc.add_paragraph('Etablissement: $ORIGINE_ETABLISSEMENT$')
doc.add_paragraph('')

doc.add_heading('7. PROFIL DE RISQUE', level=1)
doc.add_paragraph('Objectifs: $OBJECTIFS$')
doc.add_paragraph('Horizon de placement: $HORIZON$')
doc.add_paragraph('Tolerance au risque: $TOLERANCE_RISQUE$')
doc.add_paragraph('Pertes maximales acceptables: $PERTES_MAX$')
doc.add_paragraph('PROFIL CALCULE: $PROFIL_RISQUE$ (Score: $PROFIL_SCORE$)')
doc.add_paragraph('')

doc.add_heading('8. DURABILITE ESG', level=1)
doc.add_paragraph('Souhait investissement durable: $DURABILITE_SOUHAIT$')
doc.add_paragraph('Part taxonomie europeenne: $DURABILITE_TAXONOMIE$')
doc.add_paragraph('')

doc.add_paragraph('Signatures:')
doc.add_paragraph('')
doc.add_paragraph('Client                              Conseiller')
doc.save(f'{templates_dir}/KYC_TEMPLATE.docx')
print('KYC_TEMPLATE.docx cree')

# LETTRE MISSION CIF TEMPLATE
doc = Document()
doc.add_heading('LETTRE DE MISSION CIF', 0)
doc.add_paragraph('')
doc.add_paragraph('Date: $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('$TITRE_CLIENT$ $PRENOM_CLIENT$ $NOM_CLIENT$,')
doc.add_paragraph('')
doc.add_paragraph('Suite a notre entretien du $DATE_DER$, je vous confirme accepter la mission de conseil en investissements financiers que vous souhaitez me confier.')
doc.add_paragraph('')
doc.add_heading('Objectifs investissement', level=1)
doc.add_paragraph('$OBJECTIFS$')
doc.add_paragraph('')
doc.add_heading('Horizon de placement', level=1)
doc.add_paragraph('$HORIZON$')
doc.add_paragraph('')
doc.add_heading('Profil de risque retenu', level=1)
doc.add_paragraph('$PROFIL_RISQUE$')
doc.add_paragraph('')
doc.add_paragraph('Fait a Papeete, le $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Signature Client                    Signature Conseiller')
doc.save(f'{templates_dir}/LETTRE_MISSION_CIF_TEMPLATE.docx')
print('LETTRE_MISSION_CIF_TEMPLATE.docx cree')

# DECLARATION ADEQUATION TEMPLATE
doc = Document()
doc.add_heading('DECLARATION ADEQUATION', 0)
doc.add_paragraph('')
doc.add_paragraph('Date: $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Client: $TITRE_CLIENT$ $PRENOM_CLIENT$ $NOM_CLIENT$')
doc.add_paragraph('')
doc.add_heading('Profil Client', level=1)
doc.add_paragraph('Objectifs: $OBJECTIFS$')
doc.add_paragraph('Experience: $EXPERIENCE$')
doc.add_paragraph('Tolerance au risque: $TOLERANCE_RISQUE$')
doc.add_paragraph('Horizon: $HORIZON$')
doc.add_paragraph('Profil: $PROFIL_RISQUE$')
doc.add_paragraph('')
doc.add_heading('Adequation', level=1)
doc.add_paragraph('Le conseil donne est adequat au profil du client.')
doc.add_paragraph('Part du patrimoine concernee: $PATRIMOINE_CONFIE_PCT$%')
doc.add_paragraph('')
doc.add_paragraph('Fait a Papeete, le $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Signature Client                    Signature Conseiller')
doc.save(f'{templates_dir}/DECLARATION_ADEQUATION_TEMPLATE.docx')
print('DECLARATION_ADEQUATION_TEMPLATE.docx cree')

# CONVENTION RTO TEMPLATE
doc = Document()
doc.add_heading('CONVENTION DE RECEPTION TRANSMISSION ORDRES', 0)
doc.add_paragraph('')
doc.add_paragraph('Date: $DATE_JOUR$')
doc.add_paragraph('')
doc.add_heading('Entre les soussignes', level=1)
doc.add_paragraph('$CABINET_NOM$')
doc.add_paragraph('RCS: $CABINET_RCS$')
doc.add_paragraph('N ORIAS: $CABINET_ORIAS$')
doc.add_paragraph('Membre de: $CABINET_COMPAGNIE_CIF$ - N: $CABINET_NUM_CIF$')
doc.add_paragraph('')
doc.add_paragraph('Et')
doc.add_paragraph('')
doc.add_paragraph('$TITRE_CLIENT$ $PRENOM_CLIENT$ $NOM_CLIENT$')
doc.add_paragraph('Adresse: $ADRESSE_CLIENT$')
doc.add_paragraph('')
doc.add_heading('Objet', level=1)
doc.add_paragraph('La presente convention a pour objet de definir les conditions dans lesquelles le cabinet transmet les ordres de son client.')
doc.add_paragraph('')
doc.add_paragraph('Fait a Papeete, le $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Signature Client                    Signature Conseiller')
doc.save(f'{templates_dir}/CONVENTION_RTO_TEMPLATE.docx')
print('CONVENTION_RTO_TEMPLATE.docx cree')

# RAPPORT CONSEIL IAS TEMPLATE
doc = Document()
doc.add_heading('RAPPORT DE CONSEIL EN ASSURANCE', 0)
doc.add_paragraph('')
doc.add_paragraph('Date: $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Client: $TITRE_CLIENT$ $PRENOM_CLIENT$ $NOM_CLIENT$')
doc.add_paragraph('')
doc.add_heading('1. Profil Client', level=1)
doc.add_paragraph('Objectifs: $OBJECTIFS$')
doc.add_paragraph('Experience financiere: $EXPERIENCE$')
doc.add_paragraph('Tolerance au risque: $TOLERANCE_RISQUE$')
doc.add_paragraph('Horizon de placement: $HORIZON$')
doc.add_paragraph('Profil de risque: $PROFIL_RISQUE$')
doc.add_paragraph('Part maximale UC: $PROFIL_RISQUE_PCT_MAX$%')
doc.add_paragraph('')
doc.add_heading('2. Situation Patrimoniale', level=1)
doc.add_paragraph('Patrimoine liquide: $PATRIMOINE_LIQUIDE_PCT$%')
doc.add_paragraph('Patrimoine non liquide: $PATRIMOINE_NON_LIQUIDE_PCT$%')
doc.add_paragraph('')
doc.add_heading('3. Conseil', level=1)
doc.add_paragraph('Produit conseille: $PRODUIT_CONSEILLE$')
doc.add_paragraph('Assureur: $ASSUREUR$')
doc.add_paragraph('Montant: $MONTANT_INVESTI$ euros')
doc.add_paragraph('')
doc.add_paragraph('Fait a Papeete, le $DATE_JOUR$')
doc.add_paragraph('')
doc.add_paragraph('Signature Client                    Signature Conseiller')
doc.save(f'{templates_dir}/RAPPORT_CONSEIL_IAS_TEMPLATE.docx')
print('RAPPORT_CONSEIL_IAS_TEMPLATE.docx cree')

print('')
print('Tous les templates ont ete crees avec succes!')

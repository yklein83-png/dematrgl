"""
Script pour créer le template DER complet basé sur le document officiel de la Chambre des CGP
"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Configuration des styles
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10)

# ============================================
# EN-TÊTE
# ============================================
header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = header.add_run("DOCUMENT D'ENTRÉE EN RELATION")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()

# ============================================
# INTRODUCTION
# ============================================
intro = doc.add_paragraph()
intro.add_run("Madame, Monsieur,").bold = True

doc.add_paragraph(
    "Dans le cadre de notre activité de conseil en gestion de patrimoine, notre cabinet est soumis à diverses "
    "réglementations correspondant aux différents statuts que nous exerçons. Le présent document est établi en vue "
    "de vous présenter le cadre réglementaire dans lequel nous allons travailler."
)

doc.add_paragraph(
    "Nous vous adressons les informations réglementaires et les informations complémentaires préconisées par "
    "LA COMPAGNIE CIF, association agréée par l'AMF, par la CNCEF, association agréée par l'ACPR, ainsi que par "
    "LA COMPAGNIE IOBSP, association agréée par l'ACPR, dont nous sommes adhérents."
)

# ============================================
# PRÉSENTATION DES ACTIVITÉS
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Présentation des activités du cabinet").bold = True

doc.add_paragraph(
    "Notre cabinet Le Fare de l'Epargne est enregistré auprès de l'ORIAS sous le numéro 21003330 (www.orias.fr) : "
    "https://www.orias.fr/home"
)

# CIF
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Conseiller en Investissements Financiers (CIF)").bold = True

doc.add_paragraph(
    "Dans le cadre de son activité de CIF, le cabinet est susceptible de fournir des conseils en investissement "
    "de manière non indépendante :"
)

doc.add_paragraph(
    "Le conseil en investissement non indépendant, repose sur une analyse diversifiée de manière large des instruments "
    "ou services financiers disponibles sur le marché, et sans pour autant se limiter aux instruments ou services des "
    "entités partenaires. La rémunération pourra se faire sous forme de rétrocessions de commissions et/ou d'autres "
    "avantages monétaires ou non monétaires perçus ou reçus de la part de nos partenaires, cumulés éventuellement à "
    "nos honoraires. Cette rémunération est justifiée par une amélioration de la qualité de notre prestation et dans "
    "l'objectif d'agir au mieux de vos intérêts."
)

doc.add_paragraph(
    "La nature du conseil ainsi que les contours de la rémunération vous seront communiqués de manière plus précise "
    "au sein de la lettre de mission."
)

doc.add_paragraph("L'autorité de contrôle pour l'activité CIF est :")
p = doc.add_paragraph()
p.add_run("Autorité des Marchés Financiers\n17 Place de la Bourse\n75 082 PARIS Cedex 02").italic = True

# IAS
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Intermédiaire en Assurance (IAS) en qualité de courtier en assurance de catégorie C :").bold = True

doc.add_paragraph(
    "Le cabinet n'est pas soumis à une obligation contractuelle de travailler exclusivement avec une ou plusieurs "
    "entreprises d'assurance et peut communiquer sur simple demande du client, le nom de ces dernières et offre au "
    "client un conseil fondé sur une analyse objective et impartiale du marché, et analyse, pour ce faire un nombre "
    "suffisant de contrats d'assurance offerts sur le marché de façon à pouvoir recommander, en fonction de critères "
    "professionnels, le contrat qui serait adapté aux besoins du souscripteur éventuel."
)

doc.add_paragraph(
    "Dans le cadre de cette activité, le cabinet n'est pas autorisé à encaisser des fonds destinés à un assuré ou "
    "à une compagnie d'assurances, sauf dans le cadre de l'IARD."
)

doc.add_paragraph(
    "Dans le cadre de l'exercice de notre activité d'intermédiation en assurance, notre prestation est fournie sur "
    "un conseil approprié permettant de conseiller le produit le plus adapté aux besoins du client, et le formaliser "
    "au sein d'un rapport d'adéquation."
)

# IOBSP
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Intermédiaire en Opérations de Banque et Services de Paiement (IOBSP) en qualité de :").bold = True

doc.add_paragraph(
    "Courtier en opérations de banque et services de paiement exerçant l'intermédiation en vertu d'un mandat de "
    "notre client pour les catégories de crédit suivantes : crédits immobiliers."
)

doc.add_paragraph(
    "Dans le cadre de son activité d'intermédiation, et conformément à l'article L. 519-6 du Code monétaire et "
    "financier, le cabinet ne perçoit aucune somme représentative de provision, de commissions, de frais de recherche, "
    "de démarches, de constitution de dossier ou d'entremise quelconque, avant le versement effectif des fonds prêtés."
)

# Conseil crédits immobiliers
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Informations spécifiques portant sur le service de conseil de crédits immobiliers").bold = True

doc.add_paragraph("Notre cabinet propose un service de conseil portant sur des contrats de crédits immobiliers.")

doc.add_paragraph(
    "Dans le cadre de ce service, notre recommandation porte sur uniquement une gamme référencée et restreinte de "
    "produits (conseil non indépendant). Cela signifie que, pour vous conseiller, nous nous fondons sur un nombre "
    "restreint de contrats de crédit disponibles sur le marché pour les intermédiaires agissant en vertu d'un mandat "
    "délivré par un client (Article L519-1-1 du code monétaire et financier)."
)

doc.add_paragraph("L'autorité de contrôle des activités IAS et IOBSP est :")
p = doc.add_paragraph()
p.add_run("ACPR\n4 place de Budapest\nDirection du Contrôle des Pratiques Commerciales\n75436 PARIS Cedex 9").italic = True

# ============================================
# LISTE DES PARTENAIRES
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Liste des partenaires").bold = True

table = doc.add_table(rows=4, cols=2)
table.style = 'Table Grid'

# En-têtes
table.rows[0].cells[0].text = "Activité"
table.rows[0].cells[1].text = "Partenaires"
table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

# Données
table.rows[1].cells[0].text = "CIF"
table.rows[1].cells[1].text = "Sofidy, Périal, Iroko, Theoreim, Sogenial, Remake Live, Novaxia, Alderan, Advenis, Corum, Principal, PPG"

table.rows[2].cells[0].text = "IAS"
table.rows[2].cells[1].text = "Vie Plus, UAF Life, UNEP, Intencial, Selencia"

table.rows[3].cells[0].text = "IOBSP"
table.rows[3].cells[1].text = "Socredo, Banque de Polynésie, Banque de Tahiti"

# ============================================
# ORGANISATION ET BONNE CONDUITE
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Organisation et bonne conduite").bold = True

doc.add_paragraph(
    "Nous vous informons que $TITRE_CONSEILLER$ $PRENOM_CONSEILLER$ $NOM_CONSEILLER$ sera votre conseiller, "
    "sous la responsabilité du cabinet."
)

doc.add_paragraph(
    "Enfin, nous vous confirmons que conformément aux Codes de Bonne Conduite de LA COMPAGNIE CIF, La CNCEF et "
    "LA COMPAGNIE IOBSP, le cabinet a pris l'engagement pour maintenir son adhésion, de suivre des formations "
    "réglementaires obligatoires, de justifier annuellement d'une assurance responsabilité civile professionnelle, "
    "de produire le casier judiciaire des dirigeants et de déclarer immédiatement, sous peine de déchéance, tout "
    "événement susceptible de le modifier, enfin de respecter toutes les dispositions incluses dans les Codes de "
    "Bonne Conduite de LA COMPAGNIE CIF et La CNCEF et LA COMPAGNIE IOBSP."
)

# ============================================
# TRAITEMENT DES RÉCLAMATIONS
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Traitement des réclamations").bold = True

doc.add_paragraph(
    "La réclamation peut être exprimée lors d'un rendez-vous et doit être adressée par écrit à la Direction du "
    "cabinet par courrier BP4646 – 98713-Papeete-Tahiti ou par un e-mail à contact@fare-epargne.com."
)

doc.add_paragraph("Le cabinet s'engage à :")
doc.add_paragraph("• Accuser réception de la réclamation dans un délai de dix jours ouvrables ;")
doc.add_paragraph(
    "• Puis à y répondre dans un délai de deux mois maximums à compter de la date d'envoi de la réclamation, "
    "sauf survenance de circonstances particulières dûment justifiées."
)

doc.add_paragraph(
    "Si la réponse apportée à sa réclamation ne vous apparaît pas satisfaisante, vous pourrez saisir le médiateur "
    "de la consommation compétent suivant :"
)

p = doc.add_paragraph()
p.add_run("Pour les activités IAS et COBSP :").bold = True
doc.add_paragraph("CNPM Médiation Consommation\n27 avenue de la Libération\n42400 Saint Chamond\n"
                  "Tél. : 09 88 30 27 72\ncontact-admin@cnpm-mediation-consommation.eu")

p = doc.add_paragraph()
p.add_run("En second lieu, spécialement pour l'activité de CIF :").bold = True
doc.add_paragraph("Le Médiateur de l'AMF (instruments financiers et services d'investissements) :\n"
                  "Autorité des Marchés Financiers\n17, Place de la Bourse\n75 082 Paris Cedex 02\n"
                  "http://www.amf-france.org/Le-mediateur-de-l-AMF/Presentation.html")

# ============================================
# POLITIQUE DURABILITÉ
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Politique en matière de durabilité").bold = True

doc.add_paragraph(
    "Notre cabinet a mis en place un processus de sélection des instruments financiers, tenant compte des facteurs "
    "de durabilité au sens de l'article 2, point 24, du règlement (UE) 2019/2088 du Parlement européen et du Conseil "
    "du 27 novembre 2019. Pour ce faire, nous procédons à l'analyse des produits qu'il entend référencer sur la base "
    "de plusieurs critères :"
)

doc.add_paragraph("• Le fait qu'ils constituent des investissements durables au sens du Règlement SFDR ;")
doc.add_paragraph("• Le fait qu'ils constituent des investissements durables à vocation environnementale au sens du Règlement Taxonomie ;")
doc.add_paragraph("• Le fait qu'ils prennent ou non en compte les incidences négatives en matière de durabilité (PIA).")

doc.add_paragraph("Ces différents éléments vous seront expliqués dans le cadre du déroulement de notre mission.")

# ============================================
# DONNÉES PERSONNELLES
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Traitement des données personnelles et confidentialité").bold = True

doc.add_paragraph(
    "Notre travail nécessitant l'accès, la collecte, le traitement, le stockage, l'exploitation et le cas échéant, "
    "la transmission à des tiers des données à caractère personnel que vous voudrez bien nous transmettre, nous nous "
    "engageons à respecter les dispositions de la loi n° 78-17 du 6 janvier 1978 et du Règlement 2016/679 du Parlement "
    "européen et du Conseil du 27 avril 2016 relatif à la protection des personnes physiques à l'égard du traitement "
    "des données à caractère personnel et à la libre circulation de ces données."
)

doc.add_paragraph(
    "Les données personnelles que vous nous transmettez dans le cadre de nos activités et des services proposés sont "
    "collectées et traitées par Mr Pierre POHER en qualité de responsable de traitement. L'identité et les coordonnées "
    "du délégué à la protection des données personnelles au sein du cabinet sont : Mr Pierre POHER – contact@fare-epargne.com."
)

doc.add_paragraph(
    "Notre cabinet s'engage à ne collecter et traiter les données recueillies qu'au regard des finalités de traitement "
    "convenues (consentement, nécessité contractuelle, respect d'une obligation légale) entre notre cabinet et son client, "
    "à en préserver la sécurité et l'intégrité, à ne communiquer ses informations qu'aux tiers auxquels il est nécessaire "
    "de les transmettre pour la bonne exécution des missions confiées."
)

doc.add_paragraph(
    "Nous vous informons que vous bénéficiez d'un droit d'accès, de rectification, d'un droit de restriction sur "
    "l'utilisation, d'un droit d'opposition à l'utilisation, d'un droit à l'effacement ou encore d'un droit à la "
    "portabilité de vos données personnelles."
)

doc.add_paragraph("Vous pouvez ainsi bénéficier de ces droits à tout moment en nous informant directement :")
doc.add_paragraph("• par courrier électronique : contact@fare-epargne.com")
doc.add_paragraph("• par courrier postal : BP4646 – 98713-Papeete-Tahiti")

doc.add_paragraph(
    "Vous disposez également du droit d'introduire une réclamation auprès de la CNIL (Commission Nationale de "
    "l'Information et des Libertés) via leur site internet ou par voie postale : CNIL - 3 place de Fontenoy - "
    "TSA 80715 75334 PARIS CEDEX 07."
)

# ============================================
# MOYENS DE COMMUNICATION
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Moyens de communication").bold = True

doc.add_paragraph("Nous vous informons que nous communiquerons avec vous par les moyens suivants :")
doc.add_paragraph("• Par courrier")
doc.add_paragraph(
    "• Par tout autre support durable (e-mail, site internet dédié, etc…) dès lors que les clients auront opté "
    "formellement pour la fourniture des informations sur cet autre support et que ce même support se révèlera "
    "adapté au contexte dans lequel seront conduites les affaires."
)

doc.add_paragraph(
    "Les coordonnées de contact suivantes que vous nous avez communiquées pourront à tout moment être utilisées "
    "en vue de sécuriser et d'authentifier nos échanges."
)

# Case à cocher et courriel
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("☐ ").font.size = Pt(14)
p.add_run(
    "En cochant cette case, vous acceptez que les informations soient communiquées par le biais d'un support "
    "durable autre que le papier, à savoir par courriel. À tout moment, vous pouvez demander à disposer du "
    "document papier ou à revenir à une remise de document sous format papier."
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Courriel : ").bold = True
p.add_run("$CLIENT_EMAIL$")

# ============================================
# MISE À JOUR
# ============================================
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("Mise à jour des informations").bold = True

doc.add_paragraph(
    "Nous vous ferons parvenir toute mise à jour des différentes informations contenues dans le présent document "
    "en vous les communiquant par courrier, e-mail ou espace dédié. Vous pouvez également obtenir à tout moment "
    "ces informations sur simple demande auprès de notre cabinet."
)

# ============================================
# CONCLUSION ET SIGNATURES
# ============================================
doc.add_paragraph()
doc.add_paragraph("En vous remerciant de la confiance que vous voudrez bien nous accorder,")
doc.add_paragraph()
doc.add_paragraph("Nous vous prions de croire, $TITRE_CONTACT$ $NOM_CONTACT$, à l'assurance de nos sentiments distingués.")

doc.add_paragraph()
doc.add_paragraph()

# Tableau signatures
table_sig = doc.add_table(rows=4, cols=2)
table_sig.rows[0].cells[0].text = "Fait à Papeete, le $DATE_JOUR$"
table_sig.rows[0].cells[1].text = ""

table_sig.rows[1].cells[0].text = ""
table_sig.rows[1].cells[1].text = ""

table_sig.rows[2].cells[0].text = "Signature du client"
table_sig.rows[2].cells[1].text = "Signature du conseiller"
table_sig.rows[2].cells[0].paragraphs[0].runs[0].bold = True
table_sig.rows[2].cells[1].paragraphs[0].runs[0].bold = True

table_sig.rows[3].cells[0].text = "$TITRE_CONTACT$ $PRENOM_CONTACT$ $NOM_CONTACT$"
table_sig.rows[3].cells[1].text = "$TITRE_CONSEILLER$ $PRENOM_CONSEILLER$ $NOM_CONSEILLER$"

# Sauvegarder
doc.save('/app/templates/DER_TEMPLATE.docx')
print("Template DER créé avec succès!")

# Afficher les placeholders
import re
doc2 = Document('/app/templates/DER_TEMPLATE.docx')
placeholders = set()
for para in doc2.paragraphs:
    found = re.findall(r'\$[A-Z_]+\$', para.text)
    placeholders.update(found)
for table in doc2.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                found = re.findall(r'\$[A-Z_]+\$', para.text)
                placeholders.update(found)

print("\nPlaceholders dans le template:")
for p in sorted(placeholders):
    print(f"  {p}")

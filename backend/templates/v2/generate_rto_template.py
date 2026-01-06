"""
Script pour générer le template RTO (Convention de Réception et Transmission d'Ordres) v2
Avec mise en page professionnelle et placeholders {{FIELD}}
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# Mentions légales obligatoires
MENTIONS_LEGALES = """
Le Fare de l'Épargne - SARL au capital de 5 000 € - SIRET 893 599 979 00016 - RCS PAPEETE
Siège social : Immeuble Vatea, Rue des Remparts - Mission Papeete - 98714 PAPEETE
Conseiller en Investissements Financiers (CIF) - Membre de LA COMPAGNIE CIF, association agréée par l'AMF
Courtier en assurance (COA) immatriculé à l'ORIAS sous le n° 21003330 (www.orias.fr)
Courtier en opérations de banque et services de paiement (COBSP)
Responsabilité Civile Professionnelle et Garantie Financière conformes aux articles L541-3 et L512-6 et 7 du Code des Assurances
"""


def set_cell_shading(cell, color):
    """Définir la couleur de fond d'une cellule"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_style(doc, text, level=1):
    """Ajouter un titre avec style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True

    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

    return p


def add_article_title(doc, number, title):
    """Ajouter un titre d'article"""
    p = doc.add_paragraph()
    run = p.add_run(f"Article {number} : {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 51, 102)
    return p


def add_normal_paragraph(doc, text, indent=False, justify=True):
    """Ajouter un paragraphe normal"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    if indent:
        p.paragraph_format.left_indent = Cm(0.5)

    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return p


def add_bullet_list(doc, items):
    """Ajouter une liste à puces"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)


def generate_rto_template():
    """Générer le template RTO v2"""
    doc = Document()

    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ==========================================
    # EN-TÊTE
    # ==========================================

    # Logo / En-tête cabinet
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("LE FARE DE L'ÉPARGNE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Sous-titre
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_header.add_run("Conseil en Gestion de Patrimoine")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    # ==========================================
    # TITRE DU DOCUMENT
    # ==========================================

    add_heading_style(doc, "CONVENTION DE RÉCEPTION ET TRANSMISSION D'ORDRES", 1)

    # Version
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_para.add_run("Version 2025.03")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ==========================================
    # INFORMATIONS RÉFÉRENCE
    # ==========================================

    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'

    # Ligne 1
    info_table.rows[0].cells[0].text = "Date :"
    info_table.rows[0].cells[1].text = "{{DATE_SIGNATURE}}"
    info_table.rows[0].cells[2].text = "Référence :"
    info_table.rows[0].cells[3].text = "{{NUMERO_CLIENT}}"

    # Ligne 2
    info_table.rows[1].cells[0].text = "Lieu :"
    info_table.rows[1].cells[1].text = "{{LIEU_SIGNATURE}}"
    info_table.rows[1].cells[2].text = "Conseiller :"
    info_table.rows[1].cells[3].text = "{{NOM_CONSEILLER}}"

    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i % 2 == 0:  # Labels
                        run.font.bold = True
            if i % 2 == 0:
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # ==========================================
    # LES PARTIES
    # ==========================================

    add_heading_style(doc, "Entre les soussignés :", 2)
    doc.add_paragraph()

    # ==========================================
    # PARTIE 1 : LE CLIENT
    # ==========================================

    # Tableau info client
    client_table = doc.add_table(rows=1, cols=1)
    client_table.style = 'Table Grid'
    cell = client_table.rows[0].cells[0]
    set_cell_shading(cell, 'F5F5F5')

    p = cell.paragraphs[0]
    p.add_run("LE CLIENT\n\n").bold = True

    # Informations client avec placeholders
    p.add_run("{{T1_CIVILITE}} {{T1_PRENOM}} {{T1_NOM}}\n")
    p.add_run("Né(e) le {{T1_DATE_NAISSANCE}} à {{T1_LIEU_NAISSANCE}}\n")
    p.add_run("Nationalité : {{T1_NATIONALITE}}\n")
    p.add_run("Demeurant : {{T1_ADRESSE}}, {{T1_CODE_POSTAL}} {{T1_VILLE}}\n")
    p.add_run("Profession : {{T1_PROFESSION}}\n")
    p.add_run("Situation familiale : {{SITUATION_FAMILIALE}}\n")
    p.add_run("Régime matrimonial : {{REGIME_MATRIMONIAL}}\n")

    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    party1 = doc.add_paragraph()
    party1.add_run("Ci-après désigné « le Client », d'une part ;").italic = True

    doc.add_paragraph()
    and_para = doc.add_paragraph()
    and_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    and_para.add_run("et,").bold = True

    doc.add_paragraph()

    # ==========================================
    # PARTIE 2 : LE CIF
    # ==========================================

    cif_table = doc.add_table(rows=1, cols=1)
    cif_table.style = 'Table Grid'
    cell = cif_table.rows[0].cells[0]
    set_cell_shading(cell, 'E8F4FD')  # Bleu clair

    p = cell.paragraphs[0]
    p.add_run("LE CABINET - CONSEILLER EN INVESTISSEMENTS FINANCIERS\n\n").bold = True

    p.add_run("Le Fare de l'Épargne, SARL au capital de 5 000 €, ")
    p.add_run("enregistré au RCS de Papeete sous le n° 893 599 979, ")
    p.add_run("à l'ORIAS sous le numéro 21003330, ")
    p.add_run("en sa qualité de Conseiller en Investissements Financiers, ")
    p.add_run("enregistré par La Compagnie CIF, association agréée par l'Autorité des Marchés Financiers.\n\n")
    p.add_run("Siège social : Immeuble Vatea, Rue des Remparts - Mission Papeete - 98714 PAPEETE\n")
    p.add_run("Représenté par : {{NOM_CONSEILLER}}")

    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    party2 = doc.add_paragraph()
    party2.add_run("Ci-après désigné « le CIF », d'autre part,").italic = True

    doc.add_paragraph()

    parties_para = doc.add_paragraph()
    parties_para.add_run("Et désignés ensemble « les Parties ».").italic = True

    doc.add_paragraph()

    # ==========================================
    # OBJET DE LA CONVENTION
    # ==========================================

    add_heading_style(doc, "Nature et modalités de la prestation de conseil", 2)

    add_normal_paragraph(doc,
        "Les parties soussignées définissent comme suit l'objet de la convention de réception et transmission d'ordres :")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 1 : Règles de fonctionnement
    # ==========================================

    add_article_title(doc, 1, "Règles de fonctionnement et responsabilité du compte titres")

    add_normal_paragraph(doc,
        "En application de l'article 325-13 du Règlement général de l'AMF relatif à la réception et transmission "
        "d'ordres OPC par les CIF et conformément aux procédures préconisées par LA COMPAGNIE CIF, association "
        "professionnelle agréée par l'Autorité des Marchés Financiers, à laquelle nous avons adhéré et au titre "
        "de laquelle le cabinet a été inscrit sur la liste des Conseillers en Investissements Financiers "
        "(N° d'enregistrement à l'ORIAS : 21003330), il est convenu ce qui suit :")

    add_normal_paragraph(doc,
        "Le client est titulaire d'un compte de titres (préciser PEA s'il y a lieu, lister en cas de multiplicité "
        "de comptes). Seul le client est habilité à mouvementer ce(s) portefeuille(s). Le CIF ne bénéficie d'aucun "
        "mandat de gestion et ne peut procéder à aucune opération de gestion ou d'arbitrage à son initiative, ces "
        "opérations relevant du seul pouvoir du client.")

    add_normal_paragraph(doc,
        "Le Client déclare avoir les connaissances et l'aptitude à réaliser les opérations pour lesquelles il "
        "donnera un ordre. Il pourra néanmoins recueillir les conseils du CIF.")

    add_normal_paragraph(doc,
        "Le client s'engage à informer le CIF de toute modification de patrimoine ou personnelle pouvant influer "
        "le conseil qui lui serait habituellement donné et notamment de tout engagement risqué pris par ailleurs.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 2 : Rédaction des ordres
    # ==========================================

    add_article_title(doc, 2, "Rédaction des ordres")

    add_normal_paragraph(doc,
        "De préférence le client devra utiliser les modèles d'ordres qui lui auront été confiés par le CIF. "
        "À défaut, la rédaction de son ordre sur papier libre devra clairement indiquer son identité, celle de "
        "son compte, l'opération achat ou vente au regard du nom de chaque titre, accompagné de son code ISIN, "
        "et être adressée exclusivement au Cabinet par courrier postal, télécopie ou courrier e-mail.")

    add_normal_paragraph(doc,
        "Dans ces deux derniers cas, l'original du document sera expédié le même jour, à défaut dès le premier "
        "jour ouvrable suivant au cabinet par courrier postal.")

    add_normal_paragraph(doc,
        "Dans l'intérêt du client, en cas de doute, le Cabinet CIF pourra se faire confirmer l'ordre qui lui "
        "sera parvenu.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 3 : Réception de l'ordre
    # ==========================================

    add_article_title(doc, 3, "Réception de l'ordre")

    add_normal_paragraph(doc,
        "Cette prestation, exclusivement réservée aux ordres d'OPC (organismes de placements collectifs listés "
        "à l'article 214-1 du Code monétaire et financier : OPCVM, SICAF, SCPI, OPCI, etc.), est partie "
        "intégrante de la seule activité de conseil réalisée par le CIF au profit du client.")

    add_normal_paragraph(doc,
        "Le Cabinet CIF horodatera l'ordre dès sa réception pour confirmer la prise en compte de l'ordre du "
        "client : remise du courrier par la Poste, ouverture des bureaux pour les fax, ouverture du courrier e-mail.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 4 : Transmission de l'ordre
    # ==========================================

    add_article_title(doc, 4, "Transmission de l'ordre")

    add_normal_paragraph(doc,
        "À réception de l'ordre émis par le Client dans un délai maximal de 3 jours ouvrables, le Cabinet CIF "
        "transmettra l'ordre à l'établissement teneur de compte, selon les modalités de l'établissement "
        "(courrier, fax, e-mail, etc.).")

    add_normal_paragraph(doc,
        "Le Cabinet CIF ne peut être tenu responsable de toute erreur ou manquement d'exécution commis par "
        "l'établissement teneur du compte dans l'exécution de l'ordre. Par contre, il veillera à s'assurer "
        "de la correction au mieux des intérêts du client, dès qu'il aura eu connaissance de l'erreur ou du "
        "défaut d'exécution.")

    add_normal_paragraph(doc,
        "Le Cabinet CIF ne peut être tenu responsable en cas d'interruption humaine ou mécanique dans les "
        "moyens de transmission.")

    add_normal_paragraph(doc,
        "Le Cabinet CIF ne peut être tenu pour responsable d'éventuelles difficultés à exécuter l'ordre si "
        "les conditions de marché ou les conditions légales, réglementaires ou conventionnelles ne s'y prêtent "
        "pas. La passation de l'ordre par le client ne préjuge donc pas avec certitude de son exécution.")

    add_normal_paragraph(doc,
        "Dès qu'il en a connaissance ou parce qu'il le décide, le cabinet CIF informe le client par téléphone "
        "et confirme par courrier. En cas de difficulté à joindre le client, celui-ci peut être informé par "
        "fax, e-mail ou courrier. Même corrigé, l'ancien ordre signé par le client ne peut être utilisé. "
        "Si besoin est, ce dernier devra émettre un nouvel ordre signé de lui.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 5 : Confirmation de l'exécution
    # ==========================================

    add_article_title(doc, 5, "Confirmation de l'exécution de l'ordre")

    add_normal_paragraph(doc,
        "L'avis d'opéré est adressé directement dès exécution de l'ordre par l'établissement teneur de compte "
        "ou la société de gestion de portefeuilles. Il fait foi de l'exécution de l'ordre.")

    add_normal_paragraph(doc,
        "Si dans un délai de six jours ouvrables, le client n'a pas reçu d'avis d'opéré, il peut se manifester "
        "auprès du CIF qui pourra lui confirmer ou infirmer que son ordre a bien été transmis.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 6 : Rémunération
    # ==========================================

    add_article_title(doc, 6, "Rémunération")

    add_normal_paragraph(doc,
        "En aucun cas le Client n'a à s'acquitter de droits d'entrée supérieurs ou de frais supplémentaires "
        "à ce qui figure dans la fiche technique qui lui aura été remise par le CIF.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 7 : Résiliation
    # ==========================================

    add_article_title(doc, 7, "Condition de résiliation de la convention")

    add_normal_paragraph(doc,
        "La convention est conclue pour une durée d'une année. Elle est prorogée chaque année par tacite "
        "reconduction. Elle peut être résiliée à tout moment par le client, en adressant une lettre "
        "recommandée avec accusé de réception. Une copie est adressée par le client à l'établissement "
        "teneur du compte.")

    doc.add_paragraph()

    # ==========================================
    # ARTICLE 8 : Droit applicable
    # ==========================================

    add_article_title(doc, 8, "Droit applicable")

    add_normal_paragraph(doc,
        "La présente convention est soumise au droit français.")

    doc.add_paragraph()
    doc.add_paragraph()

    # ==========================================
    # COMPTES TITRES CONCERNÉS
    # ==========================================

    add_heading_style(doc, "Comptes titres concernés par la présente convention", 2)

    compte_table = doc.add_table(rows=4, cols=3)
    compte_table.style = 'Table Grid'

    # En-têtes
    headers = ["Type de compte", "Établissement", "N° de compte"]
    for i, header in enumerate(headers):
        cell = compte_table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, '003366')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Lignes vides pour remplissage
    for row_idx in range(1, 4):
        for col_idx in range(3):
            cell = compte_table.rows[row_idx].cells[col_idx]
            if row_idx == 1:
                placeholders = ["{{COMPTE_TYPE_1}}", "{{COMPTE_ETABLISSEMENT_1}}", "{{COMPTE_NUMERO_1}}"]
            elif row_idx == 2:
                placeholders = ["{{COMPTE_TYPE_2}}", "{{COMPTE_ETABLISSEMENT_2}}", "{{COMPTE_NUMERO_2}}"]
            else:
                placeholders = ["{{COMPTE_TYPE_3}}", "{{COMPTE_ETABLISSEMENT_3}}", "{{COMPTE_NUMERO_3}}"]
            cell.text = placeholders[col_idx]
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # ==========================================
    # SIGNATURES
    # ==========================================

    add_heading_style(doc, "Signatures", 2)

    # Date et lieu
    date_lieu = doc.add_paragraph()
    date_lieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_lieu.add_run("Fait à {{LIEU_SIGNATURE}}, le {{DATE_SIGNATURE}}")
    run.font.size = Pt(11)
    run.bold = True

    exemplaires = doc.add_paragraph()
    exemplaires.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = exemplaires.add_run("En {{NOMBRE_EXEMPLAIRES}} exemplaires originaux")
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Tableau des signatures
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.style = 'Table Grid'

    # En-têtes
    sig_table.rows[0].cells[0].text = "Le Client"
    sig_table.rows[0].cells[1].text = "Le Conseiller (CIF)"
    for cell in sig_table.rows[0].cells:
        set_cell_shading(cell, 'E8E8E8')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Noms
    sig_table.rows[1].cells[0].text = "{{T1_CIVILITE}} {{T1_PRENOM}} {{T1_NOM}}"
    sig_table.rows[1].cells[1].text = "{{NOM_CONSEILLER}}"
    for cell in sig_table.rows[1].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)

    # Mention "Lu et approuvé"
    sig_table.rows[2].cells[0].text = "(Signature précédée de la mention « Lu et approuvé »)"
    sig_table.rows[2].cells[1].text = ""
    for para in sig_table.rows[2].cells[0].paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.size = Pt(9)
            run.italic = True

    # Espace signature
    sig_table.rows[3].cells[0].text = "\n\n\n\n"
    sig_table.rows[3].cells[1].text = "\n\n\n\n"

    # Largeurs des colonnes
    for row in sig_table.rows:
        for cell in row.cells:
            cell.width = Cm(8)

    doc.add_paragraph()
    doc.add_paragraph()

    # ==========================================
    # MENTIONS LÉGALES (Pied de page)
    # ==========================================

    # Ligne de séparation
    sep = doc.add_paragraph()
    sep.add_run("─" * 80).font.size = Pt(6)

    # Mentions légales
    mentions = doc.add_paragraph()
    mentions.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mentions.add_run(MENTIONS_LEGALES.strip())
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Sauvegarder le template
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "RTO_V2_TEMPLATE.docx")
    doc.save(output_path)
    print(f"Template RTO v2 généré : {output_path}")

    return output_path


if __name__ == "__main__":
    generate_rto_template()

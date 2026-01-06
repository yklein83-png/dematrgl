# -*- coding: utf-8 -*-
"""
Générateur de templates DOCX pour Le Fare de l'Épargne
Crée deux documents séparés:
1. QCC_V2_TEMPLATE.docx - Questionnaire de Connaissance Client
2. PROFIL_RISQUE_V2_TEMPLATE.docx - Profil de Risque

Placeholders au format: {{CHAMP_NAME}}
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Mentions légales obligatoires
MENTIONS_LEGALES = """Le Fare de l'Epargne/ SARL au capital social de 100 000 XPF/ R.C.S. Papeete 20 346 B/ code APE 7022Z / Siège Social : route de Teroma, quartier Lauglin, Faaa - BP 4646 – 98713 Papeete – (+689) 87 77 42 08 – contact@fare-epargne.pf – faredelepargne.pf/ Enregistré à l'ORIAS sous le n° 21003330 (www.orias.fr) en qualité de : Conseiller en investissement financier adhérent de La Compagnie CIF, association agréée auprès de l'Autorité des Marchés Financiers et Intermédiaire en opérations de banque et services de paiement en qualité de courtier adhérent de La Compagnie IOBSP, association agréée auprès de l'ACPR/ Un crédit vous engage et doit être remboursé. Vérifiez vos capacités de remboursement avant de vous engager/ Intermédiaire en assurance en qualité de courtier adhérent de la CNCEF, association agréée auprès de l'ACPR/ Responsabilité Civile Professionnelle souscrite auprès de MMA IARD N° 127.110.017, 14 Boulevard Marie et Alexandre Oyon 72030 Le Mans Cedex 9/ Médiateur : CNP Médiation Consommation - 27 avenue de la Libération - 42400 Saint Chamond/ Ne peut recevoir aucun fonds, effet, ou valeur."""


def set_cell_shading(cell, color):
    """Définit la couleur de fond d'une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_checkbox(paragraph, checked=False):
    """Ajoute une case à cocher"""
    symbol = "☑" if checked else "☐"
    run = paragraph.add_run(symbol + " ")
    run.font.size = Pt(11)


def create_styled_document():
    """Crée un document avec styles prédéfinis"""
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Style titre principal
    style = doc.styles.add_style('TitrePrincipal', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(16)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 51, 102)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style.paragraph_format.space_after = Pt(12)

    # Style section
    style = doc.styles.add_style('TitreSection', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    style.font.bold = True
    style.font.color.rgb = RGBColor(0, 51, 102)
    style.paragraph_format.space_before = Pt(18)
    style.paragraph_format.space_after = Pt(6)

    # Style sous-section
    style = doc.styles.add_style('SousSection', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(12)
    style.paragraph_format.space_after = Pt(6)

    # Style normal
    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10)

    return doc


def add_header_footer(doc, title):
    """Ajoute en-tête et pied de page"""
    section = doc.sections[0]

    # En-tête
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Le Fare de l'Épargne - {title}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.style.font.size = Pt(9)
    header_para.style.font.color.rgb = RGBColor(128, 128, 128)

    # Pied de page
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Version 2025-03 | Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.style.font.size = Pt(8)


def create_two_column_table(doc, left_title, right_title):
    """Crée un tableau à 2 colonnes pour Titulaire 1 / Titulaire 2"""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes
    for i, title in enumerate([left_title, right_title]):
        cell = table.rows[0].cells[i]
        cell.text = title
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'E6E6E6')

    return table


def add_field_row(table, label, placeholder_t1, placeholder_t2=None):
    """Ajoute une ligne de champ au tableau"""
    row = table.add_row()
    row.cells[0].text = f"{label}: {placeholder_t1}"
    if placeholder_t2:
        row.cells[1].text = f"{label}: {placeholder_t2}"
    else:
        row.cells[1].text = ""


# ===========================================
# DOCUMENT 1: QUESTIONNAIRE CONNAISSANCE CLIENT
# ===========================================

def create_qcc_document():
    """Crée le document QCC complet"""
    doc = create_styled_document()
    add_header_footer(doc, "Questionnaire de Connaissance Client")

    # TITRE PRINCIPAL
    title = doc.add_paragraph("QUESTIONNAIRE DE CONNAISSANCE CLIENT", style='TitrePrincipal')
    doc.add_paragraph("V.2025-03", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Référence client
    p = doc.add_paragraph()
    p.add_run("Référence client: ").bold = True
    p.add_run("{{NUMERO_CLIENT}}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # MISE EN GARDE
    doc.add_paragraph("Mise en garde et informations préalables", style='TitreSection')

    mise_en_garde = """Dès lors qu'il fournit des conseils en investissements financiers et/ou en produits d'investissements assurantiels (contrats d'assurance vie ou de capitalisation) à son client, le cabinet doit s'enquérir des exigences et besoins de ce dernier, de sa situation financière, de ses objectifs, de ses connaissances et de son expérience en matière financière, de sa tolérance au risque, de sa capacité à supporter les pertes, ainsi que de ses préférences en matière de durabilité, afin de délivrer un conseil adapté.

À cet effet, nous vous soumettons ce recueil d'informations patrimoniales qui comporte un questionnaire de connaissance du client.

Nous attirons votre attention sur le fait que vous vous engagez à une communication exhaustive et sincère des informations vous concernant et que dans le cas contraire :
• Dans le cadre de la commercialisation d'un contrat d'assurance, le conseil délivré pourrait ne pas être totalement adapté,
• Dans le cadre de la fourniture d'un conseil en investissements financiers, le cabinet devra s'abstenir de vous recommander les opérations, instruments et services en question.

Toute modification de votre situation familiale, patrimoniale ou professionnelle peut avoir une incidence sur la prestation fournie et doit donc être signalée à votre conseiller afin d'actualiser le présent document."""

    doc.add_paragraph(mise_en_garde)

    # RGPD
    doc.add_paragraph("Protection des données personnelles", style='SousSection')

    rgpd = """Les informations recueillies dans le présent formulaire font l'objet d'un traitement destiné à établir un diagnostic de votre situation patrimoniale et de celle de votre foyer, et à vous conseiller relativement à la gestion de votre patrimoine dans le cadre des activités professionnelles de conseil en gestion de patrimoine. Ces informations sont nécessaires pour permettre au cabinet de réaliser ses missions. Le défaut de réponse peut avoir des conséquences sur la réalisation conforme des missions du cabinet. Conformément au Règlement Général sur la protection des données personnelles (RGPD), vous disposez sur ces données d'un droit d'accès, de rectification, et limitation, ainsi que d'un droit d'opposition et de portabilité conformément à la loi. Si vous souhaitez exercer ces droits, vous pouvez nous contacter par email. Vous disposez également du droit d'introduire une réclamation auprès de la CNIL."""

    doc.add_paragraph(rgpd)

    # SECTION 1: CONNAISSANCE CLIENT
    doc.add_paragraph("1 - CONNAISSANCE CLIENT", style='TitreSection')

    # Instructions
    instructions = """Ce questionnaire doit être renseigné et signé par la et/ou les personne(s) suivante(s) :
• Pour un compte ayant un seul titulaire : par le titulaire du compte ou par son représentant légal ;
• Pour un compte en indivision ou un compte joint : chaque titulaire du compte doit remplir et signer le questionnaire ;
• En cas de procuration : le titulaire du compte et son mandataire doivent chacun remplir et signer."""
    doc.add_paragraph(instructions)

    # Tableau Titulaires
    doc.add_paragraph("Identification des titulaires", style='SousSection')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    # En-têtes
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "TITULAIRE 1"
    hdr_cells[1].text = "TITULAIRE 2"
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # Civilité
    row = table.add_row()
    row.cells[0].text = "Civilité: {{T1_CIVILITE}}"
    row.cells[1].text = "Civilité: {{T2_CIVILITE}}"

    # Nom
    row = table.add_row()
    row.cells[0].text = "Nom: {{T1_NOM}}"
    row.cells[1].text = "Nom: {{T2_NOM}}"

    # Nom de jeune fille
    row = table.add_row()
    row.cells[0].text = "Nom de naissance: {{T1_NOM_JEUNE_FILLE}}"
    row.cells[1].text = "Nom de naissance: {{T2_NOM_JEUNE_FILLE}}"

    # Prénom
    row = table.add_row()
    row.cells[0].text = "Prénom(s): {{T1_PRENOM}}"
    row.cells[1].text = "Prénom(s): {{T2_PRENOM}}"

    # Date et lieu de naissance
    row = table.add_row()
    row.cells[0].text = "Né(e) le: {{T1_DATE_NAISSANCE}}\nà: {{T1_LIEU_NAISSANCE}}"
    row.cells[1].text = "Né(e) le: {{T2_DATE_NAISSANCE}}\nà: {{T2_LIEU_NAISSANCE}}"

    # Nationalité
    row = table.add_row()
    row.cells[0].text = "Nationalité: {{T1_NATIONALITE}}"
    row.cells[1].text = "Nationalité: {{T2_NATIONALITE}}"

    # Adresse
    row = table.add_row()
    row.cells[0].text = "Adresse: {{T1_ADRESSE}}\n{{T1_CODE_POSTAL}} {{T1_VILLE}}"
    row.cells[1].text = "Adresse: {{T2_ADRESSE}}\n{{T2_CODE_POSTAL}} {{T2_VILLE}}"

    # Email
    row = table.add_row()
    row.cells[0].text = "Email: {{T1_EMAIL}}"
    row.cells[1].text = "Email: {{T2_EMAIL}}"

    # Téléphone
    row = table.add_row()
    row.cells[0].text = "Téléphone: {{T1_TELEPHONE}}"
    row.cells[1].text = "Téléphone: {{T2_TELEPHONE}}"

    # Pièce d'identité
    row = table.add_row()
    row.cells[0].text = "Pièce d'identité: {{T1_PIECE_IDENTITE}}\nN°: {{T1_NUMERO_PIECE}}\nValidité: {{T1_DATE_VALIDITE_PIECE}}"
    row.cells[1].text = "Pièce d'identité: {{T2_PIECE_IDENTITE}}\nN°: {{T2_NUMERO_PIECE}}\nValidité: {{T2_DATE_VALIDITE_PIECE}}"

    # Protection juridique
    row = table.add_row()
    row.cells[0].text = "Régime de protection juridique: {{T1_REGIME_PROTECTION_JURIDIQUE}}\nForme: {{T1_REGIME_PROTECTION_FORME}}\nReprésentant légal: {{T1_REPRESENTANT_LEGAL}}"
    row.cells[1].text = "Régime de protection juridique: {{T2_REGIME_PROTECTION_JURIDIQUE}}\nForme: {{T2_REGIME_PROTECTION_FORME}}\nReprésentant légal: {{T2_REPRESENTANT_LEGAL}}"

    # Résidence fiscale
    row = table.add_row()
    row.cells[0].text = "Résidence fiscale: {{T1_RESIDENCE_FISCALE}}\nNIF: {{T1_NIF}}"
    row.cells[1].text = "Résidence fiscale: {{T2_RESIDENCE_FISCALE}}\nNIF: {{T2_NIF}}"

    # Profession
    row = table.add_row()
    row.cells[0].text = "Situation professionnelle: {{T1_SITUATION_PRO}}\nProfession: {{T1_PROFESSION}}\nSecteur d'activité: {{T1_SECTEUR_ACTIVITE}}\nEmployeur: {{T1_EMPLOYEUR}}"
    row.cells[1].text = "Situation professionnelle: {{T2_SITUATION_PRO}}\nProfession: {{T2_PROFESSION}}\nSecteur d'activité: {{T2_SECTEUR_ACTIVITE}}\nEmployeur: {{T2_EMPLOYEUR}}"

    # Chef d'entreprise
    row = table.add_row()
    row.cells[0].text = "Chef d'entreprise: {{T1_CHEF_ENTREPRISE}}\nDénomination: {{T1_ENTREPRISE_DENOMINATION}}\nForme juridique: {{T1_ENTREPRISE_FORME_JURIDIQUE}}\nSiège social: {{T1_ENTREPRISE_SIEGE_SOCIAL}}"
    row.cells[1].text = "Chef d'entreprise: {{T2_CHEF_ENTREPRISE}}\nDénomination: {{T2_ENTREPRISE_DENOMINATION}}\nForme juridique: {{T2_ENTREPRISE_FORME_JURIDIQUE}}\nSiège social: {{T2_ENTREPRISE_SIEGE_SOCIAL}}"

    # FATCA
    doc.add_paragraph("Statut FATCA", style='SousSection')

    table_fatca = doc.add_table(rows=2, cols=2)
    table_fatca.style = 'Table Grid'

    hdr = table_fatca.rows[0].cells
    hdr[0].text = "TITULAIRE 1"
    hdr[1].text = "TITULAIRE 2"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    row = table_fatca.rows[1].cells
    row[0].text = "Êtes-vous une personne américaine (US-Person)*?\n{{T1_US_PERSON}}"
    row[1].text = "Êtes-vous une personne américaine (US-Person)*?\n{{T2_US_PERSON}}"

    note_fatca = """*Critères d'américanité : passeport US / carte verte / Lieu de naissance aux USA / adresse légale ou postale US / Numéros de téléphone américains commençant par +1 / Procuration à une personne ayant une adresse aux US / Virements réguliers vers des comptes domiciliés aux US – Si vous ne répondez pas, nous serons dans l'obligation de déclarer votre/vos contrat(s) aux autorités fiscales."""

    p = doc.add_paragraph(note_fatca)
    p.style.font.size = Pt(8)
    p.style.font.italic = True

    # SITUATION FAMILIALE
    doc.add_paragraph("Situation Familiale", style='SousSection')

    p = doc.add_paragraph()
    p.add_run("Situation matrimoniale: ").bold = True
    p.add_run("{{SITUATION_FAMILIALE}}")

    p = doc.add_paragraph()
    p.add_run("Date du mariage: ").bold = True
    p.add_run("{{DATE_MARIAGE}}   ")
    p.add_run("Contrat de mariage: ").bold = True
    p.add_run("{{CONTRAT_MARIAGE}}   ")
    p.add_run("Régime: ").bold = True
    p.add_run("{{REGIME_MATRIMONIAL}}")

    p = doc.add_paragraph()
    p.add_run("Date du PACS: ").bold = True
    p.add_run("{{DATE_PACS}}   ")
    p.add_run("Convention de PACS: ").bold = True
    p.add_run("{{CONVENTION_PACS}}   ")
    p.add_run("Régime: ").bold = True
    p.add_run("{{REGIME_PACS}}")

    p = doc.add_paragraph()
    p.add_run("Date du divorce: ").bold = True
    p.add_run("{{DATE_DIVORCE}}")

    # Donations
    doc.add_paragraph("Donations réalisées", style='SousSection')

    table_don = doc.add_table(rows=3, cols=2)
    table_don.style = 'Table Grid'

    hdr = table_don.rows[0].cells
    hdr[0].text = "TITULAIRE 1"
    hdr[1].text = "TITULAIRE 2"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    row = table_don.rows[1].cells
    row[0].text = "Donation entre époux:\n{{DONATION_ENTRE_EPOUX}}\nDate: {{DONATION_ENTRE_EPOUX_DATE}}\nMontant: {{DONATION_ENTRE_EPOUX_MONTANT}} €"
    row[1].text = "Donation entre époux:\n{{T2_DONATION_ENTRE_EPOUX}}\nDate: {{T2_DONATION_ENTRE_EPOUX_DATE}}\nMontant: {{T2_DONATION_ENTRE_EPOUX_MONTANT}} €"

    row = table_don.rows[2].cells
    row[0].text = "Donation aux enfants/petits-enfants:\n{{DONATION_ENFANTS}}\nDate: {{DONATION_ENFANTS_DATE}}\nMontant: {{DONATION_ENFANTS_MONTANT}} €"
    row[1].text = "Donation aux enfants/petits-enfants:\n{{T2_DONATION_ENFANTS}}\nDate: {{T2_DONATION_ENFANTS_DATE}}\nMontant: {{T2_DONATION_ENFANTS_MONTANT}} €"

    # Enfants
    doc.add_paragraph("Enfants", style='SousSection')

    p = doc.add_paragraph()
    p.add_run("Nombre d'enfant(s): ").bold = True
    p.add_run("{{NOMBRE_ENFANTS}}   ")
    p.add_run("Dont à charge: ").bold = True
    p.add_run("{{NOMBRE_ENFANTS_CHARGE}}")

    table_enf = doc.add_table(rows=1, cols=5)
    table_enf.style = 'Table Grid'

    hdr = table_enf.rows[0].cells
    headers = ["Prénom, Nom", "Date de naissance", "Lien de parenté", "À charge fiscalement?", "Ont-ils des enfants?"]
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'E6E6E6')

    # Placeholder pour les enfants (jusqu'à 4)
    for i in range(4):
        row = table_enf.add_row()
        row.cells[0].text = f"{{{{ENFANT_{i+1}_NOM}}}}"
        row.cells[1].text = f"{{{{ENFANT_{i+1}_DATE_NAISSANCE}}}}"
        row.cells[2].text = f"{{{{ENFANT_{i+1}_LIEN}}}}"
        row.cells[3].text = f"{{{{ENFANT_{i+1}_A_CHARGE}}}}"
        row.cells[4].text = f"{{{{ENFANT_{i+1}_A_ENFANTS}}}}"

    # Informations complémentaires
    doc.add_paragraph("Informations complémentaires", style='SousSection')
    p = doc.add_paragraph("(susceptibles d'influencer votre situation patrimoniale actuelle ou future – évolution professionnelle, enfant, déménagement, achats, etc...)")
    p.style.font.italic = True

    doc.add_paragraph("{{INFORMATIONS_COMPLEMENTAIRES}}")

    # SECTION 2: SITUATION FINANCIÈRE
    doc.add_paragraph("2 - SITUATION FINANCIÈRE ET PATRIMONIALE", style='TitreSection')

    doc.add_paragraph("Flux financiers", style='SousSection')

    table_fin = doc.add_table(rows=1, cols=2)
    table_fin.style = 'Table Grid'

    hdr = table_fin.rows[0].cells
    hdr[0].text = "TITULAIRE 1"
    hdr[1].text = "TITULAIRE 2"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    # Revenus
    row = table_fin.add_row()
    row.cells[0].text = "Revenus moyens annuels globaux du foyer fiscal\n(professionnels, fonciers, mobiliers, rentes, pensions...)\n\n{{REVENUS_ANNUELS_FOYER}}"
    row.cells[1].text = "Revenus moyens annuels globaux du foyer fiscal\n(professionnels, fonciers, mobiliers, rentes, pensions...)\n\n{{T2_REVENUS_ANNUELS_FOYER}}"

    # Patrimoine
    row = table_fin.add_row()
    row.cells[0].text = "Estimation globale de votre patrimoine (dettes exclues)\n\n{{PATRIMOINE_GLOBAL}}"
    row.cells[1].text = "Estimation globale de votre patrimoine (dettes exclues)\n\n{{T2_PATRIMOINE_GLOBAL}}"

    # Charges
    row = table_fin.add_row()
    row.cells[0].text = "Engagements financiers (dettes, charges, impôts...)\n\n{{CHARGES_ANNUELLES_POURCENT}}% des revenus\nsoit {{CHARGES_ANNUELLES_MONTANT}} €"
    row.cells[1].text = "Engagements financiers (dettes, charges, impôts...)\n\n{{T2_CHARGES_ANNUELLES_POURCENT}}% des revenus\nsoit {{T2_CHARGES_ANNUELLES_MONTANT}} €"

    # Répartition patrimoine
    row = table_fin.add_row()
    row.cells[0].text = "Répartition de votre patrimoine:\n• Actifs financiers: {{PATRIMOINE_FINANCIER_POURCENT}}%\n• Actifs immobiliers: {{PATRIMOINE_IMMOBILIER_POURCENT}}%\n• Actifs professionnels: {{PATRIMOINE_PROFESSIONNEL_POURCENT}}%\n• Autres: {{PATRIMOINE_AUTRES_POURCENT}}%"
    row.cells[1].text = "Répartition de votre patrimoine:\n• Actifs financiers: {{T2_PATRIMOINE_FINANCIER_POURCENT}}%\n• Actifs immobiliers: {{T2_PATRIMOINE_IMMOBILIER_POURCENT}}%\n• Actifs professionnels: {{T2_PATRIMOINE_PROFESSIONNEL_POURCENT}}%\n• Autres: {{T2_PATRIMOINE_AUTRES_POURCENT}}%"

    # Imposition
    row = table_fin.add_row()
    row.cells[0].text = "Imposition:\n• IR: {{IMPOT_REVENU}}\n• IFI: {{IMPOT_FORTUNE_IMMOBILIERE}}"
    row.cells[1].text = "Imposition:\n• IR: {{T2_IMPOT_REVENU}}\n• IFI: {{T2_IMPOT_FORTUNE_IMMOBILIERE}}"

    # Capacité d'épargne
    row = table_fin.add_row()
    row.cells[0].text = "Capacité d'épargne estimée:\n{{CAPACITE_EPARGNE_MENSUELLE}} € / mois"
    row.cells[1].text = "Capacité d'épargne estimée:\n{{T2_CAPACITE_EPARGNE_MENSUELLE}} € / mois"

    # ORIGINE DES FONDS
    doc.add_paragraph("Origine des fonds", style='SousSection')

    table_orig = doc.add_table(rows=4, cols=2)
    table_orig.style = 'Table Grid'

    row = table_orig.rows[0].cells
    row[0].text = "Nature des avoirs à investir"
    row[1].text = "{{ORIGINE_FONDS_NATURE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_orig.rows[1].cells
    row[0].text = "Montant prévu des avoirs à investir"
    row[1].text = "{{ORIGINE_FONDS_MONTANT_PREVU}} €"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_orig.rows[2].cells
    row[0].text = "Origine économique de ces avoirs"
    origins = """☐ Revenus: {{ORIGINE_ECONOMIQUE_REVENUS}}
☐ Épargne constituée: {{ORIGINE_ECONOMIQUE_EPARGNE}}
☐ Héritage/Donation/Succession: {{ORIGINE_ECONOMIQUE_HERITAGE}}
☐ Cession(s) d'actifs professionnels: {{ORIGINE_ECONOMIQUE_CESSION_PRO}}
☐ Cession(s) immobilière(s): {{ORIGINE_ECONOMIQUE_CESSION_IMMO}}
☐ Cession(s) mobilière(s): {{ORIGINE_ECONOMIQUE_CESSION_MOBILIERE}}
☐ Gains de jeu: {{ORIGINE_ECONOMIQUE_GAINS_JEU}}
☐ Assurance-vie: {{ORIGINE_ECONOMIQUE_ASSURANCE_VIE}}
☐ Autres: {{ORIGINE_ECONOMIQUE_AUTRES}}"""
    row[1].text = origins
    set_cell_shading(row[0], 'E6E6E6')

    row = table_orig.rows[3].cells
    row[0].text = "Provenance des fonds\n(établissement bancaire d'origine)"
    row[1].text = "{{ORIGINE_FONDS_PROVENANCE_ETABLISSEMENT}}"
    set_cell_shading(row[0], 'E6E6E6')

    # SECTION 3: DOCUMENTS À REMETTRE
    doc.add_paragraph("3 - DOCUMENTS À REMETTRE", style='TitreSection')

    table_docs = doc.add_table(rows=1, cols=3)
    table_docs.style = 'Table Grid'

    headers = ["Situation familiale", "Retraite – Prévoyance", "Patrimoine – Fiscalité"]
    hdr = table_docs.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'D9E2F3')

    row = table_docs.add_row()
    row.cells[0].text = """☐ Carte identité (recto/verso) ou passeport
☐ Carte identité conjoint/partenaire
☐ Contrat de mariage/Pacs/Jugement divorce
☐ Livret de famille
☐ Actes de donation
☐ Testament
☐ Justificatifs de domicile (< 3 mois)"""

    row.cells[1].text = """☐ Derniers bulletins de salaire
☐ Relevé de carrière
☐ Relevé de point retraite
☐ Contrat de retraite
☐ Contrats de prévoyance
☐ Contrats de retraite complémentaires"""

    row.cells[2].text = """☐ Dernier avis d'imposition
☐ Dernière déclaration de revenus
☐ Dernière déclaration IFI
☐ Tableaux d'amortissement des prêts"""

    # PATRIMOINE DÉTAILLÉ
    doc.add_paragraph("Patrimoine détaillé", style='SousSection')

    # Patrimoine financier
    doc.add_paragraph("Composition du patrimoine financier", style='Normal').runs[0].bold = True
    p = doc.add_paragraph("(disponibilités, comptes livret, comptes bancaires, compte titres, PEA, contrats d'assurance vie)")
    p.style.font.italic = True

    table_pf = doc.add_table(rows=1, cols=6)
    table_pf.style = 'Table Grid'
    headers = ["Désignation", "Organisme", "Valeur", "Détenteur", "Date souscription", "Remarques"]
    for i, h in enumerate(headers):
        table_pf.rows[0].cells[i].text = h
        table_pf.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_pf.rows[0].cells[i], 'E6E6E6')

    # Lignes placeholder
    for i in range(4):
        row = table_pf.add_row()
        row.cells[0].text = f"{{{{PATRIMOINE_FIN_{i+1}_DESIGNATION}}}}"
        row.cells[1].text = f"{{{{PATRIMOINE_FIN_{i+1}_ORGANISME}}}}"
        row.cells[2].text = f"{{{{PATRIMOINE_FIN_{i+1}_VALEUR}}}}"
        row.cells[3].text = f"{{{{PATRIMOINE_FIN_{i+1}_DETENTEUR}}}}"
        row.cells[4].text = f"{{{{PATRIMOINE_FIN_{i+1}_DATE}}}}"
        row.cells[5].text = f"{{{{PATRIMOINE_FIN_{i+1}_REMARQUES}}}}"

    # Patrimoine immobilier
    doc.add_paragraph()
    doc.add_paragraph("Composition du patrimoine immobilier", style='Normal').runs[0].bold = True
    p = doc.add_paragraph("(résidence principale, résidence secondaire, immobilier d'investissement, SCI, SCPI)")
    p.style.font.italic = True

    table_pi = doc.add_table(rows=1, cols=6)
    table_pi.style = 'Table Grid'
    headers = ["Désignation", "Détenteur", "Forme propriété", "Valeur actuelle", "Revenus", "Crédits en cours"]
    for i, h in enumerate(headers):
        table_pi.rows[0].cells[i].text = h
        table_pi.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_pi.rows[0].cells[i], 'E6E6E6')

    for i in range(4):
        row = table_pi.add_row()
        row.cells[0].text = f"{{{{PATRIMOINE_IMMO_{i+1}_DESIGNATION}}}}"
        row.cells[1].text = f"{{{{PATRIMOINE_IMMO_{i+1}_DETENTEUR}}}}"
        row.cells[2].text = f"{{{{PATRIMOINE_IMMO_{i+1}_FORME}}}}"
        row.cells[3].text = f"{{{{PATRIMOINE_IMMO_{i+1}_VALEUR}}}}"
        row.cells[4].text = f"{{{{PATRIMOINE_IMMO_{i+1}_REVENUS}}}}"
        row.cells[5].text = f"{{{{PATRIMOINE_IMMO_{i+1}_CREDITS}}}}"

    # Patrimoine professionnel
    doc.add_paragraph()
    doc.add_paragraph("Composition du patrimoine professionnel", style='Normal').runs[0].bold = True
    p = doc.add_paragraph("(parts sociales, clientèle, fonds de commerce)")
    p.style.font.italic = True

    table_pp = doc.add_table(rows=1, cols=5)
    table_pp.style = 'Table Grid'
    headers = ["Désignation", "Détenteur", "Valeur/Capital", "Charges", "Remarques"]
    for i, h in enumerate(headers):
        table_pp.rows[0].cells[i].text = h
        table_pp.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_pp.rows[0].cells[i], 'E6E6E6')

    for i in range(3):
        row = table_pp.add_row()
        row.cells[0].text = f"{{{{PATRIMOINE_PRO_{i+1}_DESIGNATION}}}}"
        row.cells[1].text = f"{{{{PATRIMOINE_PRO_{i+1}_DETENTEUR}}}}"
        row.cells[2].text = f"{{{{PATRIMOINE_PRO_{i+1}_VALEUR}}}}"
        row.cells[3].text = f"{{{{PATRIMOINE_PRO_{i+1}_CHARGES}}}}"
        row.cells[4].text = f"{{{{PATRIMOINE_PRO_{i+1}_REMARQUES}}}}"

    # Détails du passif
    doc.add_paragraph()
    doc.add_paragraph("Détails du passif", style='Normal').runs[0].bold = True
    p = doc.add_paragraph("(emprunts, prêts immobiliers, crédit consommation, dettes)")
    p.style.font.italic = True

    table_passif = doc.add_table(rows=1, cols=6)
    table_passif.style = 'Table Grid'
    headers = ["Objet", "Emprunteur", "Capital emprunté", "Capital restant dû", "Échéances", "Taux/Durée"]
    for i, h in enumerate(headers):
        table_passif.rows[0].cells[i].text = h
        table_passif.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_passif.rows[0].cells[i], 'E6E6E6')

    for i in range(3):
        row = table_passif.add_row()
        row.cells[0].text = f"{{{{PASSIF_{i+1}_OBJET}}}}"
        row.cells[1].text = f"{{{{PASSIF_{i+1}_EMPRUNTEUR}}}}"
        row.cells[2].text = f"{{{{PASSIF_{i+1}_CAPITAL}}}}"
        row.cells[3].text = f"{{{{PASSIF_{i+1}_RESTANT}}}}"
        row.cells[4].text = f"{{{{PASSIF_{i+1}_ECHEANCES}}}}"
        row.cells[5].text = f"{{{{PASSIF_{i+1}_TAUX}}}}"

    # Revenus détaillés
    doc.add_paragraph()
    doc.add_paragraph("Détails des revenus", style='Normal').runs[0].bold = True

    table_rev = doc.add_table(rows=1, cols=3)
    table_rev.style = 'Table Grid'
    headers = ["Nature des revenus", "Périodicité", "Montant"]
    for i, h in enumerate(headers):
        table_rev.rows[0].cells[i].text = h
        table_rev.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_rev.rows[0].cells[i], 'E6E6E6')

    for i in range(4):
        row = table_rev.add_row()
        row.cells[0].text = f"{{{{REVENU_{i+1}_NATURE}}}}"
        row.cells[1].text = f"{{{{REVENU_{i+1}_PERIODICITE}}}}"
        row.cells[2].text = f"{{{{REVENU_{i+1}_MONTANT}}}}"

    # Charges détaillées
    doc.add_paragraph()
    doc.add_paragraph("Détails des charges", style='Normal').runs[0].bold = True

    table_chg = doc.add_table(rows=1, cols=3)
    table_chg.style = 'Table Grid'
    headers = ["Nature des charges", "Périodicité", "Montant"]
    for i, h in enumerate(headers):
        table_chg.rows[0].cells[i].text = h
        table_chg.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table_chg.rows[0].cells[i], 'E6E6E6')

    for i in range(4):
        row = table_chg.add_row()
        row.cells[0].text = f"{{{{CHARGE_{i+1}_NATURE}}}}"
        row.cells[1].text = f"{{{{CHARGE_{i+1}_PERIODICITE}}}}"
        row.cells[2].text = f"{{{{CHARGE_{i+1}_MONTANT}}}}"

    # Retraite et prévoyance
    doc.add_paragraph("Retraite et prévoyance", style='SousSection')

    table_ret = doc.add_table(rows=1, cols=2)
    table_ret.style = 'Table Grid'

    hdr = table_ret.rows[0].cells
    hdr[0].text = "TITULAIRE 1"
    hdr[1].text = "TITULAIRE 2"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    row = table_ret.add_row()
    row.cells[0].text = """Connaissez-vous la date prévisionnelle de départ à la retraite?
{{T1_DATE_RETRAITE_PREVUE}}

Disposez-vous d'un bilan de retraite?
{{T1_BILAN_RETRAITE}}

Avez-vous mis en place des solutions de retraite complémentaire?
{{T1_RETRAITE_COMPLEMENTAIRE}}"""

    row.cells[1].text = """Connaissez-vous la date prévisionnelle de départ à la retraite?
{{T2_DATE_RETRAITE_PREVUE}}

Disposez-vous d'un bilan de retraite?
{{T2_BILAN_RETRAITE}}

Avez-vous mis en place des solutions de retraite complémentaire?
{{T2_RETRAITE_COMPLEMENTAIRE}}"""

    # LCB-FT / PPE
    doc.add_paragraph("Conformité LCB-FT", style='SousSection')

    p = doc.add_paragraph()
    p.add_run("Êtes-vous une Personne Politiquement Exposée (PPE)? ").bold = True
    p.add_run("{{LCB_FT_PPE}}")

    p = doc.add_paragraph()
    p.add_run("Si oui, fonction exercée: ").bold = True
    p.add_run("{{LCB_FT_PPE_FONCTION}}")

    p = doc.add_paragraph()
    p.add_run("Êtes-vous un membre de la famille d'une PPE? ").bold = True
    p.add_run("{{LCB_FT_PPE_FAMILLE}}")

    p = doc.add_paragraph()
    p.add_run("Vérification gel des avoirs effectuée: ").bold = True
    p.add_run("{{LCB_FT_GEL_AVOIRS_VERIFIE}}   Date: {{LCB_FT_GEL_AVOIRS_DATE_VERIFICATION}}")

    p = doc.add_paragraph()
    p.add_run("Niveau de risque LCB-FT: ").bold = True
    p.add_run("{{LCB_FT_NIVEAU_RISQUE}}")

    # SIGNATURE
    doc.add_paragraph()
    doc.add_paragraph()

    declaration = """Le client déclare :
• que les réponses à ce questionnaire sont exactes et sincères, qu'elles correspondent à sa situation actuelle et s'engage à informer de toute modification significative pouvant intervenir dans le futur.
• avoir reçu le document d'information préalable présentant le cabinet.
• être pleinement informé que le cabinet peut utiliser les informations demandées dans le présent document au titre de ses obligations légales et règlementaires en matière de lutte contre le blanchiment des capitaux et le financement du terrorisme."""

    doc.add_paragraph(declaration)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Fait à ").bold = True
    p.add_run("{{LIEU_SIGNATURE}} ")
    p.add_run("le ").bold = True
    p.add_run("{{DATE_SIGNATURE}} ")
    p.add_run("en ").bold = True
    p.add_run("{{NOMBRE_EXEMPLAIRES}} ")
    p.add_run("exemplaires.").bold = True

    # Tableau signatures
    doc.add_paragraph()
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.style = 'Table Grid'

    hdr = table_sig.rows[0].cells
    hdr[0].text = "Client(s)\n(signature(s) précédée(s) de la mention « Lu et approuvé »)"
    hdr[1].text = "Votre Conseiller"
    for cell in hdr:
        set_cell_shading(cell, 'E6E6E6')
        cell.paragraphs[0].runs[0].bold = True

    row = table_sig.rows[1].cells
    row[0].text = "\n\n\n{{SIGNATURE_CLIENT}}"
    row[1].text = "\n\n\n{{SIGNATURE_CONSEILLER}}"

    # MENTIONS LÉGALES
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph(MENTIONS_LEGALES)
    p.style.font.size = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


# ===========================================
# DOCUMENT 2: PROFIL DE RISQUE
# ===========================================

def create_profil_risque_document():
    """Crée le document Profil de Risque complet"""
    doc = create_styled_document()
    add_header_footer(doc, "Profil de Risque Financier")

    # TITRE PRINCIPAL
    title = doc.add_paragraph("QUESTIONNAIRE PROFIL DE RISQUE FINANCIER", style='TitrePrincipal')
    doc.add_paragraph("V.2025-03", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Référence client
    p = doc.add_paragraph()
    p.add_run("Client: ").bold = True
    p.add_run("{{NOM_COMPLET_T1}}")
    p.add_run("   N° Client: ").bold = True
    p.add_run("{{NUMERO_CLIENT}}")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Introduction
    intro = """La mise en place d'un profil de risques permet d'élaborer une stratégie financière en adéquation avec votre connaissance et votre expérience en matière financière, votre sensibilité au risque, votre capacité de perte, votre horizon d'investissement et vos objectifs.

Nous vous invitons à prendre connaissance de la typologie des profils de risques et leurs caractéristiques ci-dessous."""

    doc.add_paragraph(intro)

    # TYPOLOGIE DES PROFILS
    doc.add_paragraph("Typologie des profils de risques", style='TitreSection')

    p = doc.add_paragraph("Les informations ci-dessous sont à titre indicatif pour comprendre la différence entre les profils. Document non contractuel.")
    p.style.font.italic = True

    # Tableau des profils
    table_profils = doc.add_table(rows=5, cols=2)
    table_profils.style = 'Table Grid'

    hdr = table_profils.rows[0].cells
    hdr[0].text = "Profil"
    hdr[1].text = "Caractéristiques"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    profils = [
        ("SÉCURITAIRE", "Souhait : Sécurisation de vos investissements. La croissance de celui-ci est secondaire. Vous souhaitez prendre des risques très limités sur vos investissements. Une part importante du capital sera investi sur le fonds en euros.\nPerte en capital possible\nInconvénient : la valeur de vos investissements pourrait diminuer sur un an et même à plus long terme."),
        ("PRUDENT", "Souhait : sécurisation de vos investissements avec une faible prise de risques en vue d'une croissance de vos avoirs sur le long terme.\nPerte en capital possible\nInconvénient : la valeur de vos investissements pourrait diminuer sur un an et même à plus long terme."),
        ("ÉQUILIBRÉ", "Souhait : croissance de vos investissements sur le long terme, avec une prise de risques modérée.\nPerte en capital possible\nInconvénient : la valeur de vos investissements pourrait diminuer durant quelques années consécutives."),
        ("DYNAMIQUE", "Souhait : croissance de vos investissements sur le long terme avec prise de risques très élevée.\nPerte en capital possible\nInconvénient : la valeur de vos investissements pourrait très fortement diminuer durant quelques années consécutives.")
    ]

    for i, (profil, desc) in enumerate(profils):
        row = table_profils.rows[i+1].cells
        row[0].text = profil
        row[0].paragraphs[0].runs[0].bold = True
        row[1].text = desc

    # SECTION 1: CONNAISSANCE ET EXPÉRIENCE
    doc.add_paragraph("1. Connaissance et expérience des produits financiers", style='TitreSection')

    # Produits monétaires
    doc.add_paragraph("Produits monétaires et fonds euros", style='SousSection')
    p = doc.add_paragraph("(Livret A, PEL, fonds euros, assurance vie)")
    p.style.font.italic = True

    table_mon = doc.add_table(rows=5, cols=2)
    table_mon.style = 'Table Grid'

    row = table_mon.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_MONETAIRES_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_mon.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_MONETAIRES_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_mon.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_MONETAIRES_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_mon.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_MONETAIRES_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_mon.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: A moyen et long terme, les produits monétaires offrent une espérance de rendement inférieure à celle de certains actifs risqués.
Réponse: {{KYC_MONETAIRES_Q1}}

Q2: A moyen et long terme, les produits monétaires font courir un risque de perte en capital plus limité que celui des actifs risqués.
Réponse: {{KYC_MONETAIRES_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Obligations
    doc.add_paragraph("Obligations et fonds obligataires", style='SousSection')
    p = doc.add_paragraph("(titre de créance, OPC obligataires)")
    p.style.font.italic = True

    table_obl = doc.add_table(rows=5, cols=2)
    table_obl.style = 'Table Grid'

    row = table_obl.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_OBLIGATIONS_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_obl.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_OBLIGATIONS_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_obl.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_OBLIGATIONS_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_obl.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_OBLIGATIONS_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_obl.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: Plus la santé financière d'un émetteur est saine, plus le coupon versé sera élevé.
Réponse: {{KYC_OBLIGATIONS_Q1}}

Q2: Un investissement sur ce type de placement présente un risque de perte en capital en raison du risque de défaut de l'émetteur.
Réponse: {{KYC_OBLIGATIONS_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Actions
    doc.add_paragraph("Actions et fonds actions", style='SousSection')
    p = doc.add_paragraph("(actions, fonds en actions, OPC actions)")
    p.style.font.italic = True

    table_act = doc.add_table(rows=5, cols=2)
    table_act.style = 'Table Grid'

    row = table_act.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_ACTIONS_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_act.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_ACTIONS_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_act.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_ACTIONS_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_act.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_ACTIONS_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_act.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: La valeur d'une action peut chuter à 0 EUR.
Réponse: {{KYC_ACTIONS_Q1}}

Q2: Un investissement sur ce type de placement présente un risque de perte en capital en raison du risque de défaut de l'émetteur.
Réponse: {{KYC_ACTIONS_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # SCPI
    doc.add_paragraph("SCPI", style='SousSection')

    table_scpi = doc.add_table(rows=5, cols=2)
    table_scpi.style = 'Table Grid'

    row = table_scpi.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_SCPI_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_scpi.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_SCPI_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_scpi.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_SCPI_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_scpi.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_SCPI_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_scpi.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: L'investissement permet de mutualiser les risques.
Réponse: {{KYC_SCPI_Q1}}

Q2: Les investisseurs qui souhaitent vendre leurs parts de SCPI doivent eux-mêmes trouver un nouvel acquéreur.
Réponse: {{KYC_SCPI_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Private Equity
    doc.add_paragraph("Private Equity (FCPI, FCPR, FIP)", style='SousSection')

    table_pe = doc.add_table(rows=5, cols=2)
    table_pe.style = 'Table Grid'

    row = table_pe.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_PE_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_pe.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_PE_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_pe.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_PE_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_pe.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_PE_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_pe.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: Investir dans ce type de produits est risqué et nécessite de conserver les parts pendant plus de 8 ans.
Réponse: {{KYC_PE_Q1}}

Q2: Un investissement sur ce type de placement présente un risque de perte en capital en raison du risque de défaut de l'émetteur.
Réponse: {{KYC_PE_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # ETF
    doc.add_paragraph("Fonds indiciels (Trackers ou ETF)", style='SousSection')

    table_etf = doc.add_table(rows=5, cols=2)
    table_etf.style = 'Table Grid'

    row = table_etf.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_ETF_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_etf.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_ETF_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_etf.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_ETF_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_etf.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_ETF_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_etf.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: Ce type d'instrument réplique exactement l'indice sur lequel il est adossé.
Réponse: {{KYC_ETF_Q1}}

Q2: Je peux acheter ou vendre ce type d'instrument à tout moment de la journée, comme une action cotée.
Réponse: {{KYC_ETF_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Produits dérivés
    doc.add_paragraph("Produits dérivés (Options, Futures, Warrants, Certificats)", style='SousSection')

    table_der = doc.add_table(rows=5, cols=2)
    table_der.style = 'Table Grid'

    row = table_der.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_DERIVES_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_der.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_DERIVES_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_der.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_DERIVES_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_der.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_DERIVES_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_der.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: L'utilisation de ce type d'instrument peut augmenter mon risque de perte en capital.
Réponse: {{KYC_DERIVES_Q1}}

Q2: Il est possible d'utiliser ce type d'instrument pour couvrir un risque spécifique dans un portefeuille.
Réponse: {{KYC_DERIVES_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Produits structurés
    doc.add_paragraph("Produits structurés", style='SousSection')

    table_str = doc.add_table(rows=5, cols=2)
    table_str.style = 'Table Grid'

    row = table_str.rows[0].cells
    row[0].text = "Détention"
    row[1].text = "{{KYC_STRUCTURES_DETENTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_str.rows[1].cells
    row[0].text = "Nombre d'opérations réalisées par an"
    row[1].text = "{{KYC_STRUCTURES_OPERATIONS}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_str.rows[2].cells
    row[0].text = "Durée de détention"
    row[1].text = "{{KYC_STRUCTURES_DUREE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_str.rows[3].cells
    row[0].text = "Volume des opérations"
    row[1].text = "{{KYC_STRUCTURES_VOLUME}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_str.rows[4].cells
    row[0].text = "Questions de connaissances"
    row[1].text = """Q1: La valeur d'un produit structuré est-elle garantie en cas de rachat avant son échéance?
Réponse: {{KYC_STRUCTURES_Q1}}

Q2: Un produit structuré présente-t-il un risque de perte en capital au cours de vie et à l'échéance?
Réponse: {{KYC_STRUCTURES_Q2}}"""
    set_cell_shading(row[0], 'E6E6E6')

    # Gestion du portefeuille
    doc.add_paragraph("Gestion du portefeuille", style='SousSection')

    table_gest = doc.add_table(rows=4, cols=2)
    table_gest.style = 'Table Grid'

    questions_gest = [
        ("Avez-vous (ou avez-vous déjà eu) un portefeuille géré sous mandat?", "{{KYC_PORTEFEUILLE_MANDAT}}"),
        ("Gérez-vous (ou avez-vous déjà géré) vous-même votre portefeuille?", "{{KYC_PORTEFEUILLE_GESTION_PERSONNELLE}}"),
        ("Gérez-vous avec l'aide d'un conseiller votre portefeuille?", "{{KYC_PORTEFEUILLE_GESTION_CONSEILLER}}"),
        ("Avez-vous exercé pendant au moins un an, dans le secteur financier, une position professionnelle exigeant une connaissance des investissements?", "{{KYC_PORTEFEUILLE_EXPERIENCE_PRO}}")
    ]

    for i, (q, placeholder) in enumerate(questions_gest):
        row = table_gest.rows[i].cells
        row[0].text = q
        row[1].text = placeholder
        set_cell_shading(row[0], 'E6E6E6')

    # Culture financière
    doc.add_paragraph("Culture financière", style='SousSection')

    table_cult = doc.add_table(rows=3, cols=2)
    table_cult.style = 'Table Grid'

    questions_cult = [
        ("Lisez-vous la presse ou l'actualité financière spécialisée?", "{{KYC_CULTURE_PRESSE_FINANCIERE}}"),
        ("Regardez-vous régulièrement les cours de la Bourse?", "{{KYC_CULTURE_SUIVI_BOURSE}}"),
        ("Regardez-vous au moins tous les mois vos relevés bancaires?", "{{KYC_CULTURE_RELEVES_BANCAIRES}}")
    ]

    for i, (q, placeholder) in enumerate(questions_cult):
        row = table_cult.rows[i].cells
        row[0].text = q
        row[1].text = placeholder
        set_cell_shading(row[0], 'E6E6E6')

    # SECTION 2: OBJECTIFS ET RISQUES
    doc.add_paragraph("2. Objectifs d'investissement et tolérance au risque", style='TitreSection')

    # Objectifs
    doc.add_paragraph("Objectifs d'investissement", style='SousSection')
    p = doc.add_paragraph("(plusieurs choix possibles à classer par ordre de priorité)")
    p.style.font.italic = True

    table_obj = doc.add_table(rows=7, cols=2)
    table_obj.style = 'Table Grid'

    objectifs = [
        ("Préservation du capital", "{{OBJECTIF_PRESERVATION}}"),
        ("Valorisation de capital", "{{OBJECTIF_VALORISATION}}"),
        ("Diversification des actifs détenus", "{{OBJECTIF_DIVERSIFICATION}}"),
        ("Recherche de revenus", "{{OBJECTIF_REVENUS}}"),
        ("Transmission", "{{OBJECTIF_TRANSMISSION}}"),
        ("Optimisation fiscale", "{{OBJECTIF_FISCAL}}"),
        ("Autres", "{{OBJECTIF_AUTRES}}")
    ]

    for i, (obj, placeholder) in enumerate(objectifs):
        row = table_obj.rows[i].cells
        row[0].text = obj
        row[1].text = f"Priorité: {placeholder}"
        set_cell_shading(row[0], 'E6E6E6')

    # Tolérance au risque
    doc.add_paragraph("Tolérance au risque", style='SousSection')

    tolerance_desc = """Les courbes ci-dessous présentent les fluctuations des cours de 3 placements hypothétiques sur 10 ans. Quel est le placement avec lequel vous seriez le plus à l'aise?

• Placement A (Risque faible): tout en privilégiant la protection de votre capital sur la durée, vous acceptez une diversification partielle de vos investissements sur des actifs plus volatils et donc plus risqués

• Placement B (Risque moyen): en acceptant de diversifier significativement vos actifs sur des supports à forte volatilité pouvant entraîner une perte en capital, vous êtes à la recherche d'une valorisation importante de votre investissement

• Placement C (Risque élevé): en contrepartie d'une perte potentielle partielle, voire totale, de votre épargne, vous cherchez avant tout à maximiser la performance de votre investissement"""

    doc.add_paragraph(tolerance_desc)

    p = doc.add_paragraph()
    p.add_run("Votre choix: ").bold = True
    p.add_run("{{TOLERANCE_RISQUE}}")

    # Expérience de perte
    doc.add_paragraph("Expérience et réaction face aux pertes", style='SousSection')

    table_perte = doc.add_table(rows=4, cols=2)
    table_perte.style = 'Table Grid'

    row = table_perte.rows[0].cells
    row[0].text = "Avez-vous déjà effectué un investissement qui a connu une baisse de valeur?"
    row[1].text = "{{EXPERIENCE_PERTE}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_perte.rows[1].cells
    row[0].text = "Si oui, niveau de baisse"
    row[1].text = "{{EXPERIENCE_PERTE_NIVEAU}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_perte.rows[2].cells
    row[0].text = "Quelle a été (ou serait) votre réaction par rapport à cette baisse?"
    row[1].text = """{{REACTION_PERTE}}
☐ Investir de nouveau pour profiter des opportunités
☐ Tout vendre pour réinvestir dans des supports moins risqués
☐ Vendre seulement une partie
☐ Ne rien changer"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_perte.rows[3].cells
    row[0].text = "Si la valeur de votre investissement augmente de 20%, comment réagissez-vous?"
    row[1].text = """{{REACTION_GAIN}}
☐ Je conserve ma position
☐ Je réinvestis un montant inférieur ou égal au montant initial
☐ Je réinvestis un montant supérieur au montant initial"""
    set_cell_shading(row[0], 'E6E6E6')

    # Horizon et liquidité
    doc.add_paragraph("Horizon de placement et liquidité", style='SousSection')

    table_hor = doc.add_table(rows=4, cols=2)
    table_hor.style = 'Table Grid'

    row = table_hor.rows[0].cells
    row[0].text = "Sur quel horizon souhaitez-vous réaliser ce placement?"
    row[1].text = """{{HORIZON_PLACEMENT}}
☐ < 1 an    ☐ Entre 1 et 3 ans    ☐ Entre 3 et 5 ans    ☐ > 5 ans"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_hor.rows[1].cells
    row[0].text = "Le critère de liquidité est-il important dans le cadre de ce placement?"
    row[1].text = """{{LIQUIDITE_IMPORTANTE}}
☐ Oui, je dois pouvoir disposer de mon argent à tout moment
☐ Non, je dispose de liquidités accessibles rapidement"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_hor.rows[2].cells
    row[0].text = "Pour le placement envisagé, quel niveau de pertes maximales êtes-vous prêt à subir?"
    row[1].text = """{{PERTES_MAXIMALES_ACCEPTABLES}}
☐ Aucune    ☐ Max 10%    ☐ Max 25%    ☐ Max 50%    ☐ Jusqu'à 100%"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_hor.rows[3].cells
    row[0].text = "Quel pourcentage de votre patrimoine total représente le montant que vous envisagez d'investir?"
    row[1].text = """{{POURCENTAGE_PATRIMOINE_INVESTI}}
☐ Moins de 10%    ☐ Entre 10% et 25%    ☐ Entre 25% et 50%    ☐ Entre 50% et 75%    ☐ Plus de 75%"""
    set_cell_shading(row[0], 'E6E6E6')

    # SECTION 3: DURABILITÉ ESG
    doc.add_paragraph("3. Investissements durables (ESG)", style='TitreSection')

    p = doc.add_paragraph()
    p.add_run("Souhaitez-vous intégrer des critères de durabilité dans vos choix d'investissements? ").bold = True
    p.add_run("{{DURABILITE_SOUHAIT}}")

    doc.add_paragraph("Si oui, veuillez répondre aux questions suivantes:", style='SousSection')

    table_esg = doc.add_table(rows=5, cols=2)
    table_esg.style = 'Table Grid'

    row = table_esg.rows[0].cells
    row[0].text = "Préférences pour la Taxonomie Européenne\n(% de votre investissement)"
    row[1].text = """{{DURABILITE_TAXONOMIE_POURCENT}}
☐ ≥5%    ☐ ≥25%    ☐ ≥50%    ☐ Aucun"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_esg.rows[1].cells
    row[0].text = "Souhaits en termes d'investissements durables\n(% de votre investissement)"
    row[1].text = """{{DURABILITE_INVESTISSEMENTS_POURCENT}}
☐ ≥5%    ☐ ≥25%    ☐ ≥50%    ☐ Aucun"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_esg.rows[2].cells
    row[0].text = "Souhaitez-vous sélectionner vos investissements en fonction de leur impact sur les facteurs de durabilité?"
    row[1].text = "{{DURABILITE_IMPACT_SELECTION}}"
    set_cell_shading(row[0], 'E6E6E6')

    row = table_esg.rows[3].cells
    row[0].text = "Effets sur lesquels vous souhaitez minimiser les incidences négatives"
    row[1].text = """{{DURABILITE_CRITERES}}
☐ Gaz à effet de serre
☐ Impact sur la biodiversité
☐ Emissions polluantes dans l'eau
☐ Génération des déchets dangereux
☐ Inefficacité énergétique (immobilier)
☐ Respect des normes internationales (OCDE, Nations unies)
☐ Égalité de rémunération (Homme/Femme)
☐ Diversité des genres au sein des conseils
☐ Exposition aux armes controversées"""
    set_cell_shading(row[0], 'E6E6E6')

    row = table_esg.rows[4].cells
    row[0].text = "Note importante"
    row[1].text = "Actuellement, peu de produits financiers ESG ont une part significative d'investissements durables alignée à la Taxonomie tels que définis par la réglementation européenne. Si votre souhait est de sélectionner les produits financiers selon ce critère, il sera difficile de vous apporter pleine satisfaction. Néanmoins, l'amélioration de l'offre dans le futur permettra de se rapprocher progressivement de vos souhaits."
    set_cell_shading(row[0], 'FFF2CC')

    # SECTION 4: CONCLUSION - PROFIL CALCULÉ
    doc.add_paragraph("Réservé au conseiller", style='TitreSection')

    doc.add_paragraph("Conclusion sur votre profil de risque", style='SousSection')

    conclusion = """Votre conseiller vous indique que, selon vos réponses au questionnaire, le degré de risques que vous êtes susceptible de tolérer est:"""
    doc.add_paragraph(conclusion)

    table_concl = doc.add_table(rows=1, cols=2)
    table_concl.style = 'Table Grid'

    row = table_concl.rows[0].cells
    row[0].text = "PROFIL DE RISQUE DÉTERMINÉ"
    row[1].text = """{{PROFIL_RISQUE_CALCULE}}

☐ SÉCURITAIRE, avec une part d'actifs à risque élevé de {{PART_RISQUE_SECURITAIRE}}% maximum
☐ PRUDENT, avec une part d'actifs à risque élevé de {{PART_RISQUE_PRUDENT}}% maximum
☐ ÉQUILIBRÉ, avec une part d'actifs à risque élevé de {{PART_RISQUE_EQUILIBRE}}% maximum
☐ DYNAMIQUE, avec une part d'actifs à risque élevé de {{PART_RISQUE_DYNAMIQUE}}% maximum

Score calculé: {{PROFIL_RISQUE_SCORE}} / 100
Date du calcul: {{PROFIL_RISQUE_DATE_CALCUL}}"""
    set_cell_shading(row[0], 'D9E2F3')
    row[0].paragraphs[0].runs[0].bold = True

    # Commentaire conseiller
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Commentaire du conseiller: ").bold = True
    doc.add_paragraph("{{PROFIL_COMMENTAIRE}}")

    # DÉCLARATION CLIENT
    doc.add_paragraph()
    doc.add_paragraph()

    declaration = """Le client déclare :

• que les réponses à ce questionnaire sont exactes et sincères, qu'elles correspondent à sa situation actuelle et s'engage à informer de toute modification significative pouvant intervenir dans le futur.
• avoir reçu le document d'information préalable présentant le cabinet.
• être pleinement informé que le cabinet peut utiliser les informations demandées dans le présent document au titre de ses obligations légales et règlementaires en matière de lutte contre le blanchiment des capitaux et le financement du terrorisme.

Il est également rappelé, conformément aux dispositions de l'article 325-8 du Règlement général de l'AMF, qu'un conseiller en investissements financiers qui n'obtiendrait pas les informations requises au titre du présent questionnaire, doit s'abstenir de recommander des instruments financiers."""

    doc.add_paragraph(declaration)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run("Fait à ").bold = True
    p.add_run("{{LIEU_SIGNATURE}} ")
    p.add_run("le ").bold = True
    p.add_run("{{DATE_SIGNATURE}} ")
    p.add_run("en ").bold = True
    p.add_run("{{NOMBRE_EXEMPLAIRES}} ")
    p.add_run("exemplaires.").bold = True

    # Tableau signatures
    doc.add_paragraph()
    table_sig = doc.add_table(rows=2, cols=2)
    table_sig.style = 'Table Grid'

    hdr = table_sig.rows[0].cells
    hdr[0].text = "Client(s)\n(signature(s) précédée(s) de la mention « Lu et approuvé »)"
    hdr[1].text = "Votre Conseiller"
    for cell in hdr:
        set_cell_shading(cell, 'E6E6E6')
        cell.paragraphs[0].runs[0].bold = True

    row = table_sig.rows[1].cells
    row[0].text = "\n\n\n{{SIGNATURE_CLIENT}}"
    row[1].text = "\n\n\n{{SIGNATURE_CONSEILLER}}"

    # MENTIONS LÉGALES
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph(MENTIONS_LEGALES)
    p.style.font.size = Pt(7)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


# ===========================================
# GÉNÉRATION DES DOCUMENTS
# ===========================================

if __name__ == "__main__":
    output_dir = os.path.dirname(os.path.abspath(__file__))

    print("Génération des templates DOCX v2...")

    # Document 1: QCC
    print("  - Création QCC_V2_TEMPLATE.docx...")
    doc_qcc = create_qcc_document()
    doc_qcc.save(os.path.join(output_dir, "QCC_V2_TEMPLATE.docx"))
    print("    ✓ QCC_V2_TEMPLATE.docx créé")

    # Document 2: Profil de Risque
    print("  - Création PROFIL_RISQUE_V2_TEMPLATE.docx...")
    doc_pr = create_profil_risque_document()
    doc_pr.save(os.path.join(output_dir, "PROFIL_RISQUE_V2_TEMPLATE.docx"))
    print("    ✓ PROFIL_RISQUE_V2_TEMPLATE.docx créé")

    print("\nTerminé! Les templates sont dans:", output_dir)

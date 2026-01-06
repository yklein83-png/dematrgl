#!/usr/bin/env python3
"""
Génère le template Convention RTO V3 avec placeholders standardisés
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_rto_template():
    doc = Document()

    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === TITRE ===
    title = doc.add_heading('CONVENTION DE RÉCEPTION ET TRANSMISSION D\'ORDRES', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === PARTIES ===
    doc.add_heading('ENTRE LES SOUSSIGNÉS', level=1)

    # --- Client ---
    p = doc.add_paragraph()
    p.add_run('Le Client :').bold = True

    p = doc.add_paragraph()
    p.add_run('{{CLIENT_CIVILITE}} {{CLIENT_PRENOM}} {{CLIENT_NOM}}').bold = True

    doc.add_paragraph('Né(e) le {{CLIENT_DATE_NAISSANCE}} à {{CLIENT_LIEU_NAISSANCE}}')
    doc.add_paragraph('Nationalité : {{CLIENT_NATIONALITE}}')
    doc.add_paragraph('Demeurant : {{CLIENT_ADRESSE}}')
    doc.add_paragraph('{{CLIENT_CODE_POSTAL}} {{CLIENT_VILLE}}')
    doc.add_paragraph('Profession : {{CLIENT_PROFESSION}}')
    doc.add_paragraph('Situation familiale : {{CLIENT_SITUATION_FAMILIALE}}')
    doc.add_paragraph('Régime matrimonial : {{CLIENT_REGIME_MATRIMONIAL}}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ci-après désigné « le Client », d\'une part,').italic = True

    doc.add_paragraph()
    doc.add_paragraph('Et,')
    doc.add_paragraph()

    # --- Cabinet ---
    p = doc.add_paragraph()
    p.add_run('Le Cabinet :').bold = True

    p = doc.add_paragraph()
    p.add_run('{{CABINET_NOM}}').bold = True

    doc.add_paragraph('{{CABINET_FORME_JURIDIQUE}}, au capital de {{CABINET_CAPITAL}}')
    doc.add_paragraph('Siège social : {{CABINET_ADRESSE}}')
    doc.add_paragraph('{{CABINET_CODE_POSTAL}} {{CABINET_VILLE}}')
    doc.add_paragraph('RCS {{CABINET_VILLE_RCS}} n° {{CABINET_NUM_RCS}}')
    doc.add_paragraph('N° ORIAS : {{CABINET_NUM_ORIAS}} (www.orias.fr)')
    doc.add_paragraph('Enregistré en qualité de Conseil en Investissements Financiers')
    doc.add_paragraph('Membre de {{CABINET_ASSOCIATION_CIF}}, association agréée par l\'AMF')
    doc.add_paragraph('N° d\'enregistrement CIF : {{CABINET_NUM_CIF}}')

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Représenté par {{CABINET_REPRESENTANT_CIVILITE}} {{CABINET_REPRESENTANT_PRENOM}} {{CABINET_REPRESENTANT_NOM}}, en qualité de {{CABINET_REPRESENTANT_QUALITE}}').italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Ci-après désigné « le CIF », d\'autre part,').italic = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Et ensemble désignés « les Parties ».').italic = True

    doc.add_paragraph()

    # === PRÉAMBULE ===
    doc.add_heading('PRÉAMBULE', level=1)

    doc.add_paragraph(
        'La présente convention a pour objet de définir les conditions dans lesquelles le CIF '
        'pourra recevoir et transmettre des ordres pour le compte du Client, conformément aux '
        'dispositions de l\'article 325-13 du Règlement général de l\'AMF relatif à la réception '
        'et transmission d\'ordres sur OPC par les CIF.'
    )

    doc.add_paragraph()
    doc.add_heading('IL A ÉTÉ CONVENU CE QUI SUIT :', level=1)

    # === ARTICLES ===

    # Article 1
    doc.add_heading('Article 1 : Objet et périmètre', level=2)
    doc.add_paragraph(
        'Le Client est titulaire du(des) compte(s) suivant(s) :'
    )
    doc.add_paragraph('• {{COMPTE_TYPE_1}} n° {{COMPTE_NUMERO_1}} ouvert auprès de {{COMPTE_ETABLISSEMENT_1}}')
    doc.add_paragraph('{{#SI_COMPTE_2}}• {{COMPTE_TYPE_2}} n° {{COMPTE_NUMERO_2}} ouvert auprès de {{COMPTE_ETABLISSEMENT_2}}{{/SI_COMPTE_2}}')

    doc.add_paragraph()
    doc.add_paragraph(
        'Seul le Client est habilité à mouvementer ce(s) portefeuille(s). Le CIF ne bénéficie '
        'd\'aucun mandat de gestion et ne peut procéder à aucune opération de gestion ou '
        'd\'arbitrage de sa propre initiative, ces opérations relevant du seul pouvoir du Client.'
    )

    # Article 2
    doc.add_heading('Article 2 : Déclarations du Client', level=2)
    doc.add_paragraph('Le Client déclare :')
    doc.add_paragraph('• Avoir pris connaissance du Document d\'Entrée en Relation (DER) du CIF ;')
    doc.add_paragraph('• Avoir complété le Questionnaire de Connaissance Client (QCC) ;')
    doc.add_paragraph('• Avoir les connaissances et l\'aptitude nécessaires pour réaliser les opérations pour lesquelles il donnera un ordre ;')
    doc.add_paragraph('• S\'engager à informer le CIF de toute modification de sa situation patrimoniale ou personnelle pouvant influer sur le conseil qui lui serait donné.')

    # Article 3
    doc.add_heading('Article 3 : Rédaction et transmission des ordres', level=2)
    doc.add_paragraph(
        'Le Client devra utiliser de préférence les modèles d\'ordres fournis par le CIF. '
        'À défaut, tout ordre sur papier libre devra clairement indiquer :'
    )
    doc.add_paragraph('• L\'identité du Client et la référence de son compte ;')
    doc.add_paragraph('• La nature de l\'opération (souscription, rachat, arbitrage) ;')
    doc.add_paragraph('• Le nom de l\'OPC concerné accompagné de son code ISIN ;')
    doc.add_paragraph('• Le montant ou le nombre de parts.')

    doc.add_paragraph()
    doc.add_paragraph(
        'L\'ordre devra être adressé au Cabinet par courrier postal, télécopie ou courrier '
        'électronique à l\'adresse suivante : {{CABINET_EMAIL}}'
    )

    # Article 4
    doc.add_heading('Article 4 : Réception de l\'ordre', level=2)
    doc.add_paragraph(
        'Cette prestation est exclusivement réservée aux ordres portant sur des OPC '
        '(Organismes de Placements Collectifs) tels que définis à l\'article L.214-1 du Code '
        'monétaire et financier : OPCVM, FIA, SCPI, OPCI, etc.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Le Cabinet CIF horodatera l\'ordre dès sa réception pour confirmer sa prise en compte.'
    )

    # Article 5
    doc.add_heading('Article 5 : Transmission de l\'ordre', level=2)
    doc.add_paragraph(
        'À réception de l\'ordre émis par le Client, le Cabinet CIF transmettra l\'ordre à '
        'l\'établissement teneur de compte dans un délai maximal de 3 jours ouvrables, selon '
        'les modalités propres à chaque établissement.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Le Cabinet CIF ne peut être tenu responsable :'
    )
    doc.add_paragraph('• De toute erreur ou manquement d\'exécution commis par l\'établissement teneur de compte ;')
    doc.add_paragraph('• En cas d\'interruption des moyens de transmission ;')
    doc.add_paragraph('• Si les conditions de marché ou réglementaires ne permettent pas l\'exécution de l\'ordre.')

    # Article 6
    doc.add_heading('Article 6 : Confirmation de l\'exécution', level=2)
    doc.add_paragraph(
        'L\'avis d\'opéré est adressé directement au Client par l\'établissement teneur de compte '
        'ou la société de gestion. Il fait foi de l\'exécution de l\'ordre.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Si dans un délai de six jours ouvrables le Client n\'a pas reçu d\'avis d\'opéré, '
        'il peut se rapprocher du CIF pour vérification.'
    )

    # Article 7
    doc.add_heading('Article 7 : Rémunération', level=2)
    doc.add_paragraph(
        'En aucun cas le Client n\'a à s\'acquitter de droits d\'entrée supérieurs ou de frais '
        'supplémentaires à ce qui figure dans le Document d\'Information Clé (DIC) qui lui aura '
        'été remis préalablement à toute souscription.'
    )
    doc.add_paragraph()
    doc.add_paragraph('Mode de rémunération du CIF : {{MISSION_REMUNERATION_MODE}}')
    doc.add_paragraph('{{#SI_HONORAIRES}}Honoraires : {{MISSION_HONORAIRES_MONTANT}} XPF - {{MISSION_HONORAIRES_DESCRIPTION}}{{/SI_HONORAIRES}}')

    # Article 8
    doc.add_heading('Article 8 : Durée et résiliation', level=2)
    doc.add_paragraph(
        'La présente convention est conclue pour une durée d\'un an à compter de sa signature. '
        'Elle est renouvelée par tacite reconduction pour des périodes successives d\'un an.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Elle peut être résiliée à tout moment par l\'une ou l\'autre des Parties par lettre '
        'recommandée avec accusé de réception, moyennant un préavis d\'un mois.'
    )

    # Article 9
    doc.add_heading('Article 9 : Réclamations et médiation', level=2)
    doc.add_paragraph(
        'Pour toute réclamation, le Client peut s\'adresser au CIF par courrier à l\'adresse '
        'du siège social ou par email à {{CABINET_EMAIL}}.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'En cas de désaccord persistant, le Client peut saisir gratuitement le médiateur de '
        'l\'AMF (Autorité des Marchés Financiers) : {{CABINET_MEDIATEUR}}'
    )

    # Article 10
    doc.add_heading('Article 10 : Droit applicable', level=2)
    doc.add_paragraph(
        'La présente convention est soumise au droit français. Tout litige sera de la '
        'compétence exclusive des tribunaux du ressort du siège social du CIF.'
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # === SIGNATURES ===
    doc.add_heading('SIGNATURES', level=1)

    p = doc.add_paragraph()
    p.add_run('Fait à {{SIGNATURE_LIEU}}, le {{SIGNATURE_DATE}}').bold = True

    doc.add_paragraph('En {{SIGNATURE_NOMBRE_EXEMPLAIRES}} exemplaires originaux.')

    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau signatures
    table = doc.add_table(rows=2, cols=2)
    table.autofit = True

    # En-têtes
    cell_client = table.rows[0].cells[0]
    cell_client.text = 'Le Client'
    cell_client.paragraphs[0].runs[0].bold = True

    cell_cif = table.rows[0].cells[1]
    cell_cif.text = 'Le CIF'
    cell_cif.paragraphs[0].runs[0].bold = True

    # Signatures
    table.rows[1].cells[0].text = '{{CLIENT_CIVILITE}} {{CLIENT_PRENOM}} {{CLIENT_NOM}}\n\n\n\n(Signature précédée de la mention "Lu et approuvé")'
    table.rows[1].cells[1].text = '{{CABINET_REPRESENTANT_CIVILITE}} {{CABINET_REPRESENTANT_PRENOM}} {{CABINET_REPRESENTANT_NOM}}\nPour {{CABINET_NOM}}\n\n\n(Signature et cachet)'

    # Sauvegarder
    output_path = os.path.join(os.path.dirname(__file__), 'CONVENTION_RTO_V3_TEMPLATE.docx')
    doc.save(output_path)
    print(f"Template créé : {output_path}")

    # Liste des placeholders utilisés
    placeholders = [
        "CLIENT_CIVILITE", "CLIENT_PRENOM", "CLIENT_NOM",
        "CLIENT_DATE_NAISSANCE", "CLIENT_LIEU_NAISSANCE", "CLIENT_NATIONALITE",
        "CLIENT_ADRESSE", "CLIENT_CODE_POSTAL", "CLIENT_VILLE",
        "CLIENT_PROFESSION", "CLIENT_SITUATION_FAMILIALE", "CLIENT_REGIME_MATRIMONIAL",
        "CABINET_NOM", "CABINET_FORME_JURIDIQUE", "CABINET_CAPITAL",
        "CABINET_ADRESSE", "CABINET_CODE_POSTAL", "CABINET_VILLE",
        "CABINET_VILLE_RCS", "CABINET_NUM_RCS", "CABINET_NUM_ORIAS",
        "CABINET_ASSOCIATION_CIF", "CABINET_NUM_CIF", "CABINET_EMAIL", "CABINET_MEDIATEUR",
        "CABINET_REPRESENTANT_CIVILITE", "CABINET_REPRESENTANT_PRENOM",
        "CABINET_REPRESENTANT_NOM", "CABINET_REPRESENTANT_QUALITE",
        "COMPTE_TYPE_1", "COMPTE_NUMERO_1", "COMPTE_ETABLISSEMENT_1",
        "COMPTE_TYPE_2", "COMPTE_NUMERO_2", "COMPTE_ETABLISSEMENT_2",
        "MISSION_REMUNERATION_MODE", "MISSION_HONORAIRES_MONTANT", "MISSION_HONORAIRES_DESCRIPTION",
        "SIGNATURE_LIEU", "SIGNATURE_DATE", "SIGNATURE_NOMBRE_EXEMPLAIRES",
    ]

    print(f"\nPlaceholders utilisés ({len(placeholders)}) :")
    for p in placeholders:
        print(f"  - {{{{{p}}}}}")

    return output_path

if __name__ == '__main__':
    create_rto_template()

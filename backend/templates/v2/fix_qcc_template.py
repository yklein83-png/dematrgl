# -*- coding: utf-8 -*-
"""
Script pour régénérer le template QCC avec corrections:
1. Header aligné avec le style DER
2. Pagination activée dans le footer
3. Situation Familiale: sections conditionnelles
4. Origine Economique: checkbox avant élément + montant
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os


# Mentions légales
MENTIONS_LEGALES = """Le Fare de l'Epargne - SARL au capital de 5 000 € - SIRET 893 599 979 00016 - RCS PAPEETE
Siège social : Immeuble Vatea, Rue des Remparts - Mission Papeete - 98714 PAPEETE
Conseiller en Investissements Financiers (CIF) - Membre de LA COMPAGNIE CIF, association agréée par l'AMF
Courtier en assurance (COA) immatriculé à l'ORIAS sous le n° 21003330 (www.orias.fr)"""


def set_cell_shading(cell, color):
    """Définit la couleur de fond d'une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_number(paragraph):
    """Ajoute un numéro de page au paragraphe"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    return run


def add_total_pages(paragraph):
    """Ajoute le nombre total de pages"""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "NUMPAGES"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    return run


def create_qcc_document():
    """Crée le document QCC avec les corrections"""
    doc = Document()

    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2.5)  # Plus de marge pour le footer
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ==========================================
    # EN-TÊTE (style DER)
    # ==========================================
    section = doc.sections[0]
    header = section.header

    # Vider l'en-tête existant
    for para in header.paragraphs:
        para.clear()

    # Logo / En-tête cabinet (comme DER)
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("LE FARE DE L'ÉPARGNE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Sous-titre
    sub_para = header.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Conseil en Gestion de Patrimoine")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(102, 102, 102)

    # Ligne de séparation (mentions légales courtes)
    mentions_para = header.add_paragraph()
    mentions_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mentions_para.add_run("CIF membre LA COMPAGNIE CIF (AMF) • ORIAS n° 21003330 • RCS PAPEETE")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.font.italic = True

    # ==========================================
    # PIED DE PAGE avec pagination
    # ==========================================
    footer = section.footer

    # Vider le footer existant
    for para in footer.paragraphs:
        para.clear()

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Version
    run = footer_para.add_run("QCC V.2025-03 | Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Numéro de page
    add_page_number(footer_para)

    run = footer_para.add_run(" / ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Total pages
    add_total_pages(footer_para)

    # ==========================================
    # TITRE DU DOCUMENT
    # ==========================================

    # Espace après header
    doc.add_paragraph()

    # Titre principal
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("QUESTIONNAIRE DE CONNAISSANCE CLIENT")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Sous-titre avec date
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Document réglementaire")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(102, 102, 102)

    # Tableau info client/conseiller (comme DER)
    doc.add_paragraph()

    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'

    # Ligne 1
    info_table.rows[0].cells[0].text = "Date :"
    info_table.rows[0].cells[1].text = "{{DATE_SIGNATURE}}"
    info_table.rows[0].cells[2].text = "Client :"
    info_table.rows[0].cells[3].text = "{{NOM_COMPLET_T1}}"

    # Ligne 2
    info_table.rows[1].cells[0].text = "Référence :"
    info_table.rows[1].cells[1].text = "{{NUMERO_CLIENT}}"
    info_table.rows[1].cells[2].text = "Conseiller :"
    info_table.rows[1].cells[3].text = "{{NOM_CONSEILLER}}"

    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i % 2 == 0:
                        run.font.bold = True
            if i % 2 == 0:
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # ==========================================
    # MISE EN GARDE
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("Mise en garde et informations préalables")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    mise_en_garde = """Dès lors qu'il fournit des conseils en investissements financiers et/ou en produits d'investissements assurantiels, le cabinet doit s'enquérir des exigences et besoins de ce dernier, de sa situation financière, de ses objectifs, de ses connaissances et de son expérience en matière financière, de sa tolérance au risque, de sa capacité à supporter les pertes, ainsi que de ses préférences en matière de durabilité, afin de délivrer un conseil adapté.

Nous attirons votre attention sur le fait que vous vous engagez à une communication exhaustive et sincère des informations vous concernant."""

    doc.add_paragraph(mise_en_garde)

    # ==========================================
    # SECTION 1: IDENTIFICATION
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("1 - CONNAISSANCE CLIENT")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Sous-section Identification
    p = doc.add_paragraph()
    run = p.add_run("Identification des titulaires")
    run.bold = True
    run.font.size = Pt(11)

    # Tableau titulaires
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "TITULAIRE 1"
    hdr_cells[1].text = "TITULAIRE 2"
    for cell in hdr_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D9E2F3')

    # Champs d'identification
    fields = [
        ("Civilité", "{{T1_CIVILITE}}", "{{T2_CIVILITE}}"),
        ("Nom", "{{T1_NOM}}", "{{T2_NOM}}"),
        ("Nom de naissance", "{{T1_NOM_JEUNE_FILLE}}", "{{T2_NOM_JEUNE_FILLE}}"),
        ("Prénom(s)", "{{T1_PRENOM}}", "{{T2_PRENOM}}"),
        ("Né(e) le / à", "{{T1_DATE_NAISSANCE}}\n{{T1_LIEU_NAISSANCE}}", "{{T2_DATE_NAISSANCE}}\n{{T2_LIEU_NAISSANCE}}"),
        ("Nationalité", "{{T1_NATIONALITE}}", "{{T2_NATIONALITE}}"),
        ("Adresse", "{{T1_ADRESSE}}\n{{T1_CODE_POSTAL}} {{T1_VILLE}}", "{{T2_ADRESSE}}\n{{T2_CODE_POSTAL}} {{T2_VILLE}}"),
        ("Email", "{{T1_EMAIL}}", "{{T2_EMAIL}}"),
        ("Téléphone", "{{T1_TELEPHONE}}", "{{T2_TELEPHONE}}"),
        ("Pièce d'identité", "{{T1_PIECE_IDENTITE}} N°{{T1_NUMERO_PIECE}}\nVal. {{T1_DATE_VALIDITE_PIECE}}", "{{T2_PIECE_IDENTITE}} N°{{T2_NUMERO_PIECE}}\nVal. {{T2_DATE_VALIDITE_PIECE}}"),
        ("Résidence fiscale / NIF", "{{T1_RESIDENCE_FISCALE}}\nNIF: {{T1_NIF}}", "{{T2_RESIDENCE_FISCALE}}\nNIF: {{T2_NIF}}"),
        ("Profession", "{{T1_PROFESSION}}\n{{T1_SECTEUR_ACTIVITE}}", "{{T2_PROFESSION}}\n{{T2_SECTEUR_ACTIVITE}}"),
        ("US Person (FATCA)", "{{T1_US_PERSON}}", "{{T2_US_PERSON}}"),
    ]

    for label, val1, val2 in fields:
        row = table.add_row()
        row.cells[0].text = f"{label}: {val1}"
        row.cells[1].text = f"{label}: {val2}"

    doc.add_paragraph()

    # ==========================================
    # SITUATION FAMILIALE (CONDITIONNELLE)
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("Situation Familiale")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("Situation matrimoniale: ").bold = True
    p.add_run("{{SITUATION_FAMILIALE}}")

    # Section Mariage (s'affiche si situation = Marié)
    p = doc.add_paragraph()
    p.add_run("{{#IF_MARIE}}")  # Conditionnel - sera traité par le générateur

    p = doc.add_paragraph()
    p.add_run("Date du mariage: ").bold = True
    p.add_run("{{DATE_MARIAGE}}   ")
    p.add_run("Régime matrimonial: ").bold = True
    p.add_run("{{REGIME_MATRIMONIAL}}")

    p = doc.add_paragraph()
    p.add_run("Contrat de mariage: ").bold = True
    p.add_run("{{CONTRAT_MARIAGE}}")

    p = doc.add_paragraph()
    p.add_run("{{/IF_MARIE}}")

    # Section PACS (s'affiche si situation = Pacsé)
    p = doc.add_paragraph()
    p.add_run("{{#IF_PACSE}}")

    p = doc.add_paragraph()
    p.add_run("Date du PACS: ").bold = True
    p.add_run("{{DATE_PACS}}   ")
    p.add_run("Régime du PACS: ").bold = True
    p.add_run("{{REGIME_PACS}}")

    p = doc.add_paragraph()
    p.add_run("{{/IF_PACSE}}")

    # Section Divorcé (s'affiche si situation = Divorcé)
    p = doc.add_paragraph()
    p.add_run("{{#IF_DIVORCE}}")

    p = doc.add_paragraph()
    p.add_run("Date du divorce: ").bold = True
    p.add_run("{{DATE_DIVORCE}}")

    p = doc.add_paragraph()
    p.add_run("{{/IF_DIVORCE}}")

    # Enfants
    p = doc.add_paragraph()
    run = p.add_run("Enfants")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run("Nombre d'enfant(s): ").bold = True
    p.add_run("{{NOMBRE_ENFANTS}}   ")
    p.add_run("Dont à charge: ").bold = True
    p.add_run("{{NOMBRE_ENFANTS_CHARGE}}")

    # Tableau enfants
    table_enf = doc.add_table(rows=1, cols=5)
    table_enf.style = 'Table Grid'

    headers = ["Prénom, Nom", "Date naissance", "Lien", "À charge?", "A des enfants?"]
    hdr = table_enf.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(hdr[i], 'E6E6E6')

    for i in range(4):
        row = table_enf.add_row()
        row.cells[0].text = f"{{{{ENFANT_{i+1}_NOM}}}}"
        row.cells[1].text = f"{{{{ENFANT_{i+1}_DATE_NAISSANCE}}}}"
        row.cells[2].text = f"{{{{ENFANT_{i+1}_LIEN}}}}"
        row.cells[3].text = f"{{{{ENFANT_{i+1}_A_CHARGE}}}}"
        row.cells[4].text = f"{{{{ENFANT_{i+1}_A_ENFANTS}}}}"

    doc.add_paragraph()

    # ==========================================
    # SECTION 2: SITUATION FINANCIÈRE
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("2 - SITUATION FINANCIÈRE ET PATRIMONIALE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    run = p.add_run("Flux financiers")
    run.bold = True
    run.font.size = Pt(11)

    table_fin = doc.add_table(rows=1, cols=2)
    table_fin.style = 'Table Grid'

    hdr = table_fin.rows[0].cells
    hdr[0].text = "TITULAIRE 1"
    hdr[1].text = "TITULAIRE 2"
    for cell in hdr:
        set_cell_shading(cell, 'D9E2F3')
        cell.paragraphs[0].runs[0].bold = True

    fin_rows = [
        ("Revenus annuels globaux", "{{REVENUS_ANNUELS_FOYER}}", "{{T2_REVENUS_ANNUELS_FOYER}}"),
        ("Patrimoine global (hors dettes)", "{{PATRIMOINE_GLOBAL}}", "{{T2_PATRIMOINE_GLOBAL}}"),
        ("Charges annuelles", "{{CHARGES_ANNUELLES_POURCENT}}% soit {{CHARGES_ANNUELLES_MONTANT}} €", "{{T2_CHARGES_ANNUELLES_POURCENT}}% soit {{T2_CHARGES_ANNUELLES_MONTANT}} €"),
        ("Capacité d'épargne mensuelle", "{{CAPACITE_EPARGNE_MENSUELLE}} €/mois", "{{T2_CAPACITE_EPARGNE_MENSUELLE}} €/mois"),
    ]

    for label, val1, val2 in fin_rows:
        row = table_fin.add_row()
        row.cells[0].text = f"{label}:\n{val1}"
        row.cells[1].text = f"{label}:\n{val2}"

    doc.add_paragraph()

    # ==========================================
    # ORIGINE DES FONDS (CORRIGÉE)
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("Origine économique des avoirs")
    run.bold = True
    run.font.size = Pt(11)

    table_orig = doc.add_table(rows=3, cols=2)
    table_orig.style = 'Table Grid'

    # Ligne 1: Nature
    row = table_orig.rows[0].cells
    row[0].text = "Nature des avoirs à investir"
    row[1].text = "{{ORIGINE_FONDS_NATURE}}"
    set_cell_shading(row[0], 'E6E6E6')

    # Ligne 2: Montant prévu
    row = table_orig.rows[1].cells
    row[0].text = "Montant prévu des avoirs à investir"
    row[1].text = "{{ORIGINE_FONDS_MONTANT_PREVU}} €"
    set_cell_shading(row[0], 'E6E6E6')

    # Ligne 3: Origine avec checkboxes CORRIGÉES
    row = table_orig.rows[2].cells
    row[0].text = "Origine économique\n(cocher et indiquer le montant)"
    set_cell_shading(row[0], 'E6E6E6')

    # Format corrigé: ☐/☑ Élément: Montant €
    origins_corrected = """{{ORIGINE_REVENUS_CHECK}} Revenus: {{ORIGINE_REVENUS_MONTANT}} €
{{ORIGINE_EPARGNE_CHECK}} Épargne constituée: {{ORIGINE_EPARGNE_MONTANT}} €
{{ORIGINE_HERITAGE_CHECK}} Héritage/Donation/Succession: {{ORIGINE_HERITAGE_MONTANT}} €
{{ORIGINE_CESSION_PRO_CHECK}} Cession d'actifs professionnels: {{ORIGINE_CESSION_PRO_MONTANT}} €
{{ORIGINE_CESSION_IMMO_CHECK}} Cession immobilière: {{ORIGINE_CESSION_IMMO_MONTANT}} €
{{ORIGINE_CESSION_MOBILIERE_CHECK}} Cession mobilière: {{ORIGINE_CESSION_MOBILIERE_MONTANT}} €
{{ORIGINE_GAINS_JEU_CHECK}} Gains de jeu: {{ORIGINE_GAINS_JEU_MONTANT}} €
{{ORIGINE_ASSURANCE_VIE_CHECK}} Assurance-vie: {{ORIGINE_ASSURANCE_VIE_MONTANT}} €
{{ORIGINE_AUTRES_CHECK}} Autres: {{ORIGINE_AUTRES_MONTANT}} €"""

    row[1].text = origins_corrected

    doc.add_paragraph()

    # Provenance bancaire
    p = doc.add_paragraph()
    p.add_run("Provenance bancaire: ").bold = True
    p.add_run("{{ORIGINE_FONDS_PROVENANCE_ETABLISSEMENT}}")

    doc.add_paragraph()

    # ==========================================
    # SIGNATURES
    # ==========================================

    p = doc.add_paragraph()
    run = p.add_run("Signatures")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.add_run("Je soussigné(e), certifie sur l'honneur l'exactitude des informations portées sur le présent document.")

    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.style = 'Table Grid'

    sig_table.rows[0].cells[0].text = "Titulaire 1\n\n\n\nSignature: {{SIGNATURE_CLIENT}}\nFait à: {{LIEU_SIGNATURE}}\nLe: {{DATE_SIGNATURE}}"
    sig_table.rows[0].cells[1].text = "Titulaire 2\n\n\n\nSignature: {{SIGNATURE_CLIENT_T2}}\nFait à: {{LIEU_SIGNATURE}}\nLe: {{DATE_SIGNATURE}}"

    sig_table.rows[1].cells[0].text = "Conseiller\n\n\n\nSignature: {{SIGNATURE_CONSEILLER}}\nFait à: {{LIEU_SIGNATURE}}\nLe: {{DATE_SIGNATURE}}"
    sig_table.rows[1].cells[1].text = ""

    for cell in sig_table.rows[0].cells:
        set_cell_shading(cell, 'F5F5F5')

    # ==========================================
    # MENTIONS LÉGALES (bas de document)
    # ==========================================

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run(MENTIONS_LEGALES)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    return doc


def main():
    """Génère le template QCC corrigé"""
    script_dir = os.path.dirname(os.path.abspath(__file__))

    # Backup de l'ancien template
    old_path = os.path.join(script_dir, "QCC_V2_TEMPLATE.docx")
    backup_path = os.path.join(script_dir, "QCC_V2_TEMPLATE_OLD.docx")

    if os.path.exists(old_path):
        import shutil
        shutil.copy(old_path, backup_path)
        print(f"Backup créé: {backup_path}")

    # Générer le nouveau template
    doc = create_qcc_document()
    output_path = os.path.join(script_dir, "QCC_V2_TEMPLATE.docx")
    doc.save(output_path)

    print(f"Template QCC généré: {output_path}")
    print("\nCorrections appliquées:")
    print("1. ✅ Header aligné avec style DER (logo + mentions)")
    print("2. ✅ Pagination activée (Page X / Y)")
    print("3. ✅ Situation Familiale avec sections conditionnelles")
    print("4. ✅ Origine Economique: checkbox + élément + montant")


if __name__ == "__main__":
    main()

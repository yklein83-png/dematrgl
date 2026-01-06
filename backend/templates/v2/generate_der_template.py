"""
Script pour générer le template DER (Document d'Entrée en Relation) v2
Avec mise en page professionnelle et placeholders {{FIELD}}
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# Mentions légales obligatoires
MENTIONS_LEGALES = """
Le Fare de l'Épargne - SARL au capital de 5 000 € - SIRET 893 599 979 00016 - RCS PAPEETE
Siège social : Immeuble Vatea, Rue des Remparts - Mission Papeete - 98714 PAPEETE
Conseiller en Investissements Financiers (CIF) - Membre de LA COMPAGNIE CIF, association agréée par l'AMF
Courtier en assurance (COA) immatriculé à l'ORIAS sous le n° 21003330 (www.orias.fr)
Courtier en opérations de banque et services de paiement (COBSP)
Démarcheur Bancaire et Financier mandaté par SELENCIA, INTENCIAL, VIE PLUS, UAF LIFE PATRIMOINE
Responsabilité Civile Professionnelle et Garantie Financière conformes aux articles L541-3 et L512-6 et 7 du Code des Assurances
"""


def set_cell_shading(cell, color):
    """Définir la couleur de fond d'une cellule"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_style(doc, text, level=1):
    """Ajouter un titre avec style"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True

    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 51, 102)  # Bleu foncé
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 51, 102)

    return p


def add_normal_paragraph(doc, text, indent=False):
    """Ajouter un paragraphe normal"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    if indent:
        p.paragraph_format.left_indent = Cm(0.5)

    return p


def add_bullet_list(doc, items):
    """Ajouter une liste à puces"""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        for run in p.runs:
            run.font.size = Pt(10)


def create_info_table(doc, rows, col_widths=None):
    """Créer un tableau d'informations"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(rows):
        # Cellule label
        cell_label = table.rows[i].cells[0]
        cell_label.text = label
        for para in cell_label.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
                run.font.bold = True
        set_cell_shading(cell_label, 'E8E8E8')

        # Cellule valeur
        cell_value = table.rows[i].cells[1]
        cell_value.text = value
        for para in cell_value.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

    # Largeurs des colonnes
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def create_partners_table(doc):
    """Créer le tableau des partenaires"""
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    # En-tête
    header_cells = table.rows[0].cells
    header_cells[0].text = "Activité"
    header_cells[1].text = "Partenaires"
    for cell in header_cells:
        set_cell_shading(cell, '003366')
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Données
    data = [
        ("CIF", "Sofidy, Périal, Iroko, Theoreim, Sogenial, Remake Live, Novaxia, Alderan, Advenis, Corum, Principal, PPG"),
        ("IAS", "Vie Plus, UAF Life, UNEP, Intencial, Selencia"),
        ("IOBSP", "Socredo, Banque de Polynésie, Banque de Tahiti")
    ]

    for i, (activity, partners) in enumerate(data):
        row = table.rows[i + 1]
        row.cells[0].text = activity
        row.cells[1].text = partners
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        set_cell_shading(row.cells[0], 'E8E8E8')

    return table


def generate_der_template():
    """Générer le template DER v2"""
    doc = Document()

    # Configuration des marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ==========================================
    # EN-TÊTE
    # ==========================================

    # Logo / En-tête cabinet
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("LE FARE DE L'ÉPARGNE")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # Sous-titre
    sub_header = doc.add_paragraph()
    sub_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_header.add_run("Conseil en Gestion de Patrimoine")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)

    doc.add_paragraph()

    # ==========================================
    # TITRE DU DOCUMENT
    # ==========================================

    add_heading_style(doc, "DOCUMENT D'ENTRÉE EN RELATION", 1)

    # Version
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version_para.add_run("Version 2025.03")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ==========================================
    # INFORMATIONS CLIENT
    # ==========================================

    info_table = doc.add_table(rows=2, cols=4)
    info_table.style = 'Table Grid'

    # Ligne 1
    info_table.rows[0].cells[0].text = "Date :"
    info_table.rows[0].cells[1].text = "{{DATE_SIGNATURE}}"
    info_table.rows[0].cells[2].text = "Client :"
    info_table.rows[0].cells[3].text = "{{NOM_COMPLET_T1}}"

    # Ligne 2
    info_table.rows[1].cells[0].text = "Référence :"
    info_table.rows[1].cells[1].text = "{{NUMERO_CLIENT}}"
    info_table.rows[1].cells[2].text = "Conseiller :"
    info_table.rows[1].cells[3].text = "{{NOM_CONSEILLER}}"

    for row in info_table.rows:
        for i, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                    if i % 2 == 0:  # Labels
                        run.font.bold = True
            if i % 2 == 0:
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # ==========================================
    # INTRODUCTION
    # ==========================================

    intro = doc.add_paragraph()
    intro.add_run("Madame, Monsieur,").bold = True

    add_normal_paragraph(doc,
        "Dans le cadre de notre activité de conseil en gestion de patrimoine, notre cabinet est soumis "
        "à diverses réglementations correspondant aux différents statuts que nous exerçons. Le présent "
        "document est établi en vue de vous présenter le cadre réglementaire dans lequel nous allons travailler.")

    add_normal_paragraph(doc,
        "Nous vous adressons les informations réglementaires et les informations complémentaires préconisées "
        "par LA COMPAGNIE CIF, association agréée par l'AMF, par la CNCEF, association agréée par l'ACPR, "
        "ainsi que par LA COMPAGNIE IOBSP, association agréée par l'ACPR, dont nous sommes adhérents.")

    doc.add_paragraph()

    # ==========================================
    # PRÉSENTATION DU CABINET
    # ==========================================

    add_heading_style(doc, "Présentation des activités du cabinet", 2)

    add_normal_paragraph(doc,
        "Notre cabinet Le Fare de l'Épargne est enregistré auprès de l'ORIAS sous le numéro 21003330 (www.orias.fr) :")

    link_para = doc.add_paragraph()
    link_para.add_run("https://www.orias.fr/home/showIntermediaire/D00020346").font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph()

    # ==========================================
    # ACTIVITÉ CIF
    # ==========================================

    add_heading_style(doc, "Conseiller en Investissements Financiers (CIF)", 2)

    add_normal_paragraph(doc,
        "Dans le cadre de son activité de CIF, le cabinet est susceptible de fournir des conseils en "
        "investissement de manière non indépendante :")

    add_normal_paragraph(doc,
        "Le conseil en investissement non indépendant repose sur une analyse diversifiée de manière large "
        "des instruments ou services financiers disponibles sur le marché, et sans pour autant se limiter "
        "aux instruments ou services des entités partenaires. La rémunération pourra se faire sous forme "
        "de rétrocessions de commissions et/ou d'autres avantages monétaires ou non monétaires perçus ou "
        "reçus de la part de nos partenaires, cumulés éventuellement à nos honoraires. Cette rémunération "
        "est justifiée par une amélioration de la qualité de notre prestation et dans l'objectif d'agir au "
        "mieux de vos intérêts.", indent=True)

    add_normal_paragraph(doc,
        "La nature du conseil ainsi que les contours de la rémunération vous seront communiquées de manière "
        "plus précise au sein de la lettre de mission.")

    # Autorité de contrôle CIF
    authority_table = doc.add_table(rows=1, cols=1)
    authority_table.style = 'Table Grid'
    cell = authority_table.rows[0].cells[0]
    set_cell_shading(cell, 'F5F5F5')
    p = cell.paragraphs[0]
    p.add_run("L'autorité de contrôle pour l'activité CIF est :\n").bold = True
    p.add_run("Autorité des Marchés Financiers\n17 Place de la Bourse\n75 082 PARIS Cedex 02")
    for run in p.runs:
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ==========================================
    # ACTIVITÉ IAS
    # ==========================================

    add_heading_style(doc, "Intermédiaire en Assurance (IAS)", 2)

    add_normal_paragraph(doc,
        "En qualité de courtier en assurance de catégorie C :")

    add_normal_paragraph(doc,
        "Le cabinet n'est pas soumis à une obligation contractuelle de travailler exclusivement avec une ou "
        "plusieurs entreprises d'assurance et peut communiquer sur simple demande du client, le nom de ces "
        "dernières et offre au client un conseil fondé sur une analyse objective et impartiale du marché, "
        "et analyse, pour ce faire un nombre suffisant de contrats d'assurance offerts sur le marché de façon "
        "à pouvoir recommander, en fonction de critères professionnels, le contrat qui serait adapté aux "
        "besoins du souscripteur éventuel.", indent=True)

    add_normal_paragraph(doc,
        "Dans le cadre de cette activité, le cabinet n'est pas autorisé à encaisser des fonds destinés à un "
        "assuré ou à une compagnie d'assurances, sauf dans le cadre de l'IARD.")

    add_normal_paragraph(doc,
        "Dans le cadre de l'exercice de notre activité d'intermédiation en assurance, notre prestation est "
        "fournie sur un conseil approprié permettant de conseiller le produit le plus adapté aux besoins du "
        "client, et le formaliser au sein d'un rapport d'adéquation.")

    doc.add_paragraph()

    # ==========================================
    # ACTIVITÉ IOBSP
    # ==========================================

    add_heading_style(doc, "Intermédiaire en Opérations de Banque et Services de Paiement (IOBSP)", 2)

    add_normal_paragraph(doc,
        "En qualité de Courtier en opérations de banque et services de paiement exerçant l'intermédiation "
        "en vertu d'un mandat de notre client pour les catégories de crédit suivantes : crédits immobiliers.")

    add_normal_paragraph(doc,
        "Dans le cadre de son activité d'intermédiation, et conformément à l'article L. 519-6 du Code monétaire "
        "et financier, le cabinet ne perçoit aucune somme représentative de provision, de commissions, de frais "
        "de recherche, de démarches, de constitution de dossier ou d'entremise quelconque, avant le versement "
        "effectif des fonds prêtés.", indent=True)

    # Conseil crédits immobiliers
    add_heading_style(doc, "Informations spécifiques portant sur le service de conseil de crédits immobiliers", 3)

    add_normal_paragraph(doc,
        "Notre cabinet propose un service de conseil portant sur des contrats de crédits immobiliers. "
        "Dans le cadre de ce service, notre recommandation porte sur uniquement une gamme référencée et "
        "restreinte de produits (conseil non indépendant).")

    add_normal_paragraph(doc,
        "Cela signifie que, pour vous conseiller, nous nous fondons sur un nombre restreint de contrats de "
        "crédit disponibles sur le marché pour les intermédiaires agissant en vertu d'un mandat délivré par "
        "un client (Article L519-1-1 du code monétaire et financier). Dans le cadre du service de conseil, "
        "notre recommandation porte sur une large gamme de contrats de crédit disponibles sur le marché.", indent=True)

    # Autorité de contrôle IAS/IOBSP
    authority_table2 = doc.add_table(rows=1, cols=1)
    authority_table2.style = 'Table Grid'
    cell = authority_table2.rows[0].cells[0]
    set_cell_shading(cell, 'F5F5F5')
    p = cell.paragraphs[0]
    p.add_run("L'autorité de contrôle des activités IAS et IOBSP est :\n").bold = True
    p.add_run("ACPR\n4 place de Budapest\nDirection du Contrôle des Pratiques Commerciales\n75436 PARIS Cedex 9")
    for run in p.runs:
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ==========================================
    # LISTE DES PARTENAIRES
    # ==========================================

    add_heading_style(doc, "Liste des partenaires", 2)

    create_partners_table(doc)

    doc.add_paragraph()

    # ==========================================
    # ORGANISATION ET BONNE CONDUITE
    # ==========================================

    add_heading_style(doc, "Organisation et bonne conduite", 2)

    conseiller_para = doc.add_paragraph()
    conseiller_para.add_run("Nous vous informons que ").font.size = Pt(10)
    conseiller_para.add_run("{{TITRE_CONSEILLER}} {{NOM_CONSEILLER}}").bold = True
    conseiller_para.add_run(" sera votre conseiller, sous la responsabilité du cabinet.").font.size = Pt(10)

    add_normal_paragraph(doc,
        "Enfin, nous vous confirmons que conformément aux Codes de Bonne Conduite de LA COMPAGNIE CIF, "
        "La CNCEF et LA COMPAGNIE IOBSP, le cabinet a pris l'engagement pour maintenir son adhésion, "
        "de suivre des formations réglementaires obligatoires, de justifier annuellement d'une assurance "
        "responsabilité civile professionnelle, de produire le casier judiciaire des dirigeants et de "
        "déclarer immédiatement, sous peine de déchéance, tout événement susceptible de le modifier, "
        "enfin de respecter toutes les dispositions incluses dans les Codes de Bonne Conduite de "
        "LA COMPAGNIE CIF et La CNCEF et LA COMPAGNIE IOBSP.")

    doc.add_paragraph()

    # ==========================================
    # TRAITEMENT DES RÉCLAMATIONS
    # ==========================================

    add_heading_style(doc, "Traitement des réclamations", 2)

    add_normal_paragraph(doc,
        "La réclamation peut être exprimée lors d'un rendez-vous et doit être adressée par écrit à la "
        "Direction du cabinet par courrier BP4646 - 98713-Papeete-Tahiti ou par e-mail à contact@fare-epargne.com.")

    add_normal_paragraph(doc, "Le cabinet s'engage à :")

    add_bullet_list(doc, [
        "Accuser réception de la réclamation dans un délai de dix jours ouvrables ;",
        "Puis à y répondre dans un délai de deux mois maximums à compter de la date d'envoi de la réclamation, "
        "sauf survenance de circonstances particulières dûment justifiées."
    ])

    add_normal_paragraph(doc,
        "Si la réponse apportée à sa réclamation ne vous apparaît pas satisfaisante, vous pourrez saisir "
        "le médiateur de la consommation compétent suivant :")

    # Médiateur IAS/COBSP
    mediator_table = doc.add_table(rows=2, cols=1)
    mediator_table.style = 'Table Grid'

    cell1 = mediator_table.rows[0].cells[0]
    set_cell_shading(cell1, 'F5F5F5')
    p1 = cell1.paragraphs[0]
    p1.add_run("Pour les activités IAS et COBSP :\n").bold = True
    p1.add_run("CNPM Médiation Consommation\n27 avenue de la Libération\n42400 Saint Chamond\n"
               "Tél. : 09 88 30 27 72\ncontact-admin@cnpm-mediation-consommation.eu\n"
               "https://www.cnpm-mediation-consommation.eu/")
    for run in p1.runs:
        run.font.size = Pt(9)

    cell2 = mediator_table.rows[1].cells[0]
    set_cell_shading(cell2, 'F5F5F5')
    p2 = cell2.paragraphs[0]
    p2.add_run("Pour l'activité de CIF :\n").bold = True
    p2.add_run("Le Médiateur de l'AMF\nAutorité des Marchés Financiers\n17, Place de la Bourse\n"
               "75 082 Paris Cedex 02\nhttp://www.amf-france.org/Le-mediateur-de-l-AMF/Presentation.html")
    for run in p2.runs:
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ==========================================
    # POLITIQUE DURABILITÉ
    # ==========================================

    add_heading_style(doc, "Politique en matière de durabilité", 2)

    add_normal_paragraph(doc,
        "Notre cabinet a mis en place un processus de sélection des instruments financiers, tenant compte "
        "des facteurs de durabilité au sens de l'article 2, point 24, du règlement (UE) 2019/2088 du "
        "Parlement européen et du Conseil du 27 novembre 2019. Pour ce faire, nous procédons à l'analyse "
        "des produits qu'il entend référencer sur la base de plusieurs critères :")

    add_bullet_list(doc, [
        "Le fait qu'ils constituent des investissements durables au sens du Règlement SFDR ;",
        "Le fait qu'ils constituent des investissements durables à vocation environnementale au sens du Règlement Taxonomie ;",
        "Le fait qu'ils prennent ou non en compte les incidences négatives en matière de durabilité (PIA)."
    ])

    add_normal_paragraph(doc, "Ces différents éléments vous seront expliqués dans le cadre du déroulement de notre mission.")

    doc.add_paragraph()

    # ==========================================
    # DONNÉES PERSONNELLES
    # ==========================================

    add_heading_style(doc, "Traitement des données personnelles et confidentialité", 2)

    add_normal_paragraph(doc,
        "Notre travail nécessitant l'accès, la collecte, le traitement, le stockage, l'exploitation et le "
        "cas échéant, la transmission à des tiers des données à caractère personnel que vous voudrez bien "
        "nous transmettre, nous nous engageons à respecter les dispositions de la loi n° 78-17 du 6 janvier "
        "1978 et du Règlement 2016/679 du Parlement européen et du Conseil du 27 avril 2016 relatif à la "
        "protection des personnes physiques à l'égard du traitement des données à caractère personnel et "
        "à la libre circulation de ces données.")

    add_normal_paragraph(doc,
        "Les données personnelles que vous nous transmettez dans le cadre de nos activités et des services "
        "proposés sont collectées et traitées par Mr Pierre POHER en qualité de responsable de traitement. "
        "L'identité et les coordonnées du délégué à la protection des données personnelles au sein du cabinet "
        "sont : Mr Pierre POHER - contact@fare-epargne.com.")

    add_normal_paragraph(doc,
        "Notre cabinet s'engage à ne collecter et traiter les données recueillies qu'au regard des finalités "
        "de traitement convenues (consentement, nécessité contractuelle, respect d'une obligation légale) entre "
        "notre cabinet et son client, à en préserver la sécurité et l'intégrité, à ne communiquer ses informations "
        "qu'aux tiers auxquels il est nécessaire de les transmettre pour la bonne exécution des missions confiées.")

    add_normal_paragraph(doc,
        "Nous vous informons que vous bénéficiez d'un droit d'accès, de rectification, d'un droit de restriction "
        "sur l'utilisation, d'un droit d'opposition à l'utilisation, d'un droit à l'effacement ou encore d'un "
        "droit à la portabilité de vos données personnelles.")

    add_normal_paragraph(doc, "Vous pouvez ainsi bénéficier de ces droits à tout moment en nous informant directement :")

    add_bullet_list(doc, [
        "par courrier électronique : contact@fare-epargne.com",
        "par courrier postal : BP4646 - 98713-Papeete-Tahiti"
    ])

    add_normal_paragraph(doc,
        "Vous disposez également du droit d'introduire une réclamation auprès de la CNIL (Commission Nationale "
        "de l'Information et des Libertés) via leur site internet ou par voie postale : CNIL - 3 place de "
        "Fontenoy - TSA 80715 75334 PARIS CEDEX 07.")

    doc.add_paragraph()

    # ==========================================
    # MOYENS DE COMMUNICATION
    # ==========================================

    add_heading_style(doc, "Moyens de communication", 2)

    add_normal_paragraph(doc, "Nous vous informons que nous communiquerons avec vous par les moyens suivants :")

    add_bullet_list(doc, [
        "Par courrier",
        "Par tout autre support durable (e-mail, site internet dédié, etc...) dès lors que les clients auront "
        "opté formellement pour la fourniture des informations sur cet autre support et que ce même support "
        "se révèlera adapté au contexte dans lequel seront conduites les affaires."
    ])

    add_normal_paragraph(doc,
        "Les coordonnées de contact suivantes que vous nous avez communiquées pourront à tout moment être "
        "utilisées en vue de sécuriser et d'authentifier nos échanges.")

    doc.add_paragraph()

    # Case à cocher support durable
    checkbox_table = doc.add_table(rows=1, cols=1)
    checkbox_table.style = 'Table Grid'
    cell = checkbox_table.rows[0].cells[0]
    set_cell_shading(cell, 'FFFDE7')  # Jaune clair
    p = cell.paragraphs[0]
    p.add_run("[ ] ").font.size = Pt(12)
    p.add_run("En cochant cette case, vous acceptez que les informations soient communiquées par le biais "
              "d'un support durable autre que le papier, à savoir par courriel. À tout moment, vous pouvez "
              "demander à disposer du document papier ou à revenir à une remise de document sous format papier.")
    for run in p.runs:
        run.font.size = Pt(10)

    doc.add_paragraph()

    # Courriel du client
    email_table = doc.add_table(rows=1, cols=2)
    email_table.style = 'Table Grid'
    email_table.rows[0].cells[0].text = "Courriel :"
    email_table.rows[0].cells[1].text = "{{T1_EMAIL}}"
    for cell in email_table.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
    set_cell_shading(email_table.rows[0].cells[0], 'E8E8E8')

    doc.add_paragraph()

    # ==========================================
    # MISE À JOUR
    # ==========================================

    add_heading_style(doc, "Mise à jour des informations", 2)

    add_normal_paragraph(doc,
        "Nous vous ferons parvenir toute mise à jour des différentes informations contenues dans le présent "
        "document en vous les communiquant par courrier, e-mail ou espace dédié. Vous pouvez également obtenir "
        "à tout moment ces informations sur simple demande auprès de notre cabinet.")

    doc.add_paragraph()
    doc.add_paragraph()

    # ==========================================
    # CONCLUSION ET SIGNATURES
    # ==========================================

    conclusion = doc.add_paragraph()
    conclusion.add_run("En vous remerciant de la confiance que vous voudrez bien nous accorder,").font.size = Pt(10)

    closing = doc.add_paragraph()
    closing.add_run("Nous vous prions de croire, ").font.size = Pt(10)
    closing.add_run("{{TITRE_CLIENT}} {{NOM_COMPLET_T1}}").bold = True
    closing.add_run(", à l'assurance de nos sentiments distingués.").font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    # Tableau des signatures
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'

    # En-têtes
    sig_table.rows[0].cells[0].text = "Signature du Client"
    sig_table.rows[0].cells[1].text = "Signature du Conseiller"
    for cell in sig_table.rows[0].cells:
        set_cell_shading(cell, 'E8E8E8')
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    # Noms
    sig_table.rows[1].cells[0].text = "{{NOM_COMPLET_T1}}"
    sig_table.rows[1].cells[1].text = "{{NOM_CONSEILLER}}"
    for cell in sig_table.rows[1].cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.size = Pt(10)

    # Espace signature
    sig_table.rows[2].cells[0].text = "\n\n\n"
    sig_table.rows[2].cells[1].text = "\n\n\n"

    # Hauteur des lignes
    for row in sig_table.rows:
        for cell in row.cells:
            cell.width = Cm(8)

    doc.add_paragraph()

    # Date et lieu
    date_lieu = doc.add_paragraph()
    date_lieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_lieu.add_run("Fait à {{LIEU_SIGNATURE}}, le {{DATE_SIGNATURE}}").font.size = Pt(10)
    date_lieu.add_run("\nEn {{NOMBRE_EXEMPLAIRES}} exemplaires originaux").font.size = Pt(9)

    # ==========================================
    # MENTIONS LÉGALES (Pied de page)
    # ==========================================

    doc.add_paragraph()
    doc.add_paragraph()

    # Ligne de séparation
    sep = doc.add_paragraph()
    sep.add_run("─" * 80).font.size = Pt(6)

    # Mentions légales
    mentions = doc.add_paragraph()
    mentions.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mentions.add_run(MENTIONS_LEGALES.strip())
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Sauvegarder le template
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "DER_V2_TEMPLATE.docx")
    doc.save(output_path)
    print(f"Template DER v2 généré : {output_path}")

    return output_path


if __name__ == "__main__":
    generate_der_template()

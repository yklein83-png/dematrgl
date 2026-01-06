"""
Service de génération de documents DOCX
Utilise python-docx et les templates TXT officiels
"""

import os
import re
from datetime import datetime
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.models.client import Client
from app.models.user import User
from app.config import settings


class DocxGenerator:
    """
    Générateur de documents DOCX
    Utilise les templates TXT de la Liasse Réglementaire CGP
    et les templates DOCX v2 avec placeholders {{FIELD}}
    """

    def __init__(self):
        """Initialise le générateur avec les chemins"""
        self.export_path = settings.EXPORT_PATH

        # Chemin vers les templates TXT - utilise le dossier templates du projet
        self.txt_templates_path = os.path.join(settings.DOCX_TEMPLATE_PATH, "txt")

        # Chemin vers les templates DOCX v2
        self.v2_templates_path = os.path.join(settings.DOCX_TEMPLATE_PATH, "v2")

        # Créer le dossier export si nécessaire
        os.makedirs(self.export_path, exist_ok=True)

        # Charger les mentions légales une seule fois
        self._mentions_legales = self._load_mentions_legales()

    # ==========================================
    # MÉTHODES UTILITAIRES
    # ==========================================

    def _load_txt_template(self, filename: str) -> str:
        """Charger un template TXT"""
        filepath = os.path.join(self.txt_templates_path, filename)
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            return ""

    def _load_mentions_legales(self) -> str:
        """Charger les mentions légales"""
        return self._load_txt_template("MentionsLegales.txt")

    def _format_date(self, date_obj) -> str:
        """Formater une date au format français"""
        if not date_obj:
            return ""
        if isinstance(date_obj, str):
            return date_obj
        return date_obj.strftime("%d/%m/%Y")

    def _format_date_long(self, date_obj=None) -> str:
        """Formater une date au format long français (ex: 27 novembre 2025)"""
        if date_obj is None:
            date_obj = datetime.now()
        mois = [
            "janvier", "février", "mars", "avril", "mai", "juin",
            "juillet", "août", "septembre", "octobre", "novembre", "décembre"
        ]
        return f"{date_obj.day} {mois[date_obj.month - 1]} {date_obj.year}"

    def _format_montant(self, montant) -> str:
        """Formater un montant en euros"""
        if not montant:
            return "0,00 €"
        try:
            value = float(montant)
            return f"{value:,.2f} €".replace(",", " ").replace(".", ",")
        except:
            return str(montant)

    def _safe(self, val, default="") -> str:
        """Retourner une valeur sécurisée"""
        return str(val) if val is not None else default

    def _checkbox(self, val) -> str:
        """Retourner une checkbox cochée ou non"""
        return "☑" if val else "☐"

    def _save_document(self, doc: Document, filename: str) -> str:
        """Sauvegarder le document"""
        filepath = os.path.join(self.export_path, filename)
        doc.save(filepath)
        return filepath

    def _generate_filename(self, doc_type: str, client) -> str:
        """
        Générer un nom de fichier standardisé
        Format: nomdudoc_nomclient_P_YYYYMMDD_HHMM.docx
        Exemple: DER_DUPONT_J_20251127_1430.docx
        """
        # Nettoyer le nom (enlever accents et caractères spéciaux)
        nom = self._safe(client.t1_nom, "CLIENT").upper()
        nom = nom.replace(" ", "_").replace("-", "_")

        # Première lettre du prénom
        prenom = self._safe(client.t1_prenom, "X")
        initiale_prenom = prenom[0].upper() if prenom else "X"

        # Date et heure
        now = datetime.now()
        date_str = now.strftime('%Y%m%d')
        heure_str = now.strftime('%H%M')

        return f"{doc_type}_{nom}_{initiale_prenom}_{date_str}_{heure_str}.docx"

    def _create_document_from_text(self, text: str, title: str = None) -> Document:
        """
        Créer un document Word à partir d'un texte brut
        Gère les paragraphes et le formatage de base
        """
        doc = Document()

        # Style par défaut
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Traiter le texte ligne par ligne
        lines = text.split('\n')

        for line in lines:
            # Ignorer les lignes vides multiples
            stripped = line.strip()

            if not stripped:
                doc.add_paragraph()
                continue

            # Détecter les titres (lignes en majuscules ou commençant par des numéros)
            if stripped.isupper() and len(stripped) > 3 and len(stripped) < 100:
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(14)
            elif stripped.startswith(('I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.', 'IX.', 'X.', 'XI.', 'XII.', 'XIII.', 'XIV.', 'XV.', 'XVI.', 'XVII.')):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
            elif stripped.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
                run.font.size = Pt(12)
            elif stripped.startswith('•') or stripped.startswith('-'):
                # Puces
                p = doc.add_paragraph(stripped, style='List Bullet')
            elif stripped.startswith('\t'):
                # Lignes avec tabulation
                p = doc.add_paragraph()
                p.add_run(stripped)
            else:
                # Paragraphe normal
                doc.add_paragraph(stripped)

        return doc

    def _add_mentions_legales(self, doc: Document) -> None:
        """Ajouter les mentions légales en pied de page"""
        if not self._mentions_legales:
            return

        # Ajouter une ligne de séparation
        doc.add_paragraph()
        doc.add_paragraph("_" * 80)

        # Ajouter les mentions légales en petit
        p = doc.add_paragraph()
        run = p.add_run(self._mentions_legales.strip())
        run.font.size = Pt(8)
        run.font.name = 'Calibri'

    def _replace_placeholders(self, text: str, replacements: Dict[str, str]) -> str:
        """Remplacer les placeholders dans le texte"""
        result = text
        for key, value in replacements.items():
            result = result.replace(key, self._safe(value))
        return result

    # ==========================================
    # GÉNÉRATION DER (Document d'Entrée en Relation)
    # ==========================================

    async def generate_der(self, client: Client, conseiller: User) -> str:
        """
        Générer le Document d'Entrée en Relation
        Utilise le template DER.txt officiel
        """
        # Charger le template
        template = self._load_txt_template("DER.txt")

        if not template:
            # Fallback si template non trouvé
            template = """
DOCUMENT D'ENTRÉE EN RELATION
V.2025-03

Madame, Monsieur,

Dans le cadre de notre activité de conseil en gestion de patrimoine, notre cabinet est soumis à diverses réglementations correspondant aux différents statuts que nous exerçons.

[Template non trouvé - veuillez vérifier le chemin vers DER.txt]
"""

        # Préparer les remplacements
        replacements = {
            "$TITRE_CONSEILLER$": "M." if conseiller.prenom else "",
            "$NOM_CONSEILLER$": self._safe(conseiller.nom, "").upper(),
            "$PRENOM_CONSEILLER$": self._safe(conseiller.prenom),
            "$TITRE_CONTACT$": self._safe(client.t1_civilite),
            "$NOM_CONTACT$": self._safe(client.t1_nom, "").upper(),
            "$DATE_JOUR$": self._format_date(datetime.now()),
            "$DATE_JOUR_JJJ_MM_AAAA$": self._format_date_long(),
        }

        # Remplacer les placeholders
        content = self._replace_placeholders(template, replacements)

        # Créer le document Word
        doc = self._create_document_from_text(content)

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("DER", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION QCC (Questionnaire Connaissance Client)
    # ==========================================

    async def generate_kyc(self, client: Client, conseiller: User) -> str:
        """
        Générer le Questionnaire Connaissance Client et Profil de Risques
        Utilise le template QCC&Risk.txt officiel
        """
        # Charger le template
        template = self._load_txt_template("QCC&Risk.txt")

        if not template:
            template = """
DOCUMENT DE CONNAISSANCE CLIENT ET PROFIL DE RISQUES
V.2025-03

[Template non trouvé - veuillez vérifier le chemin vers QCC&Risk.txt]
"""

        # Préparer les remplacements
        replacements = {
            "$TITRE_CONTACT$": self._safe(client.t1_civilite),
            "$NOM_CONTACT$": self._safe(client.t1_nom, "").upper(),
            "$PRENOM_CONTACT$": self._safe(client.t1_prenom),
            "$DATE_JOUR$": self._format_date(datetime.now()),
            "$DATE_JOUR_JJJ_MM_AAAA$": self._format_date_long(),
        }

        # Remplacer les placeholders
        content = self._replace_placeholders(template, replacements)

        # Créer le document Word
        doc = self._create_document_from_text(content)

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("QCC", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION LETTRE DE MISSION CIF
    # ==========================================

    async def generate_lettre_mission_cif(self, client: Client, conseiller: User) -> str:
        """
        Générer la Lettre de Mission CIF
        Utilise le template LettreMission.txt officiel
        """
        # Charger le template
        template = self._load_txt_template("LettreMission.txt")

        if not template:
            template = """
LETTRE DE MISSION
V.2025-03

Madame, Monsieur,

[Template non trouvé - veuillez vérifier le chemin vers LettreMission.txt]
"""

        # Préparer les remplacements
        replacements = {
            "$DATE_JOUR_JJJ_MM_AAAA$": self._format_date_long(),
            "$DATE_JOUR$": self._format_date(datetime.now()),
            "$TITRE_CLIENT$": self._safe(client.t1_civilite),
            "$NOM_CLIENT$": self._safe(client.t1_nom, "").upper(),
            "$PRENOM_CLIENT$": self._safe(client.t1_prenom),
        }

        # Remplacer les placeholders
        content = self._replace_placeholders(template, replacements)

        # Créer le document Word
        doc = self._create_document_from_text(content)

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("LETTRE_MISSION_CIF", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION DÉCLARATION D'ADÉQUATION (DA CIF)
    # ==========================================

    async def generate_declaration_adequation(self, client: Client, conseiller: User) -> str:
        """
        Générer la Déclaration d'Adéquation / Rapport d'Adéquation
        Utilise le template DA CIF.txt officiel
        """
        # Charger le template
        template = self._load_txt_template("DA CIF.txt")

        if not template:
            template = """
RAPPORT D'ADEQUATION
V.2025-03

[Template non trouvé - veuillez vérifier le chemin vers DA CIF.txt]
"""

        # Préparer les remplacements - ce document a beaucoup de champs à remplir
        replacements = {
            "$DATE_JOUR$": self._format_date(datetime.now()),
            "$DATE_JOUR_JJJ_MM_AAAA$": self._format_date_long(),
            "$TITRE_CLIENT$": self._safe(client.t1_civilite),
            "$NOM_CLIENT$": self._safe(client.t1_nom, "").upper(),
            "$PRENOM_CLIENT$": self._safe(client.t1_prenom),
            "$OBJECTIFS$": self._safe(client.objectifs_investissement),
            "$HORIZON$": self._safe(client.horizon_placement),
            "$TOLERANCE_RISQUE$": self._safe(client.tolerance_risque),
            "$PROFIL_RISQUE$": self._safe(client.profil_risque_calcule, "À déterminer"),
        }

        # Remplacer les placeholders
        content = self._replace_placeholders(template, replacements)

        # Créer le document Word
        doc = self._create_document_from_text(content)

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("DECLARATION_ADEQUATION", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION CONVENTION RTO
    # ==========================================

    async def generate_convention_rto(self, client: Client, conseiller: User) -> str:
        """
        Générer la Convention de Réception et Transmission d'Ordres
        Utilise le template RTO.txt officiel
        """
        # Charger le template
        template = self._load_txt_template("RTO.txt")

        if not template:
            template = """
CONVENTION DE RECEPTION ET TRANSMISSION D'ORDRES

[Template non trouvé - veuillez vérifier le chemin vers RTO.txt]
"""

        # Ce document nécessite des informations spécifiques du client
        # Construire la description du client
        if client.t1_civilite == "M.":
            titre_long = "Monsieur"
        elif client.t1_civilite == "Mme":
            titre_long = "Madame"
        else:
            titre_long = client.t1_civilite or ""

        client_info = f"{titre_long} {self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        if client.t1_adresse:
            client_info += f", demeurant {client.t1_adresse}"
        if client.t1_profession:
            client_info += f", {client.t1_profession}"
        if client.t1_nationalite:
            client_info += f", de nationalité {client.t1_nationalite}"

        # Préparer les remplacements
        replacements = {
            "$DATE_JOUR$": self._format_date(datetime.now()),
            "$DATE_JOUR_JJJ_MM_AAAA$": self._format_date_long(),
            "$CLIENT_INFO$": client_info,
            "$TITRE_CLIENT$": self._safe(client.t1_civilite),
            "$NOM_CLIENT$": self._safe(client.t1_nom, "").upper(),
            "$PRENOM_CLIENT$": self._safe(client.t1_prenom),
            "$ADRESSE_CLIENT$": self._safe(client.t1_adresse),
        }

        # Remplacer les placeholders
        content = self._replace_placeholders(template, replacements)

        # Créer le document Word
        doc = self._create_document_from_text(content)

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("CONVENTION_RTO", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION PROFIL DE RISQUE
    # ==========================================

    def _get_field(self, client: Client, field_name: str, default: str = "") -> str:
        """
        Récupérer un champ du client, en cherchant d'abord dans form_data puis dans les attributs directs
        """
        # Chercher dans form_data d'abord (champs supplémentaires du frontend)
        if client.form_data and field_name in client.form_data:
            val = client.form_data[field_name]
            if val is not None:
                return str(val)

        # Sinon chercher dans les attributs directs du modèle
        if hasattr(client, field_name):
            val = getattr(client, field_name)
            if val is not None:
                return str(val)

        return default

    def _get_bool_field(self, client: Client, field_name: str) -> bool:
        """Récupérer un champ booléen"""
        if client.form_data and field_name in client.form_data:
            return bool(client.form_data[field_name])
        if hasattr(client, field_name):
            return bool(getattr(client, field_name))
        return False

    async def generate_profil_risque(self, client: Client, conseiller: User) -> str:
        """
        Générer le document Profil de Risque complet
        Conforme au questionnaire réglementaire QCC
        Utilise les données du client ET du form_data pour les champs supplémentaires
        """
        doc = Document()

        # Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # ==================== TITRE ====================
        p = doc.add_paragraph()
        run = p.add_run("QUESTIONNAIRE PROFIL DE RISQUE FINANCIER")
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        run = p.add_run("V.2025-03")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # ==================== EN-TÊTE ====================
        doc.add_paragraph(f"Date : {self._format_date(datetime.now())}")
        doc.add_paragraph(f"Conseiller : {self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}")
        doc.add_paragraph(f"Numéro client : {self._safe(client.numero_client)}")
        doc.add_paragraph()

        # ==================== IDENTIFICATION ====================
        p = doc.add_paragraph()
        run = p.add_run("IDENTIFICATION DU CLIENT")
        run.bold = True
        run.font.size = Pt(12)

        client_name = f"{self._safe(client.t1_civilite)} {self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        if client.t2_nom:
            client_name += f" et {self._safe(client.t2_civilite)} {self._safe(client.t2_prenom)} {self._safe(client.t2_nom, '').upper()}"
        doc.add_paragraph(f"Client : {client_name}")
        doc.add_paragraph(f"Date de naissance : {self._format_date(client.t1_date_naissance)}")
        doc.add_paragraph()

        # ==================== OBJECTIFS D'INVESTISSEMENT ====================
        p = doc.add_paragraph()
        run = p.add_run("1. OBJECTIFS D'INVESTISSEMENT")
        run.bold = True
        run.font.size = Pt(12)

        # Objectifs avec priorités
        objectifs = []
        if self._get_bool_field(client, 'objectif_preservation'):
            prio = self._get_field(client, 'objectif_preservation_priorite', '?')
            objectifs.append(f"Préservation du capital (priorité: {prio})")
        if self._get_bool_field(client, 'objectif_valorisation'):
            prio = self._get_field(client, 'objectif_valorisation_priorite', '?')
            objectifs.append(f"Valorisation du capital (priorité: {prio})")
        if self._get_bool_field(client, 'objectif_diversification'):
            prio = self._get_field(client, 'objectif_diversification_priorite', '?')
            objectifs.append(f"Diversification des actifs (priorité: {prio})")
        if self._get_bool_field(client, 'objectif_revenus'):
            prio = self._get_field(client, 'objectif_revenus_priorite', '?')
            objectifs.append(f"Recherche de revenus (priorité: {prio})")
        if self._get_bool_field(client, 'objectif_transmission'):
            prio = self._get_field(client, 'objectif_transmission_priorite', '?')
            objectifs.append(f"Transmission (priorité: {prio})")
        if self._get_bool_field(client, 'objectif_fiscal'):
            prio = self._get_field(client, 'objectif_fiscal_priorite', '?')
            objectifs.append(f"Optimisation fiscale (priorité: {prio})")
        autre_obj = self._get_field(client, 'objectif_autre')
        if autre_obj:
            prio = self._get_field(client, 'objectif_autre_priorite', '?')
            objectifs.append(f"{autre_obj} (priorité: {prio})")

        if objectifs:
            for obj in objectifs:
                doc.add_paragraph(f"☑ {obj}")
        else:
            doc.add_paragraph("Aucun objectif sélectionné")

        doc.add_paragraph(f"Horizon de placement : {self._get_field(client, 'horizon_placement')}")
        doc.add_paragraph(f"Besoin de liquidité : {self._get_field(client, 'besoin_liquidite')}")
        doc.add_paragraph(f"% du patrimoine à investir : {self._get_field(client, 'pourcentage_patrimoine_investi')}")
        doc.add_paragraph()

        # ==================== TOLÉRANCE AU RISQUE ====================
        p = doc.add_paragraph()
        run = p.add_run("2. TOLÉRANCE AU RISQUE")
        run.bold = True
        run.font.size = Pt(12)

        placement_pref = self._get_field(client, 'placement_preference')
        if placement_pref == 'A':
            doc.add_paragraph("☑ Placement A : Risque faible, protection du capital, diversification partielle")
        elif placement_pref == 'B':
            doc.add_paragraph("☑ Placement B : Risque moyen, diversification significative, recherche de valorisation")
        elif placement_pref == 'C':
            doc.add_paragraph("☑ Placement C : Risque élevé, maximiser la performance, perte partielle/totale possible")
        else:
            doc.add_paragraph(f"Préférence de placement : {placement_pref}")

        doc.add_paragraph(f"Expérience de baisse sur investissement : {self._get_field(client, 'experience_baisse')}")
        doc.add_paragraph(f"Réaction en cas de baisse : {self._get_field(client, 'reaction_perte')}")
        doc.add_paragraph(f"Réaction si +20% de gains : {self._get_field(client, 'reaction_hausse_20pct')}")
        doc.add_paragraph(f"Perte maximale acceptable : {self._get_field(client, 'pertes_maximales_acceptables')}")
        doc.add_paragraph()

        # ==================== CONNAISSANCE DES PRODUITS ====================
        p = doc.add_paragraph()
        run = p.add_run("3. CONNAISSANCE ET EXPÉRIENCE DES PRODUITS FINANCIERS")
        run.bold = True
        run.font.size = Pt(12)

        # Fonction pour afficher un bloc produit
        def add_product_section(name: str, prefix: str):
            detention = self._get_bool_field(client, f'{prefix}_detention')
            operations = self._get_field(client, f'{prefix}_operations')
            duree = self._get_field(client, f'{prefix}_duree')
            volume = self._get_field(client, f'{prefix}_volume')
            q1 = self._get_field(client, f'{prefix}_q1')
            q2 = self._get_field(client, f'{prefix}_q2')

            p = doc.add_paragraph()
            run = p.add_run(f"{name}")
            run.bold = True
            doc.add_paragraph(f"  Détention : {self._checkbox(detention)}")
            if detention or operations or duree or volume:
                doc.add_paragraph(f"  Opérations/an : {operations}")
                doc.add_paragraph(f"  Durée détention : {duree}")
                doc.add_paragraph(f"  Volume opérations : {volume}")
            if q1:
                doc.add_paragraph(f"  Question 1 : {q1}")
            if q2:
                doc.add_paragraph(f"  Question 2 : {q2}")

        add_product_section("Produits monétaires et fonds euros", "kyc_monetaires")
        add_product_section("Obligations et fonds obligataires", "kyc_obligations")
        add_product_section("Actions et fonds actions", "kyc_actions")
        add_product_section("SCPI / OPCI", "kyc_scpi")
        add_product_section("Private Equity (FCPI, FCPR, FIP)", "kyc_pe")
        add_product_section("ETF / Trackers", "kyc_etf")
        add_product_section("Produits dérivés", "kyc_derives")
        add_product_section("Produits structurés", "kyc_structures")
        doc.add_paragraph()

        # ==================== GESTION DU PORTEFEUILLE ====================
        p = doc.add_paragraph()
        run = p.add_run("4. GESTION DU PORTEFEUILLE")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(f"Portefeuille géré sous mandat : {self._checkbox(self._get_bool_field(client, 'gestion_mandat'))}")
        doc.add_paragraph(f"Gestion autonome : {self._checkbox(self._get_bool_field(client, 'gestion_autonome'))}")
        doc.add_paragraph(f"Gestion avec conseiller : {self._checkbox(self._get_bool_field(client, 'gestion_conseiller'))}")
        doc.add_paragraph(f"Expérience professionnelle en finance (>1 an) : {self._checkbox(self._get_bool_field(client, 'experience_professionnelle_finance'))}")
        doc.add_paragraph()

        # ==================== CULTURE FINANCIÈRE ====================
        p = doc.add_paragraph()
        run = p.add_run("5. CULTURE FINANCIÈRE")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(f"Lecture presse financière : {self._checkbox(self._get_bool_field(client, 'lecture_presse_financiere'))}")
        doc.add_paragraph(f"Suivi des cours de Bourse : {self._checkbox(self._get_bool_field(client, 'suivi_bourse'))}")
        doc.add_paragraph(f"Suivi mensuel des relevés bancaires : {self._checkbox(self._get_bool_field(client, 'suivi_releves'))}")
        doc.add_paragraph()

        # ==================== INVESTISSEMENTS DURABLES (ESG) ====================
        p = doc.add_paragraph()
        run = p.add_run("6. INVESTISSEMENTS DURABLES (ESG)")
        run.bold = True
        run.font.size = Pt(12)

        durabilite = self._get_bool_field(client, 'durabilite_integration')
        doc.add_paragraph(f"Souhait d'intégrer des critères ESG : {self._checkbox(durabilite)}")

        if durabilite:
            doc.add_paragraph(f"Part alignée Taxonomie européenne : {self._get_field(client, 'durabilite_taxonomie_part')}")
            doc.add_paragraph(f"Part en investissements durables : {self._get_field(client, 'durabilite_investissement_part')}")

            impact = self._get_bool_field(client, 'durabilite_impact')
            if impact:
                doc.add_paragraph("Incidences négatives à minimiser :")
                if self._get_bool_field(client, 'esg_gaz_effet_serre'):
                    doc.add_paragraph("  ☑ Gaz à effet de serre")
                if self._get_bool_field(client, 'esg_biodiversite'):
                    doc.add_paragraph("  ☑ Impact sur la biodiversité")
                if self._get_bool_field(client, 'esg_emissions_eau'):
                    doc.add_paragraph("  ☑ Émissions polluantes dans l'eau")
                if self._get_bool_field(client, 'esg_dechets'):
                    doc.add_paragraph("  ☑ Génération de déchets dangereux")
                if self._get_bool_field(client, 'esg_energie'):
                    doc.add_paragraph("  ☑ Inefficacité énergétique (immobilier)")
                if self._get_bool_field(client, 'esg_normes_internationales'):
                    doc.add_paragraph("  ☑ Non-respect normes internationales (OCDE, ONU)")
                if self._get_bool_field(client, 'esg_controle_normes'):
                    doc.add_paragraph("  ☑ Absence de processus de contrôle des normes")
                if self._get_bool_field(client, 'esg_egalite_remuneration'):
                    doc.add_paragraph("  ☑ Inégalité de rémunération H/F")
                if self._get_bool_field(client, 'esg_diversite_genres'):
                    doc.add_paragraph("  ☑ Manque de diversité au conseil d'administration")
                if self._get_bool_field(client, 'esg_armes_controversees'):
                    doc.add_paragraph("  ☑ Armes controversées (exclusion)")
        doc.add_paragraph()

        # ==================== PROFIL DE RISQUE DÉTERMINÉ ====================
        p = doc.add_paragraph()
        run = p.add_run("CONCLUSION - PROFIL DE RISQUE")
        run.bold = True
        run.font.size = Pt(14)

        profil = self._get_field(client, 'profil_risque_calcule', "Non déterminé")
        part_actifs = self._get_field(client, 'profil_part_actifs_risques')
        commentaire = self._get_field(client, 'profil_commentaire')

        p = doc.add_paragraph()
        run = p.add_run(f"PROFIL DÉTERMINÉ : {profil.upper()}")
        run.bold = True
        run.font.size = Pt(14)

        if part_actifs:
            doc.add_paragraph(f"Part maximale d'actifs à risque élevé : {part_actifs}%")

        if commentaire:
            doc.add_paragraph(f"Commentaires : {commentaire}")

        doc.add_paragraph()

        # ==================== DÉCLARATION CLIENT ====================
        p = doc.add_paragraph()
        run = p.add_run("DÉCLARATION DU CLIENT")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(
            "Le client déclare que les réponses à ce questionnaire sont exactes et sincères, "
            "qu'elles correspondent à sa situation actuelle et s'engage à informer de toute "
            "modification significative pouvant intervenir dans le futur."
        )
        doc.add_paragraph()

        # ==================== SIGNATURES ====================
        p = doc.add_paragraph()
        run = p.add_run("SIGNATURES")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph(f"Fait à _________________, le {self._format_date(datetime.now())}")
        doc.add_paragraph()

        # Tableau signatures
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Le Client"
        table.cell(0, 1).text = "Le Conseiller"
        table.cell(1, 0).text = f"{self._safe(client.t1_civilite)} {self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        table.cell(1, 1).text = f"{self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}"
        table.cell(2, 0).text = "(signature précédée de \"Lu et approuvé\")"
        table.cell(2, 1).text = ""

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("PROFIL_RISQUE", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION RAPPORT CONSEIL IAS
    # ==========================================

    async def generate_rapport_conseil_ias(self, client: Client, conseiller: User) -> str:
        """
        Générer le Rapport de Conseil en Assurance (IAS)
        Document généré dynamiquement
        """
        doc = Document()

        # Style
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Titre
        p = doc.add_paragraph()
        run = p.add_run("RAPPORT DE CONSEIL EN ASSURANCE")
        run.bold = True
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        run = p.add_run("Intermédiation en Assurance (IAS)")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # En-tête
        doc.add_paragraph(f"Date : {self._format_date(datetime.now())}")
        doc.add_paragraph(f"Conseiller : {self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}")
        doc.add_paragraph()

        # Client
        p = doc.add_paragraph()
        run = p.add_run("CLIENT")
        run.bold = True

        client_name = f"{self._safe(client.t1_civilite)} {self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        if client.t2_nom:
            client_name += f" et {self._safe(client.t2_civilite)} {self._safe(client.t2_prenom)} {self._safe(client.t2_nom, '').upper()}"
        doc.add_paragraph(f"Nom : {client_name}")
        doc.add_paragraph(f"Numéro client : {self._safe(client.numero_client)}")
        doc.add_paragraph()

        # Analyse des besoins
        p = doc.add_paragraph()
        run = p.add_run("1. ANALYSE DES BESOINS ET EXIGENCES")
        run.bold = True

        doc.add_paragraph(f"Objectifs : {self._safe(client.objectifs_investissement)}")
        doc.add_paragraph(f"Horizon de placement : {self._safe(client.horizon_placement)}")
        doc.add_paragraph(f"Tolérance au risque : {self._safe(client.tolerance_risque)}")
        doc.add_paragraph(f"Profil de risque : {self._safe(client.profil_risque_calcule, 'À déterminer')}")
        doc.add_paragraph(f"Besoin de liquidité : {self._checkbox(client.liquidite_importante)}")
        doc.add_paragraph()

        # Situation financière
        p = doc.add_paragraph()
        run = p.add_run("2. SITUATION FINANCIÈRE")
        run.bold = True

        doc.add_paragraph(f"Revenus annuels du foyer : {self._safe(client.revenus_annuels_foyer)}")
        doc.add_paragraph(f"Patrimoine global : {self._safe(client.patrimoine_global)}")
        doc.add_paragraph(f"Capacité d'épargne mensuelle : {self._format_montant(client.capacite_epargne_mensuelle)}")
        doc.add_paragraph()

        # Expérience
        p = doc.add_paragraph()
        run = p.add_run("3. CONNAISSANCES ET EXPÉRIENCE")
        run.bold = True

        experience = "Expert" if client.kyc_portefeuille_experience_pro else "Particulier non averti"
        doc.add_paragraph(f"Niveau d'expérience : {experience}")
        doc.add_paragraph(f"Gestion sous mandat : {self._checkbox(client.kyc_portefeuille_mandat)}")
        doc.add_paragraph(f"Gestion personnelle : {self._checkbox(client.kyc_portefeuille_gestion_personnelle)}")
        doc.add_paragraph()

        # Préférences ESG
        p = doc.add_paragraph()
        run = p.add_run("4. PRÉFÉRENCES EN MATIÈRE DE DURABILITÉ")
        run.bold = True

        doc.add_paragraph(f"Souhait ESG : {self._checkbox(client.durabilite_souhait)}")
        doc.add_paragraph(f"Niveau de préférence : {self._safe(client.durabilite_niveau_preference)}")
        doc.add_paragraph()

        # Recommandation
        p = doc.add_paragraph()
        run = p.add_run("5. RECOMMANDATION")
        run.bold = True

        doc.add_paragraph(
            "Sur la base de l'analyse de vos besoins, exigences, situation financière "
            "et préférences, nous vous recommandons le produit suivant :"
        )
        doc.add_paragraph()
        doc.add_paragraph("Produit recommandé : [À COMPLÉTER]")
        doc.add_paragraph("Assureur : [À COMPLÉTER]")
        doc.add_paragraph("Type de contrat : Assurance-vie")
        doc.add_paragraph("Montant : [À COMPLÉTER]")
        doc.add_paragraph()

        # Justification
        p = doc.add_paragraph()
        run = p.add_run("6. JUSTIFICATION DE L'ADÉQUATION")
        run.bold = True

        doc.add_paragraph("Le produit recommandé est cohérent avec :")
        doc.add_paragraph("• Vos objectifs d'investissement")
        doc.add_paragraph("• Votre horizon de placement")
        doc.add_paragraph("• Votre tolérance au risque")
        doc.add_paragraph("• Votre situation financière")
        doc.add_paragraph("• Vos préférences ESG")
        doc.add_paragraph()

        # Coûts
        p = doc.add_paragraph()
        run = p.add_run("7. INFORMATIONS SUR LES COÛTS")
        run.bold = True

        doc.add_paragraph(
            "Les frais détaillés du contrat vous seront communiqués dans le Document "
            "d'Information Clé (DIC) et les conditions générales du contrat."
        )
        doc.add_paragraph()

        # Signatures
        p = doc.add_paragraph()
        run = p.add_run("SIGNATURES")
        run.bold = True

        doc.add_paragraph(f"Fait à Papeete, le {self._format_date(datetime.now())}")
        doc.add_paragraph()
        doc.add_paragraph("Je reconnais avoir reçu ce rapport de conseil avant la souscription.")
        doc.add_paragraph()

        # Tableau signatures
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = "Le Client"
        table.cell(0, 1).text = "Le Conseiller"

        # Ajouter les mentions légales
        self._add_mentions_legales(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("RAPPORT_CONSEIL_IAS", client)
        return self._save_document(doc, filename)

    # ==========================================
    # GÉNÉRATION AVEC TEMPLATES V2 (DOCX avec placeholders)
    # ==========================================

    def _replace_placeholder_in_element(self, element, placeholder: str, value: str):
        """Remplace un placeholder dans un élément (paragraphe, run, etc.)"""
        if hasattr(element, 'text') and element.text:
            if placeholder in element.text:
                element.text = element.text.replace(placeholder, value)

    def _replace_placeholders_in_doc(self, doc: Document, replacements: Dict[str, str]):
        """
        Remplace tous les placeholders {{FIELD}} dans un document DOCX
        Parcourt les paragraphes, tables, en-têtes et pieds de page
        """
        # Parcourir les paragraphes
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text:
                    for placeholder, value in replacements.items():
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value) if value else "")

            # Aussi vérifier le texte complet du paragraphe (cas de placeholders fragmentés)
            full_text = para.text
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    # Reconstruire le texte si nécessaire
                    if para.runs:
                        # Trouver le run qui contient le début du placeholder
                        for run in para.runs:
                            for placeholder, value in replacements.items():
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(value) if value else "")

        # Parcourir les tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text:
                                for placeholder, value in replacements.items():
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(value) if value else "")

        # Parcourir les en-têtes et pieds de page
        for section in doc.sections:
            # En-tête
            for para in section.header.paragraphs:
                for run in para.runs:
                    if run.text:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value) if value else "")
            # Pied de page
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if run.text:
                        for placeholder, value in replacements.items():
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value) if value else "")

        # ==========================================
        # NETTOYAGE FINAL: Supprimer tous les {{...}} restants
        # ==========================================
        placeholder_pattern = re.compile(r'\{\{[^}]+\}\}')

        # Nettoyer les paragraphes
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text and placeholder_pattern.search(run.text):
                    run.text = placeholder_pattern.sub('', run.text)

        # Nettoyer les tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if run.text and placeholder_pattern.search(run.text):
                                run.text = placeholder_pattern.sub('', run.text)

        # Nettoyer les en-têtes et pieds de page
        for section in doc.sections:
            for para in section.header.paragraphs:
                for run in para.runs:
                    if run.text and placeholder_pattern.search(run.text):
                        run.text = placeholder_pattern.sub('', run.text)
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if run.text and placeholder_pattern.search(run.text):
                        run.text = placeholder_pattern.sub('', run.text)

    def _build_client_replacements(self, client: Client, conseiller: User) -> Dict[str, str]:
        """
        Construit le dictionnaire de remplacement pour un client
        Mappe tous les champs vers les placeholders {{FIELD}}
        """
        replacements = {}

        # Récupérer les dates depuis form_data (contexteMission)
        date_remise_der = None
        date_signature_mission = None
        lieu_signature = "Papeete"
        if client.form_data:
            contexte_mission = client.form_data.get('contexteMission', {})
            date_remise_der = contexte_mission.get('dateRemiseDER')
            date_signature_mission = contexte_mission.get('dateSignature')
            lieu_signature = contexte_mission.get('lieuSignature', 'Papeete')

        # Informations générales
        replacements["{{NUMERO_CLIENT}}"] = self._safe(client.numero_client)
        replacements["{{DATE_REMISE_DER}}"] = self._format_date(date_remise_der) if date_remise_der else ""
        replacements["{{DATE_SIGNATURE}}"] = self._format_date(date_signature_mission) if date_signature_mission else self._format_date(datetime.now())
        replacements["{{DATE_JOUR}}"] = self._format_date(datetime.now())
        replacements["{{LIEU_SIGNATURE}}"] = lieu_signature
        replacements["{{NOMBRE_EXEMPLAIRES}}"] = "2"

        # Nom complet
        nom_complet_t1 = f"{self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        replacements["{{NOM_COMPLET_T1}}"] = nom_complet_t1

        # Titulaire 1
        replacements["{{T1_CIVILITE}}"] = self._safe(client.t1_civilite)
        replacements["{{T1_NOM}}"] = self._safe(client.t1_nom, "").upper()
        replacements["{{T1_PRENOM}}"] = self._safe(client.t1_prenom)
        replacements["{{T1_NOM_JEUNE_FILLE}}"] = self._get_field(client, 't1_nom_jeune_fille')
        replacements["{{T1_DATE_NAISSANCE}}"] = self._format_date(client.t1_date_naissance)
        replacements["{{T1_LIEU_NAISSANCE}}"] = self._get_field(client, 't1_lieu_naissance')
        replacements["{{T1_NATIONALITE}}"] = self._get_field(client, 't1_nationalite', 'Française')
        replacements["{{T1_ADRESSE}}"] = self._safe(client.t1_adresse)
        replacements["{{T1_CODE_POSTAL}}"] = self._get_field(client, 't1_code_postal')
        replacements["{{T1_VILLE}}"] = self._get_field(client, 't1_ville')
        replacements["{{T1_EMAIL}}"] = self._safe(client.t1_email)
        replacements["{{T1_TELEPHONE}}"] = self._safe(client.t1_telephone)
        replacements["{{T1_PIECE_IDENTITE}}"] = self._get_field(client, 't1_piece_identite')
        replacements["{{T1_NUMERO_PIECE}}"] = self._get_field(client, 't1_numero_piece')
        replacements["{{T1_DATE_VALIDITE_PIECE}}"] = self._format_date(self._get_field(client, 't1_date_validite_piece'))
        replacements["{{T1_RESIDENCE_FISCALE}}"] = self._get_field(client, 't1_residence_fiscale', 'France')
        replacements["{{T1_NIF}}"] = self._get_field(client, 't1_nif')
        replacements["{{T1_US_PERSON}}"] = "Oui" if self._get_bool_field(client, 't1_us_person') else "Non"
        replacements["{{T1_REGIME_PROTECTION_JURIDIQUE}}"] = "Oui" if self._get_bool_field(client, 't1_regime_protection_juridique') else "Non"
        replacements["{{T1_REGIME_PROTECTION_FORME}}"] = self._get_field(client, 't1_regime_protection_forme')
        replacements["{{T1_REPRESENTANT_LEGAL}}"] = self._get_field(client, 't1_representant_legal')
        replacements["{{T1_SITUATION_PRO}}"] = self._get_field(client, 't1_situation_pro')
        replacements["{{T1_PROFESSION}}"] = self._safe(client.t1_profession)
        replacements["{{T1_SECTEUR_ACTIVITE}}"] = self._get_field(client, 't1_secteur_activite')
        replacements["{{T1_EMPLOYEUR}}"] = self._get_field(client, 't1_employeur')
        replacements["{{T1_CHEF_ENTREPRISE}}"] = "Oui" if self._get_bool_field(client, 't1_chef_entreprise') else "Non"
        replacements["{{T1_ENTREPRISE_DENOMINATION}}"] = self._get_field(client, 't1_entreprise_denomination')
        replacements["{{T1_ENTREPRISE_FORME_JURIDIQUE}}"] = self._get_field(client, 't1_entreprise_forme_juridique')
        replacements["{{T1_ENTREPRISE_SIEGE_SOCIAL}}"] = self._get_field(client, 't1_entreprise_siege_social')

        # Titulaire 2
        if client.t2_nom:
            replacements["{{T2_CIVILITE}}"] = self._safe(client.t2_civilite)
            replacements["{{T2_NOM}}"] = self._safe(client.t2_nom, "").upper()
            replacements["{{T2_PRENOM}}"] = self._safe(client.t2_prenom)
            replacements["{{T2_NOM_JEUNE_FILLE}}"] = self._get_field(client, 't2_nom_jeune_fille')
            replacements["{{T2_DATE_NAISSANCE}}"] = self._format_date(client.t2_date_naissance)
            replacements["{{T2_LIEU_NAISSANCE}}"] = self._get_field(client, 't2_lieu_naissance')
            replacements["{{T2_NATIONALITE}}"] = self._get_field(client, 't2_nationalite')
            replacements["{{T2_ADRESSE}}"] = self._safe(client.t2_adresse)
            replacements["{{T2_CODE_POSTAL}}"] = self._get_field(client, 't2_code_postal')
            replacements["{{T2_VILLE}}"] = self._get_field(client, 't2_ville')
            replacements["{{T2_EMAIL}}"] = self._safe(client.t2_email)
            replacements["{{T2_TELEPHONE}}"] = self._safe(client.t2_telephone)
            replacements["{{T2_PIECE_IDENTITE}}"] = self._get_field(client, 't2_piece_identite')
            replacements["{{T2_NUMERO_PIECE}}"] = self._get_field(client, 't2_numero_piece')
            replacements["{{T2_DATE_VALIDITE_PIECE}}"] = self._format_date(self._get_field(client, 't2_date_validite_piece'))
            replacements["{{T2_RESIDENCE_FISCALE}}"] = self._get_field(client, 't2_residence_fiscale')
            replacements["{{T2_US_PERSON}}"] = "Oui" if self._get_bool_field(client, 't2_us_person') else "Non"
            replacements["{{T2_PROFESSION}}"] = self._safe(client.t2_profession)
        else:
            # Titulaire 2 vide
            for field in ['T2_CIVILITE', 'T2_NOM', 'T2_PRENOM', 'T2_NOM_JEUNE_FILLE', 'T2_DATE_NAISSANCE',
                          'T2_LIEU_NAISSANCE', 'T2_NATIONALITE', 'T2_ADRESSE', 'T2_CODE_POSTAL', 'T2_VILLE',
                          'T2_EMAIL', 'T2_TELEPHONE', 'T2_PIECE_IDENTITE', 'T2_NUMERO_PIECE',
                          'T2_DATE_VALIDITE_PIECE', 'T2_RESIDENCE_FISCALE', 'T2_US_PERSON', 'T2_PROFESSION']:
                replacements[f"{{{{{field}}}}}"] = ""

        # Situation familiale
        replacements["{{SITUATION_FAMILIALE}}"] = self._safe(client.situation_familiale)
        replacements["{{DATE_MARIAGE}}"] = self._format_date(client.date_mariage)
        replacements["{{CONTRAT_MARIAGE}}"] = "Oui" if self._get_bool_field(client, 'contrat_mariage') else "Non"
        replacements["{{REGIME_MATRIMONIAL}}"] = self._get_field(client, 'regime_matrimonial')
        replacements["{{DATE_PACS}}"] = self._format_date(client.date_pacs)
        replacements["{{CONVENTION_PACS}}"] = "Oui" if self._get_bool_field(client, 'convention_pacs') else "Non"
        replacements["{{REGIME_PACS}}"] = self._get_field(client, 'regime_pacs')
        replacements["{{DATE_DIVORCE}}"] = self._format_date(client.date_divorce)
        replacements["{{NOMBRE_ENFANTS}}"] = str(client.nombre_enfants or 0)
        replacements["{{NOMBRE_ENFANTS_CHARGE}}"] = str(client.nombre_enfants_charge or 0)

        # Donations
        replacements["{{DONATION_ENTRE_EPOUX}}"] = "Oui" if self._get_bool_field(client, 'donation_entre_epoux') else "Non"
        replacements["{{DONATION_ENTRE_EPOUX_DATE}}"] = self._format_date(client.donation_entre_epoux_date)
        replacements["{{DONATION_ENTRE_EPOUX_MONTANT}}"] = self._format_montant(client.donation_entre_epoux_montant)
        replacements["{{DONATION_ENFANTS}}"] = "Oui" if self._get_bool_field(client, 'donation_enfants') else "Non"
        replacements["{{DONATION_ENFANTS_DATE}}"] = self._format_date(client.donation_enfants_date)
        replacements["{{DONATION_ENFANTS_MONTANT}}"] = self._format_montant(client.donation_enfants_montant)

        # Sections conditionnelles situation familiale
        situation = self._safe(client.situation_familiale, '').lower()
        replacements["{{#IF_MARIE}}"] = "" if 'marié' in situation else "<!-- MASQUÉ -->"
        replacements["{{/IF_MARIE}}"] = "" if 'marié' in situation else "<!-- /MASQUÉ -->"
        replacements["{{#IF_PACSE}}"] = "" if 'pacsé' in situation or 'pacs' in situation else "<!-- MASQUÉ -->"
        replacements["{{/IF_PACSE}}"] = "" if 'pacsé' in situation or 'pacs' in situation else "<!-- /MASQUÉ -->"
        replacements["{{#IF_DIVORCE}}"] = "" if 'divorcé' in situation else "<!-- MASQUÉ -->"
        replacements["{{/IF_DIVORCE}}"] = "" if 'divorcé' in situation else "<!-- /MASQUÉ -->"

        # Situation financière
        replacements["{{REVENUS_ANNUELS_FOYER}}"] = self._safe(client.revenus_annuels_foyer)
        replacements["{{PATRIMOINE_GLOBAL}}"] = self._safe(client.patrimoine_global)
        replacements["{{CHARGES_ANNUELLES_POURCENT}}"] = str(client.charges_annuelles_pourcent or "")
        replacements["{{CHARGES_ANNUELLES_MONTANT}}"] = self._format_montant(client.charges_annuelles_montant)
        replacements["{{CAPACITE_EPARGNE_MENSUELLE}}"] = self._format_montant(client.capacite_epargne_mensuelle)
        replacements["{{IMPOT_REVENU}}"] = "Oui" if self._get_bool_field(client, 'impot_revenu') else "Non"
        replacements["{{IMPOT_FORTUNE_IMMOBILIERE}}"] = "Oui" if self._get_bool_field(client, 'impot_fortune_immobiliere') else "Non"
        replacements["{{PATRIMOINE_FINANCIER_POURCENT}}"] = str(client.patrimoine_financier_pourcent or "")
        replacements["{{PATRIMOINE_IMMOBILIER_POURCENT}}"] = str(client.patrimoine_immobilier_pourcent or "")
        replacements["{{PATRIMOINE_PROFESSIONNEL_POURCENT}}"] = str(client.patrimoine_professionnel_pourcent or "")
        replacements["{{PATRIMOINE_AUTRES_POURCENT}}"] = str(client.patrimoine_autres_pourcent or "")

        # Origine des fonds
        replacements["{{ORIGINE_FONDS_NATURE}}"] = self._get_field(client, 'origine_fonds_nature')
        replacements["{{ORIGINE_FONDS_MONTANT_PREVU}}"] = self._format_montant(self._get_field(client, 'origine_fonds_montant_prevu'))

        # Ancien format (compatibilité)
        replacements["{{ORIGINE_ECONOMIQUE_REVENUS}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_revenus'))
        replacements["{{ORIGINE_ECONOMIQUE_EPARGNE}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_epargne'))
        replacements["{{ORIGINE_ECONOMIQUE_HERITAGE}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_heritage'))
        replacements["{{ORIGINE_ECONOMIQUE_CESSION_PRO}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_cession_pro'))
        replacements["{{ORIGINE_ECONOMIQUE_CESSION_IMMO}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_cession_immo'))
        replacements["{{ORIGINE_ECONOMIQUE_CESSION_MOBILIERE}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_cession_mobiliere'))
        replacements["{{ORIGINE_ECONOMIQUE_GAINS_JEU}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_gains_jeu'))
        replacements["{{ORIGINE_ECONOMIQUE_ASSURANCE_VIE}}"] = self._checkbox(self._get_bool_field(client, 'origine_economique_assurance_vie'))
        replacements["{{ORIGINE_ECONOMIQUE_AUTRES}}"] = self._get_field(client, 'origine_economique_autres')
        replacements["{{ORIGINE_FONDS_PROVENANCE_ETABLISSEMENT}}"] = self._get_field(client, 'origine_fonds_provenance_etablissement')

        # Nouveau format: checkbox + montant (pour QCC v2)
        origins = [
            ('REVENUS', 'origine_economique_revenus', 'origine_revenus_montant'),
            ('EPARGNE', 'origine_economique_epargne', 'origine_epargne_montant'),
            ('HERITAGE', 'origine_economique_heritage', 'origine_heritage_montant'),
            ('CESSION_PRO', 'origine_economique_cession_pro', 'origine_cession_pro_montant'),
            ('CESSION_IMMO', 'origine_economique_cession_immo', 'origine_cession_immo_montant'),
            ('CESSION_MOBILIERE', 'origine_economique_cession_mobiliere', 'origine_cession_mobiliere_montant'),
            ('GAINS_JEU', 'origine_economique_gains_jeu', 'origine_gains_jeu_montant'),
            ('ASSURANCE_VIE', 'origine_economique_assurance_vie', 'origine_assurance_vie_montant'),
            ('AUTRES', 'origine_economique_autres', 'origine_autres_montant'),
        ]
        for key, check_field, montant_field in origins:
            is_checked = self._get_bool_field(client, check_field)
            replacements[f"{{{{ORIGINE_{key}_CHECK}}}}"] = "☑" if is_checked else "☐"
            replacements[f"{{{{ORIGINE_{key}_MONTANT}}}}"] = self._format_montant(self._get_field(client, montant_field)) if is_checked else ""

        # LCB-FT
        replacements["{{LCB_FT_PPE}}"] = "Oui" if self._get_bool_field(client, 'lcb_ft_ppe') else "Non"
        replacements["{{LCB_FT_PPE_FONCTION}}"] = self._get_field(client, 'lcb_ft_ppe_fonction')
        replacements["{{LCB_FT_PPE_FAMILLE}}"] = "Oui" if self._get_bool_field(client, 'lcb_ft_ppe_famille') else "Non"
        replacements["{{LCB_FT_GEL_AVOIRS_VERIFIE}}"] = "Oui" if self._get_bool_field(client, 'lcb_ft_gel_avoirs_verifie') else "Non"
        replacements["{{LCB_FT_GEL_AVOIRS_DATE_VERIFICATION}}"] = self._format_date(self._get_field(client, 'lcb_ft_gel_avoirs_date_verification'))
        replacements["{{LCB_FT_NIVEAU_RISQUE}}"] = self._get_field(client, 'lcb_ft_niveau_risque')

        # KYC - Produits financiers
        for product in ['monetaires', 'obligations', 'actions', 'scpi', 'pe', 'etf', 'derives', 'structures']:
            prefix = f'kyc_{product}'
            replacements[f"{{{{KYC_{product.upper()}_DETENTION}}}}"] = "Oui" if self._get_bool_field(client, f'{prefix}_detention') else "Non"
            replacements[f"{{{{KYC_{product.upper()}_OPERATIONS}}}}"] = self._get_field(client, f'{prefix}_operations')
            replacements[f"{{{{KYC_{product.upper()}_DUREE}}}}"] = self._get_field(client, f'{prefix}_duree')
            replacements[f"{{{{KYC_{product.upper()}_VOLUME}}}}"] = self._get_field(client, f'{prefix}_volume')
            replacements[f"{{{{KYC_{product.upper()}_Q1}}}}"] = self._get_field(client, f'{prefix}_q1')
            replacements[f"{{{{KYC_{product.upper()}_Q2}}}}"] = self._get_field(client, f'{prefix}_q2')

        # Gestion portefeuille
        replacements["{{KYC_PORTEFEUILLE_MANDAT}}"] = "Oui" if self._get_bool_field(client, 'kyc_portefeuille_mandat') else "Non"
        replacements["{{KYC_PORTEFEUILLE_GESTION_PERSONNELLE}}"] = "Oui" if self._get_bool_field(client, 'kyc_portefeuille_gestion_personnelle') else "Non"
        replacements["{{KYC_PORTEFEUILLE_GESTION_CONSEILLER}}"] = "Oui" if self._get_bool_field(client, 'kyc_portefeuille_gestion_conseiller') else "Non"
        replacements["{{KYC_PORTEFEUILLE_EXPERIENCE_PRO}}"] = "Oui" if self._get_bool_field(client, 'kyc_portefeuille_experience_pro') else "Non"

        # Culture financière
        replacements["{{KYC_CULTURE_PRESSE_FINANCIERE}}"] = "Oui" if self._get_bool_field(client, 'kyc_culture_presse_financiere') else "Non"
        replacements["{{KYC_CULTURE_SUIVI_BOURSE}}"] = "Oui" if self._get_bool_field(client, 'kyc_culture_suivi_bourse') else "Non"
        replacements["{{KYC_CULTURE_RELEVES_BANCAIRES}}"] = "Oui" if self._get_bool_field(client, 'kyc_culture_releves_bancaires') else "Non"

        # Profil de risque
        replacements["{{OBJECTIFS_INVESTISSEMENT}}"] = self._safe(client.objectifs_investissement)
        replacements["{{HORIZON_PLACEMENT}}"] = self._safe(client.horizon_placement)
        replacements["{{TOLERANCE_RISQUE}}"] = self._safe(client.tolerance_risque)
        replacements["{{PERTES_MAXIMALES_ACCEPTABLES}}"] = self._safe(client.pertes_maximales_acceptables)
        replacements["{{EXPERIENCE_PERTE}}"] = "Oui" if self._get_bool_field(client, 'experience_perte') else "Non"
        replacements["{{EXPERIENCE_PERTE_NIVEAU}}"] = self._get_field(client, 'experience_perte_niveau')
        replacements["{{REACTION_PERTE}}"] = self._get_field(client, 'reaction_perte')
        replacements["{{REACTION_GAIN}}"] = self._get_field(client, 'reaction_gain')
        replacements["{{LIQUIDITE_IMPORTANTE}}"] = "Oui" if self._get_bool_field(client, 'liquidite_importante') else "Non"
        replacements["{{POURCENTAGE_PATRIMOINE_INVESTI}}"] = self._get_field(client, 'pourcentage_patrimoine_investi')

        # Durabilité ESG
        replacements["{{DURABILITE_SOUHAIT}}"] = "Oui" if self._get_bool_field(client, 'durabilite_souhait') else "Non"
        replacements["{{DURABILITE_TAXONOMIE_POURCENT}}"] = self._get_field(client, 'durabilite_taxonomie_pourcent')
        replacements["{{DURABILITE_INVESTISSEMENTS_POURCENT}}"] = self._get_field(client, 'durabilite_investissements_pourcent')
        replacements["{{DURABILITE_IMPACT_SELECTION}}"] = "Oui" if self._get_bool_field(client, 'durabilite_impact_selection') else "Non"
        replacements["{{DURABILITE_CRITERES}}"] = self._get_field(client, 'durabilite_criteres')

        # Profil calculé
        replacements["{{PROFIL_RISQUE_CALCULE}}"] = self._safe(client.profil_risque_calcule, "Non déterminé")
        replacements["{{PROFIL_RISQUE_SCORE}}"] = str(client.profil_risque_score or "")
        replacements["{{PROFIL_RISQUE_DATE_CALCUL}}"] = self._format_date(client.profil_risque_date_calcul)
        replacements["{{PROFIL_COMMENTAIRE}}"] = self._get_field(client, 'profil_commentaire')

        # ==========================================
        # TABLEAUX JSONB - Enfants
        # ==========================================
        enfants = client.enfants or []
        for i in range(1, 11):  # Support jusqu'à 10 enfants
            if i <= len(enfants):
                enfant = enfants[i-1]
                replacements[f"{{{{ENFANT_{i}_NOM}}}}"] = f"{enfant.get('prenom', '')} {enfant.get('nom', '')}"
                replacements[f"{{{{ENFANT_{i}_DATE_NAISSANCE}}}}"] = self._format_date(enfant.get('date_naissance'))
                replacements[f"{{{{ENFANT_{i}_LIEN}}}}"] = enfant.get('lien_parente', '')
                replacements[f"{{{{ENFANT_{i}_A_CHARGE}}}}"] = "Oui" if enfant.get('a_charge') else "Non"
                replacements[f"{{{{ENFANT_{i}_A_ENFANTS}}}}"] = "Oui" if enfant.get('a_enfants') else "Non"
            else:
                # Lignes vides pour les enfants non existants
                replacements[f"{{{{ENFANT_{i}_NOM}}}}"] = ""
                replacements[f"{{{{ENFANT_{i}_DATE_NAISSANCE}}}}"] = ""
                replacements[f"{{{{ENFANT_{i}_LIEN}}}}"] = ""
                replacements[f"{{{{ENFANT_{i}_A_CHARGE}}}}"] = ""
                replacements[f"{{{{ENFANT_{i}_A_ENFANTS}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Patrimoine Financier
        # ==========================================
        patrimoine_fin = client.patrimoine_financier or []
        for i in range(1, 11):  # Support jusqu'à 10 actifs financiers
            if i <= len(patrimoine_fin):
                actif = patrimoine_fin[i-1]
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DESIGNATION}}}}"] = actif.get('designation', '')
                replacements[f"{{{{PATRIMOINE_FIN_{i}_ORGANISME}}}}"] = actif.get('organisme', '')
                replacements[f"{{{{PATRIMOINE_FIN_{i}_VALEUR}}}}"] = self._format_montant(actif.get('valeur'))
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DETENTEUR}}}}"] = actif.get('detenteur', 'Titulaire 1')
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DATE}}}}"] = self._format_date(actif.get('date_souscription'))
                replacements[f"{{{{PATRIMOINE_FIN_{i}_REMARQUES}}}}"] = actif.get('remarques', '')
            else:
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DESIGNATION}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_FIN_{i}_ORGANISME}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_FIN_{i}_VALEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DETENTEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_FIN_{i}_DATE}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_FIN_{i}_REMARQUES}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Patrimoine Immobilier
        # ==========================================
        patrimoine_immo = client.patrimoine_immobilier or []
        for i in range(1, 11):  # Support jusqu'à 10 biens immobiliers
            if i <= len(patrimoine_immo):
                bien = patrimoine_immo[i-1]
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_DESIGNATION}}}}"] = bien.get('designation', '')
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_DETENTEUR}}}}"] = bien.get('detenteur', 'Titulaire 1')
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_FORME}}}}"] = bien.get('forme_propriete', '')
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_VALEUR}}}}"] = self._format_montant(bien.get('valeur_actuelle'))
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_REVENUS}}}}"] = self._format_montant(bien.get('revenus_annuels'))
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_CREDITS}}}}"] = "Oui" if bien.get('credit_en_cours') else "Non"
            else:
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_DESIGNATION}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_DETENTEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_FORME}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_VALEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_REVENUS}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_IMMO_{i}_CREDITS}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Patrimoine Professionnel
        # ==========================================
        patrimoine_pro = client.patrimoine_professionnel or []
        for i in range(1, 11):  # Support jusqu'à 10 actifs professionnels
            if i <= len(patrimoine_pro):
                actif = patrimoine_pro[i-1]
                replacements[f"{{{{PATRIMOINE_PRO_{i}_DESIGNATION}}}}"] = actif.get('designation', '')
                replacements[f"{{{{PATRIMOINE_PRO_{i}_DETENTEUR}}}}"] = actif.get('detenteur', '')
                replacements[f"{{{{PATRIMOINE_PRO_{i}_VALEUR}}}}"] = self._format_montant(actif.get('valeur_patrimoniale'))
                replacements[f"{{{{PATRIMOINE_PRO_{i}_CHARGES}}}}"] = self._format_montant(actif.get('charges'))
                replacements[f"{{{{PATRIMOINE_PRO_{i}_REMARQUES}}}}"] = actif.get('remarques', '')
            else:
                replacements[f"{{{{PATRIMOINE_PRO_{i}_DESIGNATION}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_PRO_{i}_DETENTEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_PRO_{i}_VALEUR}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_PRO_{i}_CHARGES}}}}"] = ""
                replacements[f"{{{{PATRIMOINE_PRO_{i}_REMARQUES}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Emprunts (Passif)
        # ==========================================
        emprunts = client.patrimoine_emprunts or []
        for i in range(1, 11):
            if i <= len(emprunts):
                emprunt = emprunts[i-1]
                replacements[f"{{{{EMPRUNT_{i}_DESIGNATION}}}}"] = emprunt.get('designation', '')
                replacements[f"{{{{EMPRUNT_{i}_ORGANISME}}}}"] = emprunt.get('organisme', '')
                replacements[f"{{{{EMPRUNT_{i}_CAPITAL}}}}"] = self._format_montant(emprunt.get('capital_restant'))
                replacements[f"{{{{EMPRUNT_{i}_MENSUALITE}}}}"] = self._format_montant(emprunt.get('mensualite'))
                replacements[f"{{{{EMPRUNT_{i}_FIN}}}}"] = self._format_date(emprunt.get('date_fin'))
            else:
                replacements[f"{{{{EMPRUNT_{i}_DESIGNATION}}}}"] = ""
                replacements[f"{{{{EMPRUNT_{i}_ORGANISME}}}}"] = ""
                replacements[f"{{{{EMPRUNT_{i}_CAPITAL}}}}"] = ""
                replacements[f"{{{{EMPRUNT_{i}_MENSUALITE}}}}"] = ""
                replacements[f"{{{{EMPRUNT_{i}_FIN}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Revenus détaillés
        # ==========================================
        revenus = client.patrimoine_revenus or []
        for i in range(1, 11):
            if i <= len(revenus):
                revenu = revenus[i-1]
                replacements[f"{{{{REVENU_{i}_TYPE}}}}"] = revenu.get('type', '')
                replacements[f"{{{{REVENU_{i}_MONTANT}}}}"] = self._format_montant(revenu.get('montant_annuel'))
                replacements[f"{{{{REVENU_{i}_BENEFICIAIRE}}}}"] = revenu.get('beneficiaire', '')
            else:
                replacements[f"{{{{REVENU_{i}_TYPE}}}}"] = ""
                replacements[f"{{{{REVENU_{i}_MONTANT}}}}"] = ""
                replacements[f"{{{{REVENU_{i}_BENEFICIAIRE}}}}"] = ""

        # ==========================================
        # TABLEAUX JSONB - Charges détaillées
        # ==========================================
        charges = client.patrimoine_charges or []
        for i in range(1, 11):
            if i <= len(charges):
                charge = charges[i-1]
                replacements[f"{{{{CHARGE_{i}_TYPE}}}}"] = charge.get('type', '')
                replacements[f"{{{{CHARGE_{i}_MONTANT}}}}"] = self._format_montant(charge.get('montant_annuel'))
                replacements[f"{{{{CHARGE_{i}_REMARQUE}}}}"] = charge.get('remarque', '')
            else:
                replacements[f"{{{{CHARGE_{i}_TYPE}}}}"] = ""
                replacements[f"{{{{CHARGE_{i}_MONTANT}}}}"] = ""
                replacements[f"{{{{CHARGE_{i}_REMARQUE}}}}"] = ""

        # Conseiller
        replacements["{{SIGNATURE_CONSEILLER}}"] = f"{self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}"
        replacements["{{SIGNATURE_CLIENT}}"] = nom_complet_t1

        return replacements

    async def generate_qcc_v2(self, client: Client, conseiller: User) -> str:
        """
        Générer le QCC avec le template DOCX v2
        Document professionnel avec placeholders remplacés
        """
        template_path = os.path.join(self.v2_templates_path, "QCC_V2_TEMPLATE.docx")

        if not os.path.exists(template_path):
            # Fallback vers l'ancienne méthode
            return await self.generate_kyc(client, conseiller)

        # Charger le template
        doc = Document(template_path)

        # Construire les remplacements
        replacements = self._build_client_replacements(client, conseiller)

        # Remplacer les placeholders
        self._replace_placeholders_in_doc(doc, replacements)

        # Supprimer les sections conditionnelles masquées
        self._remove_hidden_sections(doc)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("QCC", client)
        return self._save_document(doc, filename)

    def _remove_hidden_sections(self, doc: Document):
        """
        Supprime les paragraphes contenant les marqueurs de sections masquées
        Les marqueurs <!-- MASQUÉ --> et <!-- /MASQUÉ --> indiquent des sections à supprimer
        """
        paragraphs_to_remove = []
        in_hidden_section = False

        for para in doc.paragraphs:
            text = para.text
            if "<!-- MASQUÉ -->" in text:
                in_hidden_section = True
                paragraphs_to_remove.append(para)
            elif "<!-- /MASQUÉ -->" in text:
                in_hidden_section = False
                paragraphs_to_remove.append(para)
            elif in_hidden_section:
                paragraphs_to_remove.append(para)

        # Supprimer les paragraphes marqués
        for para in paragraphs_to_remove:
            p = para._element
            p.getparent().remove(p)

        # Aussi traiter les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        if "<!-- MASQUÉ -->" in text or "<!-- /MASQUÉ -->" in text:
                            para.clear()

    def _translate_value(self, value: str, translations: Dict[str, str]) -> str:
        """Traduit une valeur en utilisant un dictionnaire de traductions"""
        if value and value in translations:
            return translations[value]
        return value or ""

    def _build_profil_risque_replacements(self, client: Client, conseiller: User) -> Dict[str, str]:
        """
        Construit le dictionnaire de remplacement spécifique au Profil de Risque
        Inclut les traductions des valeurs et les objectifs avec priorités
        """
        # Commencer avec les remplacements génériques
        replacements = self._build_client_replacements(client, conseiller)

        # Traductions pour reaction_perte
        reaction_perte_translations = {
            'acheter_plus': 'Investir de nouveau pour profiter des opportunités',
            'vendre_tout': 'Tout vendre pour réinvestir dans des supports moins risqués',
            'vendre_partie': 'Vendre seulement une partie',
            'ne_rien_changer': 'Ne rien changer',
            'rien_changer': 'Ne rien changer',
        }

        # Traductions pour reaction_gain
        reaction_gain_translations = {
            'conserver': 'Je conserve ma position',
            'reinvestir_egal': 'Je réinvestis un montant inférieur ou égal au montant initial',
            'reinvestir_plus': 'Je réinvestis un montant supérieur au montant initial',
        }

        # Traductions pour horizon_placement
        horizon_translations = {
            '< 1 an': 'Moins d\'1 an',
            '1-3 ans': 'Entre 1 et 3 ans',
            '3-5 ans': 'Entre 3 et 5 ans',
            '> 5 ans': 'Plus de 5 ans',
        }

        # Traductions pour pertes_maximales
        pertes_translations = {
            '0': 'Aucune',
            '10': 'Max 10%',
            '25': 'Max 25%',
            '50': 'Max 50%',
            '100': 'Jusqu\'à 100%',
        }

        # Traductions pour pourcentage_patrimoine
        patrimoine_translations = {
            '< 10%': 'Moins de 10%',
            '10-25%': 'Entre 10% et 25%',
            '25-50%': 'Entre 25% et 50%',
            '50-75%': 'Entre 50% et 75%',
            '> 75%': 'Plus de 75%',
        }

        # Appliquer les traductions
        reaction_perte = self._get_field(client, 'reaction_perte')
        replacements["{{REACTION_PERTE}}"] = self._translate_value(reaction_perte, reaction_perte_translations)

        reaction_gain = self._get_field(client, 'reaction_gain')
        replacements["{{REACTION_GAIN}}"] = self._translate_value(reaction_gain, reaction_gain_translations)

        horizon = self._get_field(client, 'horizon_placement')
        replacements["{{HORIZON_PLACEMENT}}"] = self._translate_value(horizon, horizon_translations)

        pertes = self._get_field(client, 'pertes_maximales_acceptables')
        replacements["{{PERTES_MAXIMALES_ACCEPTABLES}}"] = self._translate_value(pertes, pertes_translations)

        patrimoine = self._get_field(client, 'pourcentage_patrimoine_investi')
        replacements["{{POURCENTAGE_PATRIMOINE_INVESTI}}"] = self._translate_value(patrimoine, patrimoine_translations)

        # Objectifs d'investissement avec priorités
        objectifs = ['preservation', 'valorisation', 'diversification', 'revenus', 'transmission', 'fiscal']
        objectifs_labels = {
            'preservation': 'Préservation du capital',
            'valorisation': 'Valorisation de capital',
            'diversification': 'Diversification des actifs détenus',
            'revenus': 'Recherche de revenus',
            'transmission': 'Transmission',
            'fiscal': 'Optimisation fiscale',
        }

        for obj in objectifs:
            is_selected = self._get_bool_field(client, f'objectif_{obj}')
            priority = self._get_field(client, f'objectif_{obj}_priorite')
            replacements[f"{{{{OBJECTIF_{obj.upper()}_SELECTED}}}}"] = "☑" if is_selected else "☐"
            replacements[f"{{{{OBJECTIF_{obj.upper()}_PRIORITE}}}}"] = priority if is_selected else ""

        # Construire le texte des objectifs sélectionnés avec priorités
        objectifs_text = []
        for obj in objectifs:
            if self._get_bool_field(client, f'objectif_{obj}'):
                priority = self._get_field(client, f'objectif_{obj}_priorite')
                label = objectifs_labels.get(obj, obj)
                if priority:
                    objectifs_text.append(f"{label} (Priorité {priority})")
                else:
                    objectifs_text.append(label)
        replacements["{{OBJECTIFS_LISTE}}"] = ", ".join(objectifs_text) if objectifs_text else ""

        # ESG - critères sélectionnés
        esg_criteres = ['gaz_effet_serre', 'biodiversite', 'emissions_eau', 'dechets', 'energie',
                        'normes_internationales', 'egalite_remuneration', 'diversite_genres', 'armes_controversees']
        esg_labels = {
            'gaz_effet_serre': 'Gaz à effet de serre',
            'biodiversite': 'Impact sur la biodiversité',
            'emissions_eau': 'Émissions polluantes dans l\'eau',
            'dechets': 'Génération des déchets dangereux',
            'energie': 'Inefficacité énergétique (immobilier)',
            'normes_internationales': 'Respect des normes internationales (OCDE, Nations unies)',
            'egalite_remuneration': 'Égalité de rémunération (Homme/Femme)',
            'diversite_genres': 'Diversité des genres au sein des conseils',
            'armes_controversees': 'Exposition aux armes controversées',
        }

        for critere in esg_criteres:
            is_selected = self._get_bool_field(client, f'esg_{critere}')
            replacements[f"{{{{ESG_{critere.upper()}_SELECTED}}}}"] = "☑" if is_selected else "☐"

        # Liste des critères ESG sélectionnés
        esg_text = []
        for critere in esg_criteres:
            if self._get_bool_field(client, f'esg_{critere}'):
                esg_text.append(esg_labels.get(critere, critere))
        replacements["{{ESG_CRITERES_LISTE}}"] = ", ".join(esg_text) if esg_text else "Aucun critère sélectionné"

        return replacements

    async def generate_profil_risque_v2(self, client: Client, conseiller: User) -> str:
        """
        Générer le Profil de Risque avec le template DOCX v2
        Document professionnel avec placeholders remplacés
        """
        template_path = os.path.join(self.v2_templates_path, "PROFIL_RISQUE_V2_TEMPLATE.docx")

        if not os.path.exists(template_path):
            # Fallback vers l'ancienne méthode
            return await self.generate_profil_risque(client, conseiller)

        # Charger le template
        doc = Document(template_path)

        # Construire les remplacements spécifiques au Profil de Risque
        replacements = self._build_profil_risque_replacements(client, conseiller)

        # Remplacer les placeholders
        self._replace_placeholders_in_doc(doc, replacements)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("PROFIL_RISQUE", client)
        return self._save_document(doc, filename)

    async def generate_der_v2(self, client: Client, conseiller: User) -> str:
        """
        Générer le DER (Document d'Entrée en Relation) avec le template DOCX v2
        Document professionnel avec mise en page structurée et placeholders remplacés
        """
        template_path = os.path.join(self.v2_templates_path, "DER_V2_TEMPLATE.docx")

        if not os.path.exists(template_path):
            # Fallback vers l'ancienne méthode
            return await self.generate_der(client, conseiller)

        # Charger le template
        doc = Document(template_path)

        # Construire les remplacements spécifiques au DER
        replacements = self._build_der_replacements(client, conseiller)

        # Remplacer les placeholders
        self._replace_placeholders_in_doc(doc, replacements)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("DER", client)
        return self._save_document(doc, filename)

    def _build_der_replacements(self, client: Client, conseiller: User) -> Dict[str, str]:
        """
        Construit le dictionnaire de remplacement pour le DER
        Moins de champs que le QCC, focalisé sur les informations essentielles
        """
        replacements = {}

        # Récupérer la date de remise du DER depuis form_data (contexteMission)
        date_remise_der = None
        if client.form_data:
            contexte_mission = client.form_data.get('contexteMission', {})
            date_remise_der = contexte_mission.get('dateRemiseDER')

        # Informations générales
        replacements["{{DATE_REMISE_DER}}"] = self._format_date(date_remise_der) if date_remise_der else self._format_date(datetime.now())
        replacements["{{DATE_SIGNATURE}}"] = self._format_date(datetime.now())
        replacements["{{LIEU_SIGNATURE}}"] = "Papeete"
        replacements["{{NOMBRE_EXEMPLAIRES}}"] = "2"
        replacements["{{NUMERO_CLIENT}}"] = self._safe(client.numero_client)

        # Nom complet client
        nom_complet_t1 = f"{self._safe(client.t1_prenom)} {self._safe(client.t1_nom, '').upper()}"
        replacements["{{NOM_COMPLET_T1}}"] = nom_complet_t1

        # Titulaire 1
        replacements["{{TITRE_CLIENT}}"] = self._safe(client.t1_civilite)
        replacements["{{T1_EMAIL}}"] = self._safe(client.t1_email)

        # Conseiller
        nom_conseiller = f"{self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}"
        replacements["{{NOM_CONSEILLER}}"] = nom_conseiller
        replacements["{{TITRE_CONSEILLER}}"] = "M." if conseiller.prenom else ""

        return replacements

    async def generate_rto_v2(self, client: Client, conseiller: User) -> str:
        """
        Générer la Convention RTO (Réception et Transmission d'Ordres) avec le template DOCX v2
        Document professionnel avec mise en page structurée et placeholders remplacés
        """
        template_path = os.path.join(self.v2_templates_path, "RTO_V2_TEMPLATE.docx")

        if not os.path.exists(template_path):
            # Fallback vers l'ancienne méthode
            return await self.generate_convention_rto(client, conseiller)

        # Charger le template
        doc = Document(template_path)

        # Construire les remplacements spécifiques au RTO
        replacements = self._build_rto_replacements(client, conseiller)

        # Remplacer les placeholders
        self._replace_placeholders_in_doc(doc, replacements)

        # Sauvegarder avec nouveau format de nom
        filename = self._generate_filename("CONVENTION_RTO", client)
        return self._save_document(doc, filename)

    def _build_rto_replacements(self, client: Client, conseiller: User) -> Dict[str, str]:
        """
        Construit le dictionnaire de remplacement pour le RTO
        Inclut les informations client détaillées et les comptes titres
        """
        replacements = {}

        # Informations générales
        replacements["{{DATE_SIGNATURE}}"] = self._format_date(datetime.now())
        replacements["{{LIEU_SIGNATURE}}"] = "Papeete"
        replacements["{{NOMBRE_EXEMPLAIRES}}"] = "2"
        replacements["{{NUMERO_CLIENT}}"] = self._safe(client.numero_client)

        # Titulaire 1 - Informations complètes
        replacements["{{T1_CIVILITE}}"] = self._safe(client.t1_civilite)
        replacements["{{T1_NOM}}"] = self._safe(client.t1_nom, "").upper()
        replacements["{{T1_PRENOM}}"] = self._safe(client.t1_prenom)
        replacements["{{T1_DATE_NAISSANCE}}"] = self._format_date(client.t1_date_naissance)
        replacements["{{T1_LIEU_NAISSANCE}}"] = self._get_field(client, 't1_lieu_naissance')
        replacements["{{T1_NATIONALITE}}"] = self._get_field(client, 't1_nationalite', 'Française')
        replacements["{{T1_ADRESSE}}"] = self._safe(client.t1_adresse)
        replacements["{{T1_CODE_POSTAL}}"] = self._get_field(client, 't1_code_postal')
        replacements["{{T1_VILLE}}"] = self._get_field(client, 't1_ville')
        replacements["{{T1_PROFESSION}}"] = self._safe(client.t1_profession)

        # Situation familiale
        replacements["{{SITUATION_FAMILIALE}}"] = self._safe(client.situation_familiale)
        replacements["{{REGIME_MATRIMONIAL}}"] = self._get_field(client, 'regime_matrimonial')

        # Conseiller
        nom_conseiller = f"{self._safe(conseiller.prenom)} {self._safe(conseiller.nom, '').upper()}"
        replacements["{{NOM_CONSEILLER}}"] = nom_conseiller

        # Comptes titres (à remplir manuellement ou via form_data)
        for i in range(1, 4):
            replacements[f"{{{{COMPTE_TYPE_{i}}}}}"] = self._get_field(client, f'compte_type_{i}')
            replacements[f"{{{{COMPTE_ETABLISSEMENT_{i}}}}}"] = self._get_field(client, f'compte_etablissement_{i}')
            replacements[f"{{{{COMPTE_NUMERO_{i}}}}}"] = self._get_field(client, f'compte_numero_{i}')

        return replacements
